#!/usr/bin/env python3
"""
Maggie AI Assistant - Main Script
=================================
Entry point for the Maggie AI Assistant application.

This script initializes and starts the Maggie AI Assistant with optimizations
for AMD Ryzen 9 5900X and NVIDIA GeForce RTX 3080 hardware.
It handles command-line arguments, logging configuration, system verification,
and application startup.

The application architecture follows a modular design pattern with the following components:
- Core engine (MaggieAI class) that orchestrates all subsystems
- Event-driven communication using an event bus pattern
- Hardware-aware optimizations for CPU, memory, and GPU resources
- Extension system for adding new capabilities
- Multimodal input/output (speech, text, GUI)

Features
--------
- Voice interaction with wake word detection
- Natural language processing with Mistral 7B LLM
- Text-to-speech synthesis with customizable voices
- Graphical user interface (optional, can run headless)
- Extension system for custom functionality
- Hardware-optimized performance for Ryzen 9 5900X and RTX 3080

Examples
--------
Standard startup:
    $ python main.py

Start with debug logging:
    $ python main.py --debug

Verify system without starting:
    $ python main.py --verify

Create recipe template:
    $ python main.py --create-template

Apply hardware optimizations:
    $ python main.py --optimize

Run without GUI (headless mode):
    $ python main.py --headless

Notes
-----
The application requires Python 3.10.x specifically and will not work with
other Python versions. It is optimized for systems with NVIDIA RTX 3080 GPU
and AMD Ryzen 9 5900X CPU, but will adapt to other hardware configurations.

See config.yaml for detailed configuration options.
"""

# Standard library imports
import os
import argparse
import sys
import platform
import multiprocessing
import time
import yaml
from typing import Dict, Any, Optional, List, Tuple

# Third-party imports
from loguru import logger

# Local imports (updated paths)
# These imports will be available after installation is complete
# from maggie.core import MaggieAI

__all__ = ['main', 'parse_arguments', 'setup_logging', 'verify_system']


# At the top of main.py (after imports)
def error_publisher(message):
    """
    Publish error messages to the event bus.
    
    This function serves as a bridge between the logging system and the event bus,
    allowing error messages to be published as events that other components can
    subscribe to. It is used as a loguru sink to capture ERROR and CRITICAL level
    log messages.
    
    Parameters
    ----------
    message : str or dict
        The formatted error message or error data dictionary containing
        message, source, file, line, and level information
        
    Returns
    -------
    None
        This function does not return a value but publishes an event
        
    Notes
    -----
    This function relies on a global attribute 'event_bus' that must be set
    after the MaggieAI instance is created. The event bus is accessed through
    this attribute to publish the "error_logged" event.
    
    The event_bus attribute is set in the start_maggie function after
    initializing the MaggieAI instance.
    
    Example
    -------
    >>> error_publisher("Critical error in speech recognition")
    # Publishes "error_logged" event with the message
    
    >>> error_publisher({"message": "Init failed", "source": "tts", "level": "ERROR"})
    # Publishes structured error data
    """
    # Access event_bus through the MaggieAI instance
    # This function will be called after the instance is created
    if hasattr(error_publisher, 'event_bus'):
        error_publisher.event_bus.publish("error_logged", message)


def create_event_bus_handler(event_bus):
    """
    Create a loguru handler function that forwards error logs to the event bus.
    
    This factory function creates a custom handler for the loguru logging system
    that captures ERROR and CRITICAL level log messages and publishes them to
    the event bus. This allows for centralized error tracking and enables other
    components (like the GUI) to subscribe to error events.
    
    Parameters
    ----------
    event_bus : EventBus
        The event bus instance to publish error messages to. This should be
        the same event bus that other components use for communication.
        
    Returns
    -------
    callable
        A function that can be used as a loguru handler with the signature
        handler(record) -> None
        
    Notes
    -----
    The returned handler filters log records to only process ERROR and CRITICAL
    levels. It extracts relevant information from the log record and publishes
    a structured data dictionary to the event bus.
    
    The published event has the topic "error_logged" and includes metadata
    about the error source, file, line number, and severity level.
    
    Example
    -------
    >>> from loguru import logger
    >>> from maggie.utils.event_bus import EventBus
    >>> event_bus = EventBus()
    >>> error_handler = create_event_bus_handler(event_bus)
    >>> logger.add(error_handler, level="ERROR")
    >>> logger.error("Database connection failed")  # This will trigger an event
    """
    def handler(record):
        """
        Process log records and publish errors to the event bus.
        
        This inner function serves as the actual loguru handler. It filters
        records by log level and transforms them into structured event data.
        
        Parameters
        ----------
        record : dict
            The loguru record dictionary containing details about the log event
            including message, level, file, line, etc.
            
        Returns
        -------
        None
            This handler does not return a value but may publish an event
            
        Notes
        -----
        Only ERROR and CRITICAL level messages are processed and published
        to the event bus as "error_logged" events.
        """
        if record["level"].name == "ERROR" or record["level"].name == "CRITICAL":
            error_data = {
                "message": record["message"],
                "source": record["name"],
                "file": record["file"].name if hasattr(record["file"], "name") else str(record["file"]),
                "line": record["line"],
                "level": record["level"].name
            }
            event_bus.publish("error_logged", error_data)
    
    return handler


def parse_arguments() -> argparse.Namespace:
    """
    Parse command line arguments for Maggie AI Assistant.
    
    This function sets up the argument parser with all available command line options
    for configuring and running the Maggie AI Assistant. It defines default values,
    help text, and argument types.
    
    Parameters
    ----------
    None
    
    Returns
    -------
    argparse.Namespace
        Parsed command line arguments with the following fields:
        - config : str
            Path to configuration file (default: "config.yaml")
        - debug : bool
            Boolean flag for enabling debug-level logging (default: False)
        - verify : bool
            Boolean flag for running system verification without starting (default: False)
        - create_template : bool
            Boolean flag for creating the recipe template file (default: False)
        - optimize : bool
            Boolean flag for applying hardware optimizations (default: False)
        - headless : bool
            Boolean flag for running without GUI in headless mode (default: False)
    
    Notes
    -----
    The arguments are processed in order of precedence. For example, if both
    --verify and --create-template are specified, verification will be performed
    first, followed by template creation, and then the program will exit.
    
    The --config option allows specifying a non-default configuration file path.
    This is useful for testing different configurations without modifying the
    main config.yaml file.
    
    Example
    -------
    >>> args = parse_arguments()
    >>> if args.debug:
    ...     # Set up debug logging
    >>> if args.verify:
    ...     # Run system verification
    """
    parser = argparse.ArgumentParser(
        description="Maggie AI Assistant",
        formatter_class=argparse.ArgumentDefaultsHelpFormatter
    )
    parser.add_argument(
        "--config", 
        type=str, 
        default="config.yaml",
        help="Path to configuration file"
    )
    parser.add_argument(
        "--debug", 
        action="store_true",
        help="Enable debug logging"
    )
    parser.add_argument(
        "--verify", 
        action="store_true",
        help="Verify system configuration without starting the assistant"
    )
    parser.add_argument(
        "--create-template",
        action="store_true",
        help="Create the recipe template file if it doesn't exist"
    )
    parser.add_argument(
        "--optimize",
        action="store_true",
        help="Optimize configuration for detected hardware"
    )
    parser.add_argument(
        "--headless",
        action="store_true",
        help="Run in headless mode without GUI"
    )
    return parser.parse_args()


def setup_logging(debug: bool = False) -> None:
    """
    Set up logging configuration using loguru.
    
    This function configures the loguru logger with console and file handlers
    using appropriate formatting and log levels. It ensures the logs directory
    exists and sets up log rotation to prevent excessive disk usage.
    
    Parameters
    ----------
    debug : bool, optional
        Enable debug logging if True, showing more detailed logs in both
        console and file output, by default False
        
    Returns
    -------
    None
        This function doesn't return a value but configures the global logger
        
    Notes
    -----
    This function configures two logging destinations:
    1. Console output with colored formatting and level-based highlighting
    2. File output with rotation (10 MB max size) and retention (1 week)
    
    The log format includes:
    - Timestamp in YYYY-MM-DD HH:mm:ss format
    - Log level (padded to 8 characters)
    - Source location (module:function:line)
    - Log message with level-specific coloring (console only)
    
    After configuring logging, it calls log_system_info() to record basic
    system information at startup.
    
    Example
    -------
    >>> setup_logging(debug=True)  # Enable verbose debug logging
    >>> setup_logging()  # Use standard INFO level logging
    """
    # Ensure logs directory exists
    os.makedirs("logs", exist_ok=True)
    
    # Configure loguru
    log_level = "DEBUG" if debug else "INFO"
    logger.configure(
        handlers=[
            {"sink": sys.stdout, "level": log_level, "format": "<green>{time:YYYY-MM-DD HH:mm:ss}</green> | <level>{level: <8}</level> | <cyan>{name}</cyan>:<cyan>{function}</cyan>:<cyan>{line}</cyan> - <level>{message}</level>"},
            {"sink": "logs/maggie.log", "rotation": "10 MB", "retention": "1 week", "format": "{time:YYYY-MM-DD HH:mm:ss} | {level: <8} | {name}:{function}:{line} - {message}"}
        ]
    )
    
    # Log system info
    log_system_info()


def log_system_info() -> None:
    """
    Log detailed information about the system hardware and environment.
    
    This function collects and logs comprehensive information about the system,
    including operating system, Python version, CPU specifications, RAM capacity,
    and GPU details. The information is logged at INFO level to help with
    diagnostics, troubleshooting, and performance optimization.
    
    Parameters
    ----------
    None
    
    Returns
    -------
    None
        This function doesn't return a value but logs information
        
    Notes
    -----
    The function attempts to gather as much information as possible, but
    gracefully handles cases where certain information can't be obtained
    (e.g., if a required package is missing).
    
    Information collected includes:
    - Operating system name and version
    - Python version
    - CPU model, core count, and thread count
    - Total RAM capacity
    - GPU model, VRAM capacity, and CUDA information (if available)
    
    Special detection is included for AMD Ryzen 9 5900X CPU and NVIDIA RTX 3080 GPU
    to apply specific optimizations for these components.
    
    Dependencies:
    - psutil: For CPU and RAM information
    - torch: For GPU information and CUDA detection
    
    Example
    -------
    >>> log_system_info()
    # INFO: System: Windows 11 Pro 10.0.22621
    # INFO: Python: 3.10.8
    # INFO: CPU: AMD Ryzen 9 5900X 12-Core Processor
    # INFO: CPU Cores: 12 physical, 24 logical
    # INFO: RAM: 32.00 GB
    # INFO: GPU 0: NVIDIA GeForce RTX 3080
    # INFO: GPU Memory: 10.00 GB
    # INFO: RTX 3080 detected - Tensor Cores available
    """
    try:
        # System and Python info
        logger.info(f"System: {platform.system()} {platform.release()}")
        logger.info(f"Python: {platform.python_version()}")
        
        # Log CPU info
        log_cpu_info()
        
        # Log RAM info
        log_ram_info()
        
        # Log GPU info
        log_gpu_info()
            
    except ImportError as e:
        logger.warning(f"System info modules not available: {e}")
    except Exception as e:
        logger.warning(f"Error gathering system information: {e}")


def log_cpu_info() -> None:
    """
    Log detailed information about the CPU.
    
    This function gathers and logs information about the system's CPU,
    including model name, physical cores, and logical cores (threads).
    It uses the psutil package to collect hardware information.
    
    Parameters
    ----------
    None
    
    Returns
    -------
    None
        This function doesn't return a value but logs CPU information
        
    Notes
    -----
    The function attempts to gather the following CPU information:
    - CPU model/name (from platform.processor())
    - Physical core count (excluding hyperthreading/SMT)
    - Logical core count (including hyperthreading/SMT)
    
    If the psutil package is not available, the function logs a warning
    and fails gracefully without crashing the application.
    
    The CPU information is particularly important for optimizing thread usage
    and parallel processing capabilities in the application. The Maggie AI
    Assistant has specific optimizations for the AMD Ryzen 9 5900X CPU.
    
    Example
    -------
    >>> log_cpu_info()
    # INFO: CPU: AMD Ryzen 9 5900X 12-Core Processor
    # INFO: CPU Cores: 12 physical, 24 logical
    """
    try:
        import psutil
        
        cpu_info = platform.processor()
        cpu_cores = psutil.cpu_count(logical=False)
        cpu_threads = psutil.cpu_count(logical=True)
        
        logger.info(f"CPU: {cpu_info}")
        logger.info(f"CPU Cores: {cpu_cores} physical, {cpu_threads} logical")
    except ImportError:
        logger.warning("psutil not available for CPU information")
    except Exception as e:
        logger.warning(f"Error getting CPU information: {e}")


def log_ram_info() -> None:
    """
    Log information about system memory (RAM).
    
    This function collects and logs information about the system's physical
    memory (RAM) using the psutil package. It provides insight into the total
    amount of RAM available to the application.
    
    Parameters
    ----------
    None
    
    Returns
    -------
    None
        This function doesn't return a value but logs memory information
        
    Notes
    -----
    The function logs the total RAM capacity in gigabytes (GB) with 2 decimal
    precision. This information is valuable for:
    
    1. Determining if the system meets minimum memory requirements (16GB)
    2. Applying optimal configuration for high-memory systems (32GB+)
    3. Setting appropriate memory allocation limits for AI models
    
    Maggie has specific optimizations for systems with 32GB of RAM, allocating
    memory thresholds appropriately for model loading/unloading.
    
    If the psutil package is not available, the function logs a warning and
    continues execution to maintain application stability.
    
    Example
    -------
    >>> log_ram_info()
    # INFO: RAM: 32.00 GB
    """
    try:
        import psutil
        
        ram_gb = psutil.virtual_memory().total / (1024 ** 3)
        logger.info(f"RAM: {ram_gb:.2f} GB")
    except ImportError:
        logger.warning("psutil not available for RAM information")
    except Exception as e:
        logger.warning(f"Error getting RAM information: {e}")


def log_gpu_info() -> None:
    """
    Log detailed information about GPU hardware and CUDA capabilities.
    
    This function attempts to detect and log information about available GPUs
    using PyTorch's CUDA interface. It logs GPU models, available VRAM,
    and detects specific optimizations for the NVIDIA RTX 3080.
    
    Parameters
    ----------
    None
    
    Returns
    -------
    None
        This function doesn't return a value but logs GPU information
        
    Notes
    -----
    The function collects the following GPU information when available:
    - GPU detection status and CUDA availability
    - Each GPU's name/model
    - VRAM capacity in gigabytes (GB)
    - Special capabilities for RTX 3080 (Tensor Cores, Ampere architecture)
    
    The information is valuable for:
    1. Determining if GPU acceleration is available
    2. Applying architecture-specific optimizations (e.g., for RTX 3080)
    3. Making memory allocation decisions based on available VRAM
    
    Maggie has specific optimizations for the NVIDIA RTX 3080 GPU with 10GB VRAM,
    including tensor core utilization and half-precision (FP16) optimizations.
    
    If PyTorch is not installed or CUDA is not available, the function logs
    appropriate warnings without crashing the application.
    
    Example
    -------
    >>> log_gpu_info()
    # INFO: GPU 0: NVIDIA GeForce RTX 3080
    # INFO: GPU Memory: 10.00 GB
    # INFO: RTX 3080 detected - Tensor Cores available
    # INFO: Optimizing for 10GB VRAM and Ampere architecture
    """
    try:
        import torch
        if torch.cuda.is_available():
            for i in range(torch.cuda.device_count()):
                gpu_name = torch.cuda.get_device_name(i)
                memory_gb = torch.cuda.get_device_properties(i).total_memory / (1024 ** 3)
                logger.info(f"GPU {i}: {gpu_name}")
                logger.info(f"GPU Memory: {memory_gb:.2f} GB")
                
                # Log RTX 3080 specific capabilities
                if "3080" in gpu_name:
                    logger.info(f"RTX 3080 detected - Tensor Cores available")
                    logger.info(f"Optimizing for 10GB VRAM and Ampere architecture")
        else:
            logger.warning("CUDA not available, GPU acceleration disabled")
    except ImportError:
        logger.warning("PyTorch not installed, GPU detection skipped")
    except Exception as e:
        logger.warning(f"Error detecting GPU information: {e}")


def verify_system() -> bool:
    """
    Verify the system meets all requirements for running Maggie AI Assistant.
    
    This function performs a comprehensive verification of system capabilities
    and requirements before starting the application. It checks Python version,
    hardware capabilities, directory structure, dependencies, and memory
    configuration to ensure optimal operation.
    
    Parameters
    ----------
    None
    
    Returns
    -------
    bool
        True if system meets all critical requirements, False if critical
        issues are found that would prevent proper operation
        
    Notes
    -----
    The function performs the following verification steps:
    
    1. Python Version Check:
       - Requires exactly Python 3.10.x
       - Other versions are incompatible due to dependency requirements
    
    2. GPU Verification:
       - Detects CUDA availability for GPU acceleration
       - Checks VRAM capacity (minimum 4GB, recommended 8GB+)
       - Performs specific verification for RTX 3080 (10GB VRAM)
       - Tests basic CUDA operations to ensure functionality
    
    3. Directory Structure Check:
       - Verifies all required directories exist
       - Creates missing directories if needed
    
    4. Dependency Checks:
       - Verifies critical dependencies are installed
       - Checks for PyAudio which is often problematic
    
    5. Memory Configuration Check:
       - Verifies minimum RAM (8GB required, 16GB recommended)
       - Checks available memory for operation
    
    The function returns False only if critical issues are found that would
    prevent the application from running properly. Warnings about non-critical
    issues (like suboptimal hardware) are logged but don't cause verification
    to fail.
    
    Example
    -------
    >>> if verify_system():
    ...     start_application()
    ... else:
    ...     display_error_and_exit()
    """
    verification_issues = []
    logger.info("Verifying system configuration...")
    
    # Check Python version - require exactly 3.10.x
    python_version_valid = check_python_version()
    if not python_version_valid:
        verification_issues.append("Incompatible Python version")
    
    # Enhanced GPU verification with VRAM checks for RTX 3080
    gpu_valid, gpu_warnings = check_gpu_compatibility()
    if not gpu_valid:
        verification_issues.append("GPU compatibility issues detected")
    
    # Check for required directories and create if necessary
    dirs_valid = check_required_directories()
    if not dirs_valid:
        verification_issues.append("Directory issues detected")
    
    # Check for required dependencies
    deps_valid = check_dependencies()
    if not deps_valid:
        verification_issues.append("Missing dependencies")
    
    # Check for sufficient memory
    memory_valid = check_memory_configuration()
    if not memory_valid:
        verification_issues.append("Insufficient memory configuration")
    
    # Final verification result
    if verification_issues:
        logger.error("System verification failed with the following issues:")
        for issue in verification_issues:
            logger.error(f"  - {issue}")
        return False
    
    logger.info("System verification completed successfully")
    return True


def check_gpu_compatibility() -> Tuple[bool, List[str]]:
    """
    Check if GPU is compatible and properly configured for Maggie.
    
    This function performs comprehensive GPU compatibility checks including
    CUDA availability, VRAM capacity, compute capability, and functional testing.
    It has specific optimizations for the NVIDIA RTX 3080 GPU.
    
    Parameters
    ----------
    None
    
    Returns
    -------
    Tuple[bool, List[str]]
        A tuple containing:
        - Boolean indicating if GPU is compatible (True) or has critical issues (False)
        - List of warning messages for non-critical issues
        
    Notes
    -----
    The function performs these detailed checks:
    
    1. CUDA Availability:
       - Verifies PyTorch can detect and use CUDA
       - Logs CUDA version and runtime information
    
    2. GPU Hardware Detection:
       - Identifies GPU model(s)
       - Measures total and available VRAM
    
    3. RTX 3080 Specific Checks (if detected):
       - Verifies expected 10GB VRAM capacity
       - Checks compute capability (should be 8.6)
       - Verifies CUDA version compatibility (11.x recommended)
       - Ensures CUDA runtime and compiled versions match
       - Checks for sufficient available VRAM (8GB+ recommended)
    
    4. Functional Testing:
       - Performs a basic CUDA operation (matrix multiplication)
       - Ensures the operation completes successfully
       - Cleans up memory after the test
    
    The function can run in fallback modes:
    - Returns True with warnings for non-RTX 3080 GPUs
    - Returns True with warnings for CPU-only operation if no GPU is available
    - Only returns False for critical issues that would prevent operation
    
    Example
    -------
    >>> compatible, warnings = check_gpu_compatibility()
    >>> if not compatible:
    ...     print("Critical GPU issues found, cannot continue")
    ... elif warnings:
    ...     print(f"GPU will work but with limitations: {warnings}")
    ... else:
    ...     print("GPU configuration is optimal")
    """
    warnings = []
    
    try:
        import torch
        if torch.cuda.is_available():
            # Get GPU information
            device_name = torch.cuda.get_device_name(0)
            logger.info(f"CUDA available: {torch.version.cuda}")
            logger.info(f"GPU detected: {device_name}")
            
            # Get VRAM information
            gpu_memory = torch.cuda.get_device_properties(0).total_memory / (1024**3)
            logger.info(f"VRAM: {gpu_memory:.2f}GB")
            
            # Check for RTX 3080 specifically
            is_rtx_3080 = "3080" in device_name
            if is_rtx_3080:
                logger.info("RTX 3080 detected - optimal hardware configuration")
                
                # Check VRAM size for RTX 3080 (should be around 10GB)
                if gpu_memory < 9.5:
                    warnings.append(f"RTX 3080 VRAM ({gpu_memory:.1f}GB) is less than expected 10GB")
                    logger.warning(f"RTX 3080 VRAM ({gpu_memory:.1f}GB) is less than expected 10GB")
                
                # Get compute capability (should be 8.6 for RTX 3080)
                compute_capability = torch.cuda.get_device_capability(0)
                cc_version = f"{compute_capability[0]}.{compute_capability[1]}"
                if cc_version != "8.6":
                    warnings.append(f"RTX 3080 compute capability ({cc_version}) is not 8.6")
                    logger.warning(f"RTX 3080 compute capability ({cc_version}) is not 8.6")
                
                # Check CUDA version (11.x recommended for RTX 3080)
                cuda_version = torch.version.cuda
                if not cuda_version.startswith("11."):
                    warnings.append(f"CUDA version {cuda_version} - version 11.x recommended for RTX 3080")
                    logger.warning(f"CUDA version {cuda_version} - version 11.x recommended for RTX 3080")
                
                # Validate if CUDA compiled version matches runtime version
                if hasattr(torch.version, 'cuda_compiled_version') and torch.version.cuda != torch.version.cuda_compiled_version:
                    warnings.append(f"CUDA runtime version ({torch.version.cuda}) differs from compiled version ({torch.version.cuda_compiled_version})")
                    logger.warning(f"CUDA runtime version ({torch.version.cuda}) differs from compiled version ({torch.version.cuda_compiled_version})")
                
                # Check available VRAM
                available_vram = (torch.cuda.get_device_properties(0).total_memory - torch.cuda.memory_allocated() - torch.cuda.memory_reserved()) / (1024**3)
                if available_vram < 8.0:
                    warnings.append(f"Only {available_vram:.1f}GB VRAM available - less than 8GB may cause issues with large models")
                    logger.warning(f"Only {available_vram:.1f}GB VRAM available - less than 8GB may cause issues with large models")
            else:
                if gpu_memory < 8:
                    warnings.append(f"GPU memory ({gpu_memory:.1f}GB) is less than recommended 8GB")
                    logger.warning(f"GPU memory ({gpu_memory:.1f}GB) is less than recommended 8GB")
                # Not a critical issue, can run with lower specs
            
            # Test CUDA operations to ensure functional GPU
            try:
                test_tensor = torch.ones(1000, 1000, device='cuda')
                test_result = torch.matmul(test_tensor, test_tensor)
                del test_tensor, test_result
                torch.cuda.empty_cache()
                logger.info("CUDA operations test successful")
            except Exception as e:
                warnings.append(f"CUDA operations test failed: {e}")
                logger.error(f"CUDA operations test failed: {e}")
                return False, warnings
                
            return True, warnings
        else:
            logger.warning("CUDA not available, GPU acceleration disabled")
            logger.warning("Performance may be significantly reduced without GPU acceleration")
            # Not a critical issue, can run without GPU
            warnings.append("CUDA not available - running in CPU-only mode will be slow")
            return True, warnings
    except ImportError:
        logger.error("PyTorch not installed, GPU acceleration unavailable")
        logger.error("Install PyTorch with: pip install torch==2.0.1+cu118 --extra-index-url https://download.pytorch.org/whl/cu118")
        # Not a critical issue, can run without GPU
        warnings.append("PyTorch not available - GPU acceleration disabled")
        return True, warnings
    except Exception as e:
        logger.error(f"Error checking GPU compatibility: {e}")
        # Not a critical issue
        warnings.append(f"Error checking GPU: {e}")
        return True, warnings


def check_memory_configuration() -> bool:
    """
    Check if system memory configuration is sufficient for Maggie.
    
    This function verifies that the system has enough physical memory (RAM)
    to run Maggie AI Assistant effectively. It checks total capacity and
    available memory against minimum and recommended thresholds.
    
    Parameters
    ----------
    None
    
    Returns
    -------
    bool
        True if memory configuration is sufficient, False if critical issues found
        that would prevent proper operation
    
    Notes
    -----
    The function performs these verification steps:
    
    1. Total RAM Check:
       - Critical minimum: 8GB (application will not start with less)
       - Recommended minimum: 16GB (will work but with limitations)
       - Optimal: 32GB (full functionality with parallel processing)
    
    2. Available Memory Check:
       - Critical minimum: 2GB available (prevents immediate memory errors)
       - Warns if available memory is low, suggesting closing other applications
    
    Memory requirements are based on:
    - Large language model (LLM) memory usage (~4-7GB)
    - Speech processing memory requirements (~1-2GB)
    - Cache and buffer allocations
    - Operating system overhead
    
    The function uses psutil to gather memory information when available,
    but will continue execution (returning True) if psutil is not installed,
    relying on runtime memory errors to catch critical issues.
    
    Example
    -------
    >>> if check_memory_configuration():
    ...     proceed_with_startup()
    ... else:
    ...     show_error_message("Insufficient system memory")
    ...     exit(1)
    """
    try:
        import psutil
        
        # Check total physical memory
        memory = psutil.virtual_memory()
        total_gb = memory.total / (1024**3)
        available_gb = memory.available / (1024**3)
        
        logger.info(f"Total RAM: {total_gb:.2f}GB, Available: {available_gb:.2f}GB")
        
        # Critical check: Require at least 8GB RAM to function at all
        if total_gb < 8:
            logger.error(f"Critical: Insufficient RAM: {total_gb:.2f}GB (minimum 8GB required)")
            return False
            
        # Warning level checks
        if total_gb < 16:
            logger.warning(f"Insufficient RAM: {total_gb:.2f}GB (minimum 16GB recommended)")
            # Continue but log warning - not a critical failure
        elif total_gb < 32:
            logger.warning(f"RAM: {total_gb:.2f}GB (32GB recommended for optimal performance)")
        else:
            logger.info(f"RAM: {total_gb:.2f}GB (optimal)")
        
        # Check available memory - critical if extremely low
        if available_gb < 2:
            logger.error(f"Critical: Only {available_gb:.2f}GB RAM available. Close other applications.")
            return False
            
        return True
        
    except ImportError:
        logger.warning("psutil not available for memory checks")
        logger.warning("Install psutil for comprehensive memory verification")
        return True  # Continue anyway - not critical
    except Exception as e:
        logger.error(f"Error checking memory configuration: {e}")
        logger.warning("Memory check failed, proceeding with caution")
        return True  # Continue anyway - not critical


def check_python_version() -> bool:
    """
    Check if Python version is compatible with Maggie AI Assistant.
    
    This function verifies that the Python interpreter running the application
    is exactly version 3.10.x, which is the only supported version. Other versions
    (including 3.9.x and 3.11.x) are not compatible due to dependency constraints.
    
    Parameters
    ----------
    None
    
    Returns
    -------
    bool
        True if Python version is 3.10.x, False otherwise
    
    Notes
    -----
    Maggie AI Assistant requires Python 3.10.x specifically because:
    
    1. Some dependencies (like PyTorch with CUDA) have version-specific wheels
    2. Some libraries use Python 3.10-specific features
    3. Certain dependencies have not been updated for Python 3.11+
    4. Python 3.9 lacks some features used by the codebase
    
    When an incompatible Python version is detected, the function logs detailed
    error messages explaining the requirement and pointing to Python 3.10
    installation resources.
    
    Example
    -------
    >>> if check_python_version():
    ...     print("Python version compatible")
    ... else:
    ...     print("Please install Python 3.10.x")
    ...     sys.exit(1)
    """
    python_version = platform.python_version_tuple()
    if int(python_version[0]) != 3 or int(python_version[1]) != 10:
        error_msg = f"Unsupported Python version: {platform.python_version()}"
        logger.error(error_msg)
        logger.error("Maggie requires Python 3.10.x specifically. Other versions are not compatible.")
        logger.error("Please install Python 3.10 and try again.")
        return False
    else:
        logger.info(f"Python version {platform.python_version()} is compatible")
        return True


def check_cuda_availability() -> bool:
    """
    Check if CUDA is available and compatible for GPU acceleration.
    
    This function verifies CUDA availability through PyTorch and logs
    information about detected GPUs. It performs basic compatibility checking
    but is less comprehensive than check_gpu_compatibility().
    
    Parameters
    ----------
    None
    
    Returns
    -------
    bool
        True if CUDA is available and compatible OR if CPU-only operation
        is acceptable; False only on critical issues that would prevent
        even basic operation
    
    Notes
    -----
    This function serves as a lightweight CUDA check compared to the more
    comprehensive check_gpu_compatibility() function. It:
    
    1. Attempts to import PyTorch and check CUDA availability
    2. Gets basic information about detected GPU(s)
    3. Checks for RTX 3080 specifically for optimizations
    4. Verifies VRAM capacity against minimum requirements
    
    Unlike check_gpu_compatibility(), this function:
    - Does not perform functional testing of CUDA operations
    - Returns True for both GPU and CPU-only configurations
    - Only returns False for critical issues, not for suboptimal setups
    - Provides less detailed diagnostic information
    
    This function is used in contexts where a quick check is needed
    without the full verification suite, or as a fallback when the main
    verification function encounters issues.
    
    Example
    -------
    >>> if check_cuda_availability():
    ...     # CUDA might be available or CPU-only is acceptable
    ...     proceed_with_configuration()
    ... else:
    ...     # Critical issue detected
    ...     show_error_and_exit()
    """
    try:
        import torch
        if torch.cuda.is_available():
            logger.info(f"CUDA available: {torch.version.cuda}")
            logger.info(f"GPU detected: {torch.cuda.get_device_name(0)}")
            
            # Check for RTX 3080 specifically
            is_rtx_3080 = "3080" in torch.cuda.get_device_name(0)
            if is_rtx_3080:
                logger.info("RTX 3080 detected - optimal hardware configuration")
            else:
                gpu_memory = torch.cuda.get_device_properties(0).total_memory / (1024**3)
                if gpu_memory < 8:
                    logger.warning(f"GPU memory ({gpu_memory:.1f}GB) is less than recommended 10GB")
                    logger.warning("Performance may be degraded")
                    # Not a critical issue
        else:
            logger.warning("CUDA not available, GPU acceleration disabled")
            logger.warning("Performance may be significantly reduced without GPU acceleration")
            # Not a critical issue
        return True
    except ImportError:
        logger.error("PyTorch not installed, GPU acceleration unavailable")
        logger.error("Install PyTorch with: pip install torch")
        # Not a critical issue, can run without GPU
        return True
    except Exception as e:
        logger.error(f"Error checking CUDA: {e}")
        # Not a critical issue
        return True


def check_required_directories() -> bool:
    """
    Check if required directories exist and create them if necessary.
    
    This function verifies that all directories needed by Maggie AI Assistant
    exist, and attempts to create any missing directories. It ensures the
    filesystem structure is ready for application operation.
    
    Parameters
    ----------
    None
    
    Returns
    -------
    bool
        True if all directories exist or were successfully created,
        False if directory creation failed
    
    Notes
    -----
    The function checks and creates (if necessary) these directories:
    - models/ - For storing downloaded AI models
      - models/tts/ - Text-to-speech voice models
    - logs/ - Application logs with rotation
    - recipes/ - Output directory for recipe extension
    - templates/ - Template files for extensions
    
    Directory creation failures are usually caused by:
    - Permission issues (insufficient write access)
    - Disk space limitations
    - Path length limitations (on Windows)
    - Filesystem corruption
    
    Each directory creation is logged for debugging purposes,
    and detailed error information is provided if creation fails.
    
    Example
    -------
    >>> if check_required_directories():
    ...     print("Directory structure is ready")
    ... else:
    ...     print("Failed to create required directories")
    ...     sys.exit(1)
    """
    required_dirs = ["models", "models/tts", "logs", "recipes", "templates"]
    for directory in required_dirs:
        if not os.path.exists(directory):
            try:
                os.makedirs(directory, exist_ok=True)
                logger.info(f"Created directory: {directory}")
            except Exception as e:
                error_msg = f"Failed to create directory {directory}: {e}"
                logger.error(error_msg)
                return False
    return True


def check_dependencies() -> bool:
    """
    Check if required Python package dependencies are installed.
    
    This function verifies that all critical Python package dependencies
    required by Maggie AI Assistant are properly installed and importable.
    It performs specific checks for commonly problematic packages like PyAudio.
    
    Parameters
    ----------
    None
    
    Returns
    -------
    bool
        True if all critical dependencies are installed and importable,
        False if any critical dependency is missing
    
    Notes
    -----
    The function checks these critical dependencies:
    
    1. PyAudio - Audio input capture (checked separately due to common issues)
    2. pvporcupine - Wake word detection
    3. faster_whisper - Speech-to-text conversion
    4. ctransformers - LLM inference
    5. transitions - State machine for application flow
    6. docx - Document generation for recipe extension
    7. PySide6 - GUI framework
    
    When missing dependencies are detected, the function:
    1. Logs detailed error messages
    2. Provides installation instructions
    3. Returns False to indicate verification failure
    
    These dependency checks help prevent runtime errors by ensuring all
    required packages are available before attempting to use them.
    
    Example
    -------
    >>> if check_dependencies():
    ...     print("All dependencies are installed")
    ... else:
    ...     print("Missing dependencies, please run: pip install -r requirements.txt")
    ...     sys.exit(1)
    """
    # Check for PyAudio (common source of issues)
    pyaudio_installed = check_pyaudio()
    
    # Check for other critical dependencies
    critical_deps = [
        "pvporcupine", "faster_whisper", "ctransformers", 
        "transitions", "docx", "PySide6"
    ]
    missing_deps = []
    
    for dep in critical_deps:
        try:
            module_name = dep.replace("-", "_")
            if dep == "PySide6":
                module_name = "PySide6.QtCore"
            __import__(module_name)
        except ImportError:
            missing_deps.append(dep)
    
    if missing_deps:
        deps_str = ", ".join(missing_deps)
        logger.error(f"Missing critical dependencies: {deps_str}")
        logger.error("Install with: pip install -r requirements.txt")
        return False
    
    return True


def check_pyaudio() -> bool:
    """
    Check if PyAudio is installed and importable.
    
    This function specifically checks for the PyAudio package, which is
    commonly problematic to install due to its dependency on PortAudio
    and platform-specific build requirements.
    
    Parameters
    ----------
    None
    
    Returns
    -------
    bool
        True if PyAudio is successfully imported, False otherwise
    
    Notes
    -----
    PyAudio is required for:
    - Capturing audio input from microphones
    - Wake word detection
    - Voice command recognition
    
    Common installation issues:
    - Windows: Missing C++ build tools or improper wheel
    - Linux: Missing portaudio19-dev system package
    - macOS: Missing PortAudio via Homebrew
    
    When PyAudio is not found, the function logs platform-specific
    installation instructions to guide users through proper installation.
    This helps prevent cryptic import errors during runtime.
    
    Example
    -------
    >>> if check_pyaudio():
    ...     print("PyAudio found, audio input will be available")
    ... else:
    ...     print("PyAudio not found, voice features will be disabled")
    """
    try:
        import pyaudio
        logger.info("PyAudio found")
        return True
    except ImportError:
        error_msg = "PyAudio not installed"
        logger.error(error_msg)
        if platform.system() == "Windows":
            logger.error("On Windows, install with: pip install PyAudio")
        else:
            logger.error("On Linux, first install portaudio19-dev, then install pyaudio")
        return False


def create_recipe_template() -> bool:
    """
    Create the recipe template file if it doesn't exist.
    
    This function generates a Microsoft Word (.docx) template document for
    use with the recipe creator extension. The template provides a standardized
    structure for recipes with sections for ingredients, instructions, etc.
    
    Parameters
    ----------
    None
    
    Returns
    -------
    bool
        True if template created successfully or already exists,
        False if template creation failed
        
    Notes
    -----
    The function creates a Word document template with the following structure:
    
    1. Recipe Name (Title)
    2. Recipe Information (Table)
       - Preparation Time
       - Cooking Time
       - Servings
    3. Ingredients (Bulleted List)
    4. Instructions (Numbered Steps)
    5. Notes (Free Text)
    6. Nutrition Information (Key-Value List)
    
    The template is saved to "templates/recipe_template.docx" and is used by
    the recipe_creator extension to generate properly formatted recipe documents.
    
    The function requires python-docx to be installed. If missing, it logs an
    appropriate error message but doesn't crash the application.
    
    The template is only created if it doesn't already exist, allowing users
    to customize it without overwriting their changes on restart.
    
    Example
    -------
    >>> if create_recipe_template():
    ...     print("Recipe template ready at templates/recipe_template.docx")
    ... else:
    ...     print("Could not create recipe template")
    """
    template_path = "templates/recipe_template.docx"
    
    if os.path.exists(template_path):
        logger.info(f"Recipe template already exists at {template_path}")
        return True
    
    try:
        from docx import Document
        
        os.makedirs("templates", exist_ok=True)
        
        doc = Document()
        
        # Create a more detailed template
        doc.add_heading("Recipe Name", level=1)
        
        # Add metadata section
        doc.add_heading("Recipe Information", level=2)
        info_table = doc.add_table(rows=3, cols=2)
        info_table.style = 'Table Grid'
        info_table.cell(0, 0).text = "Preparation Time"
        info_table.cell(0, 1).text = "00 minutes"
        info_table.cell(1, 0).text = "Cooking Time"
        info_table.cell(1, 1).text = "00 minutes"
        info_table.cell(2, 0).text = "Servings"
        info_table.cell(2, 1).text = "0 servings"
        
        # Add ingredients section with better formatting
        doc.add_heading("Ingredients", level=2)
        doc.add_paragraph("• Ingredient 1", style='ListBullet')
        doc.add_paragraph("• Ingredient 2", style='ListBullet')
        doc.add_paragraph("• Ingredient 3", style='ListBullet')
        
        # Add steps section with numbered steps
        doc.add_heading("Instructions", level=2)
        doc.add_paragraph("1. Step 1", style='ListNumber')
        doc.add_paragraph("2. Step 2", style='ListNumber')
        doc.add_paragraph("3. Step 3", style='ListNumber')
        
        # Add notes section
        doc.add_heading("Notes", level=2)
        doc.add_paragraph("Add any additional notes, tips, or variations here.")
        
        # Add nutrition section if available
        doc.add_heading("Nutrition Information (per serving)", level=2)
        doc.add_paragraph("Calories: 000")
        doc.add_paragraph("Protein: 00g")
        doc.add_paragraph("Carbohydrates: 00g")
        doc.add_paragraph("Fat: 00g")
        
        doc.save(template_path)
        logger.info(f"Created recipe template at {template_path}")
        return True
    except ImportError:
        logger.error("python-docx not installed, cannot create template")
        return False
    except Exception as e:
        logger.error(f"Failed to create recipe template: {e}")
        return False


def optimize_system() -> bool:
    """
    Optimize system settings for best Maggie AI Assistant performance.
    
    This function applies various system-level optimizations to improve
    performance of the Maggie AI Assistant. It applies platform-specific
    optimizations for Windows and Linux, as well as common optimizations
    for all platforms.
    
    Parameters
    ----------
    None
    
    Returns
    -------
    bool
        True if optimizations succeeded, False if critical optimizations failed
        
    Notes
    -----
    The function applies these platform-specific optimizations:
    
    For Windows:
    - Sets process priority to HIGH_PRIORITY_CLASS
    - Configures CPU affinity for optimized core usage
    - For Ryzen 9 5900X: Optimizes for the first 8 physical cores (16 threads)
    
    For Linux:
    - Sets CPU governor to "performance" mode (requires root)
    - Sets process nice level to -10 (higher priority)
    - Configures process scheduling parameters
    
    Common optimizations for all platforms:
    - Sets process name for better identification
    - Configures Python-specific optimizations
    - Sets int_max_str_digits to limit memory usage in certain operations
    
    These optimizations help ensure:
    - Responsive wake word detection
    - Faster LLM inference
    - Improved speech processing performance
    - Better multitasking capabilities
    
    Example
    -------
    >>> if optimize_system():
    ...     print("System optimized for best performance")
    ... else:
    ...     print("Could not apply all optimizations")
    """
    try:
        # Apply platform-specific optimizations
        if platform.system() == "Windows":
            windows_optimizations()
        elif platform.system() == "Linux":
            linux_optimizations()
        
        # Apply common optimizations
        common_optimizations()
        
        return True
    except Exception as e:
        logger.error(f"Failed to optimize system: {e}")
        return False


def windows_optimizations() -> None:
    """
    Apply Windows-specific optimizations for Maggie AI Assistant.
    
    This function applies Windows-specific performance optimizations by
    adjusting process priority and CPU affinity. It is particularly optimized
    for systems with AMD Ryzen 9 5900X processors.
    
    Parameters
    ----------
    None
    
    Returns
    -------
    None
        This function doesn't return a value but applies system optimizations
    
    Notes
    -----
    The function applies these Windows-specific optimizations:
    
    1. Process Priority:
       - Sets the current process priority to HIGH_PRIORITY_CLASS
       - This gives the application higher CPU scheduling priority
       - Improves responsiveness for wake word detection and speech processing
    
    2. CPU Affinity Optimization:
       - For systems with 16+ logical processors (like Ryzen 9 5900X)
       - Sets affinity to use the first 16 logical processors
       - Ryzen processors have a CCX (Core Complex) architecture where
         keeping threads on the same CCX reduces inter-core latency
       - This improves cache coherency and reduces NUMA effects
    
    These optimizations require the psutil package. If it's not available,
    the function logs a warning but continues execution without crashing.
    
    The function also handles exceptions gracefully to prevent application
    crashes if optimizations cannot be applied.
    
    Example
    -------
    >>> windows_optimizations()
    # INFO: Set process priority to high
    # INFO: Set CPU affinity to first 8 physical cores: [0, 1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15]
    """
    try:
        import psutil
        # Get current process
        process = psutil.Process(os.getpid())
        
        # Set process priority to high
        process.nice(psutil.HIGH_PRIORITY_CLASS)
        logger.info("Set process priority to high")
        
        # CPU affinity optimization for Ryzen 9 5900X
        # Use the first 8 cores (16 threads for Ryzen)
        cpu_count = psutil.cpu_count(logical=True)
        if cpu_count >= 16:  # Likely Ryzen 9 with hyperthreading
            # Create affinity mask for first 8 physical cores (16 threads)
            # For Ryzen, logical processors are arranged as:
            # 0, 2, 4, ... are first 12 cores, 1, 3, 5, ... are their hyperthreaded pairs
            affinity = list(range(16))
            process.cpu_affinity(affinity)
            logger.info(f"Set CPU affinity to first 8 physical cores: {affinity}")
    except ImportError:
        logger.warning("psutil not available for Windows optimizations")
    except Exception as e:
        logger.warning(f"Failed to apply Windows performance optimizations: {e}")


def linux_optimizations() -> None:
    """
    Apply Linux-specific optimizations for Maggie AI Assistant.
    
    This function applies Linux-specific performance optimizations by
    adjusting CPU governor settings and process priority. Some optimizations
    require root privileges to apply.
    
    Parameters
    ----------
    None
    
    Returns
    -------
    None
        This function doesn't return a value but applies system optimizations
    
    Notes
    -----
    The function applies these Linux-specific optimizations:
    
    1. CPU Governor Setting (requires root):
       - Sets the CPU scaling governor to "performance" mode
       - This disables frequency scaling and keeps CPU at maximum frequency
       - Improves computational performance for AI inference and speech processing
       - Increases power consumption and heat generation
    
    2. Process Priority:
       - Sets process nice level to -10 (higher priority)
       - This gives the application higher CPU scheduling priority
       - Improves responsiveness for wake word detection
    
    The function checks if it's running with root privileges and applies
    governor settings only if root access is available. Non-root users
    still get process priority optimizations.
    
    These optimizations affect:
    - AI model inference speed (faster responses)
    - Speech recognition latency (quicker recognition)
    - Wake word detection reliability (fewer missed activations)
    
    Example
    -------
    >>> linux_optimizations()
    # INFO: Set CPU governor to performance mode  (if running as root)
    # INFO: Set process nice level to -10
    """
    try:
        # Check if running as root (required for some optimizations)
        is_root = os.geteuid() == 0
        
        if is_root:
            # Set CPU governor to performance
            os.system("echo performance | tee /sys/devices/system/cpu/cpu*/cpufreq/scaling_governor")
            logger.info("Set CPU governor to performance mode")
        else:
            logger.warning("Not running as root, skipping some system optimizations")
            
        # Set process nice level
        os.nice(-10)  # Higher priority (lower nice value)
        logger.info("Set process nice level to -10")
    except Exception as e:
        logger.warning(f"Failed to apply Linux performance optimizations: {e}")


def common_optimizations() -> None:
    """
    Apply common optimizations that work across all platforms.
    
    This function applies performance and usability optimizations that work
    consistently across Windows, Linux, and other supported platforms.
    These optimizations complement the platform-specific ones.
    
    Parameters
    ----------
    None
    
    Returns
    -------
    None
        This function doesn't return a value but applies system optimizations
    
    Notes
    -----
    The function applies these cross-platform optimizations:
    
    1. Process Name Setting:
       - Sets the visible process name to "maggie-ai-assistant"
       - Improves identifiability in process lists and monitoring tools
       - Uses the setproctitle package when available
    
    2. Python Integer String Conversion Limit:
       - Sets sys.set_int_max_str_digits(4096) in Python 3.10+
       - Limits memory allocation in arbitrary precision integer conversions
       - Prevents potential memory exhaustion from extreme integer operations
       - This is particularly useful for cryptographic operations
    
    The function handles missing dependencies gracefully by catching ImportError
    exceptions and continuing with available optimizations.
    
    The optimizations are chosen to be safe and beneficial across all platforms
    without requiring specific hardware or OS capabilities.
    
    Example
    -------
    >>> common_optimizations()
    # INFO: Set process title to 'maggie-ai-assistant'
    """
    try:
        # Set process name for better identification
        import setproctitle
        setproctitle.setproctitle("maggie-ai-assistant")
        logger.info("Set process title to 'maggie-ai-assistant'")
    except ImportError:
        pass
    
    # Set Python-specific optimizations
    if hasattr(sys, 'set_int_max_str_digits'):
        # Limit maximum string conversion size (Python 3.10+)
        sys.set_int_max_str_digits(4096)


def main() -> int:
    """
    Main entry point for the Maggie AI Assistant application.
    
    This function serves as the primary entry point and orchestrates the
    entire startup sequence for the Maggie AI Assistant. It handles command
    line parsing, system verification, configuration, and application startup.
    
    Parameters
    ----------
    None
    
    Returns
    -------
    int
        Exit code (0 for success, non-zero for error)
        
    Notes
    -----
    The function implements this startup sequence:
    
    1. Command-Line Processing:
       - Parses arguments using parse_arguments()
       - Handles flags for debug, verify, create-template, optimize, headless
    
    2. Directory Setup:
       - Ensures core directories (logs, models) exist
    
    3. Logging Configuration:
       - Sets up logging with appropriate levels (debug or info)
       - Logs system information for diagnostics
    
    4. Multiprocessing Support:
       - Configures multiprocessing start method ('spawn')
       - Ensures proper process isolation
    
    5. System Optimization (if --optimize specified):
       - Applies platform-specific performance optimizations
    
    6. Template Creation (if --create-template specified):
       - Creates recipe template document if not exists
    
    7. System Verification (if --verify specified):
       - Performs comprehensive system compatibility check
       - Exits with appropriate code after verification
    
    8. Application Startup:
       - Starts Maggie AI Assistant core services
       - Initializes GUI (unless in headless mode)
       - Sets up event handlers and signal handlers
    
    Critical errors at any stage result in appropriate error logs and
    non-zero exit codes. The function attempts to provide clear error
    messages to guide troubleshooting.
    
    Example
    -------
    >>> exit_code = main()
    >>> sys.exit(exit_code)
    """
    # Parse arguments
    args = parse_arguments()
    
    # Ensure necessary directories exist
    os.makedirs("logs", exist_ok=True)
    os.makedirs("models", exist_ok=True)
    
    # Set up logging
    setup_logging(args.debug)
    
    # Log startup and system information
    logger.info("Starting Maggie AI Assistant")
    logger.info(f"Running on Python {platform.python_version()}")
    logger.info(f"Process ID: {os.getpid()}")
    
    # Enable multiprocessing support
    setup_multiprocessing()
    
    # Optimize system if requested
    if args.optimize:
        if optimize_system():
            logger.info("System optimized for performance")
        else:
            logger.warning("System optimization failed")
    
    # Create recipe template if requested
    if args.create_template:
        create_recipe_template()
        if not args.verify:  # Exit if only create-template was specified
            return 0
    
    # Verify system if requested
    if args.verify:
        if verify_system():
            logger.info("System verification successful")
            return 0
        else:
            logger.error("System verification failed")
            return 1
    
    # Quick system check before starting
    if not verify_system():
        logger.warning("System verification failed, but attempting to start anyway")
    
    # Start the Maggie AI Assistant
    return start_maggie(args)


def setup_multiprocessing() -> None:
    """
    Set up Python multiprocessing module with appropriate settings.
    
    This function configures the Python multiprocessing module with settings
    appropriate for Maggie AI Assistant, ensuring stable and efficient
    parallel processing across multiple components.
    
    Parameters
    ----------
    None
    
    Returns
    -------
    None
        This function doesn't return a value but configures multiprocessing
    
    Notes
    -----
    The function applies these multiprocessing configurations:
    
    1. Start Method Setting:
       - Sets start method to 'spawn' instead of the default
       - 'spawn' creates a completely new Python process
       - This prevents issues with forked processes inheriting resources
       - Particularly important for processes using GPU resources
    
    This configuration is critical for:
    - Stable operation of speech recognition processes
    - Proper resource isolation between components
    - Clean process termination
    - Avoiding deadlocks and resource leaks
    
    The function handles the case where the start method has already been set
    (RuntimeError) by catching the exception and continuing execution.
    
    Example
    -------
    >>> setup_multiprocessing()
    # INFO: Set multiprocessing start method to 'spawn'
    """
    try:
        multiprocessing.set_start_method('spawn')
        logger.info("Set multiprocessing start method to 'spawn'")
    except RuntimeError:
        # Already set
        pass


def start_maggie(args: argparse.Namespace) -> int:
    """
    Initialize and start the Maggie AI Assistant.
    
    This function is responsible for initializing and starting the main
    Maggie AI Assistant application after verification and setup steps.
    It handles configuration loading, component initialization, and
    starts either the GUI or headless mode as specified.
    
    Parameters
    ----------
    args : argparse.Namespace
        Command-line arguments containing configuration options and flags.
        The relevant fields are:
        - config : str
            Path to configuration file (default: "config.yaml")
        - headless : bool
            Boolean flag for running without GUI in headless mode
        
    Returns
    -------
    int
        Exit code (0 for success, non-zero for error)
    
    Notes
    -----
    The function performs these initialization steps:
    
    1. Component Import:
       - Imports necessary modules with error handling
       - Falls back to headless mode if GUI components fail to import
    
    2. Configuration Loading:
       - Loads YAML configuration from the specified file
       - Handles missing or empty configuration gracefully
    
    3. Core Initialization:
       - Creates MaggieAI instance with loaded configuration
       - Sets up signal handlers for graceful shutdown
       - Configures error logging to event bus
    
    4. Application Startup:
       - GUI Mode: Creates Qt application and main window
       - Headless Mode: Runs in command-line with event loop
    
    5. Error Handling:
       - Catches and logs exceptions during startup
       - Returns appropriate error codes
    
    The function is designed to fail gracefully with informative error
    messages when encountering issues, rather than crashing with a
    traceback.
    
    Example
    -------
    >>> args = parse_arguments()
    >>> exit_code = start_maggie(args)
    >>> sys.exit(exit_code)
    """
    try:
        # Check if in headless mode
        if not args.headless:
            # Ensure PySide6 paths are added for GUI components 
            add_pyside6_paths()

            try:
                from PySide6.QtWidgets import QApplication
                from maggie.utils.gui import MainWindow
            except ImportError as e:
                logger.error(f"GUI components import failed: {e}")
                logger.info("Falling back to headless mode")
                args.headless = True

        # Import required components with better error handling
        from maggie.core import MaggieAI, State

    except ImportError as e:
        logger.error(f"Failed to import required module: {e}")
        return 1

    # Load config from file
    config = {}
    if os.path.exists(args.config):
        try:
            with open(args.config, 'r') as f:
                config = yaml.safe_load(f)
                if config is None:
                    logger.warning(f"Empty config file: {args.config}, using defaults")
                    config = {}
        except Exception as e:
            logger.error(f"Error reading config file: {e}")
            # Continue with empty config
    else:
        logger.warning(f"Config file not found: {args.config}, using defaults")

    # Initialize and start Maggie AI
    maggie = MaggieAI(config)
    
    # Register signal handlers for graceful shutdown
    register_signal_handlers(maggie)
    
    # Give the error_publisher function access to the event bus
    error_publisher.event_bus = maggie.event_bus
    
    # Add the error publisher as a sink with correct formatting
    logger.add(
        error_publisher,
        format="[{time:%H:%M:%S}] {level.name}: {message}",
        level="ERROR"  # Only capture ERROR and above
    )
    
    # Start Maggie core services
    success = maggie.start()
    if not success:
        logger.error("Failed to start Maggie AI core services")
        return 1
        
    # Initialize and start GUI if not in headless mode
    if not args.headless:
        try:
            app = QApplication(sys.argv)
            window = MainWindow(maggie)
            maggie.set_gui(window)  # Set bidirectional reference
            window.show()
            return app.exec()
        except Exception as e:
            logger.error(f"Error starting GUI: {e}")
            maggie.shutdown()
            return 1
    else:
        # Headless operation - keep running until signal received
        if args.headless:
            logger.info("Running in headless mode")
            try:
                # Main thread waits here
                while maggie.state != State.SHUTDOWN:
                    time.sleep(1)
                return 0
            except KeyboardInterrupt:
                logger.info("Keyboard interrupt received, shutting down")
                maggie.shutdown()
                return 0


def register_signal_handlers(maggie) -> None:
    """
    Register signal handlers for graceful shutdown of Maggie AI Assistant.
    
    This function sets up signal handlers for SIGINT (Ctrl+C) and SIGTERM
    (termination request) to ensure that the Maggie AI Assistant shuts down
    gracefully when receiving these signals, properly releasing resources
    and saving state.
    
    Parameters
    ----------
    maggie : MaggieAI
        Instance of the Maggie AI Assistant that needs to be shutdown
        gracefully when signals are received
    
    Returns
    -------
    None
        This function doesn't return a value but registers signal handlers
    
    Notes
    -----
    The function sets up handlers for:
    
    1. SIGINT Signal:
       - Sent when user presses Ctrl+C in terminal
       - Common during development and interactive use
    
    2. SIGTERM Signal:
       - Sent by process managers (systemd, Docker, etc.)
       - Used for orderly shutdown requests
    
    The registered handler:
    1. Logs the received signal
    2. Calls maggie.shutdown() to perform orderly cleanup:
       - Stops active processes
       - Saves current state if needed
       - Releases hardware resources (GPU memory, audio devices)
       - Closes open files and connections
    3. Exits the application with code 0 (success)
    
    The function handles exceptions during registration gracefully,
    logging warnings but not crashing the application if signal
    handlers cannot be registered.
    
    Example
    -------
    >>> maggie_instance = MaggieAI(config)
    >>> register_signal_handlers(maggie_instance)
    # INFO: Registered signal handlers for graceful shutdown
    """
    try:
        import signal
        def signal_handler(sig, frame):
            logger.info(f"Received signal {sig}, shutting down gracefully")
            maggie.shutdown()
            sys.exit(0)
            
        signal.signal(signal.SIGINT, signal_handler)
        signal.signal(signal.SIGTERM, signal_handler)
        logger.info("Registered signal handlers for graceful shutdown")
    except Exception as e:
        logger.warning(f"Failed to register signal handlers: {e}")


def add_pyside6_paths():
    """
    Find PySide6 paths dynamically across platforms and add them to sys.path.
    
    This function attempts to locate PySide6 installation directories using
    multiple fallback methods across different platforms, then adds them to
    the Python sys.path to ensure PySide6 modules can be imported correctly.
    
    Parameters
    ----------
    None
    
    Returns
    -------
    None
        This function doesn't return a value but modifies sys.path
    
    Notes
    -----
    The function uses three methods to locate PySide6, in order:
    
    1. Pip Package Information:
       - Uses pip to find the package installation location
       - Parses the output to extract the base path
       - Constructs PySide6 paths based on typical directory structure
    
    2. Site-Packages Search:
       - Uses site.getsitepackages() to find potential locations
       - Checks for PySide6 directory in each site-packages location
       - Adds relevant subdirectories to the path
    
    3. Virtual Environment Detection:
       - Detects if running in a virtual environment
       - Checks platform-specific locations within the venv
       - Handles differences between Windows and Linux paths
    
    The function handles errors gracefully at each stage, continuing to
    the next method if the current one fails. This ensures maximum
    compatibility across different Python environments and installations.
    
    This workaround is necessary because:
    - PySide6 has a complex directory structure with DLLs/SOs
    - Some files need to be in PATH/LD_LIBRARY_PATH
    - Different platforms have different directory structures
    - Installation methods (pip, conda, system package) vary
    
    Example
    -------
    >>> add_pyside6_paths()
    >>> from PySide6.QtWidgets import QApplication  # Now imports successfully
    """
    # Method 1: Use pip to find the package location
    try:
        import subprocess
        import json
        import re
        
        # Run pip show command for PySide6
        pip_show_pyside6 = subprocess.run(
            [sys.executable, "-m", "pip", "show", "pyside6"],
            capture_output=True,
            text=True
        ).stdout.strip()

        # Extract file path for pyside6
        fp = re.search(
            r"location:\s*(.+)", 
            pip_show_pyside6.lower()
        )

        # Extract base path for pyside6
        bp = None if not fp else fp.group(1).strip()

        # Create default PySide6 paths
        pyside6_paths = [] if not bp else [
            os.path.join(bp, "PySide6"),
            os.path.join(bp, "PySide6", "Qt6"),
            os.path.join(bp, "PySide6", "Qt6", "bin")
        ]

        result = [p for p in pyside6_paths if os.path.exists(p)]

        if result and len(result) > 0:
            for p in result:
                if p not in sys.path:
                    sys.path.append(p)

    except Exception as e:
        print(f"Error finding PySide6 paths using pip: {e}")
   
    # Method 2: Try to find in site-packages
    try:
        import site
        # Check in site-packages
        for site_dir in site.getsitepackages():
            pyside_dir = os.path.join(site_dir, "PySide6")
            if os.path.exists(pyside_dir):
                pyside_paths = [
                    pyside_dir,
                    os.path.join(pyside_dir, "Qt6"),
                    os.path.join(pyside_dir, "Qt6", "bin")
                ]
                result = [p for p in pyside_paths if os.path.exists(p)]

                if result and len(result) > 0:
                    for p in result:
                        if p not in sys.path:
                            sys.path.append(p)

    except Exception as e:
        print(f"Error finding PySide6 paths using site-packages: {e}")
                
    # Method 3: Check in virtual environment
    try:
        venv_dir = os.path.dirname(os.path.dirname(sys.executable))
        if os.path.exists(venv_dir):
            # Different structures in different platforms
            potential_paths = [
                os.path.join(venv_dir, "Lib", "site-packages", "PySide6"),  # Windows
                os.path.join(venv_dir, "lib", "python3.10", "site-packages", "PySide6"),  # Linux
                os.path.join(venv_dir, "lib", "site-packages", "PySide6")  # Alternative
            ]
            
            for pp in potential_paths:
                if os.path.exists(pp):
                    pyside_paths = [
                        pp,
                        os.path.join(pp, "Qt6"),
                        os.path.join(pp, "Qt6", "bin")
                    ]
                    result = [p for p in pyside_paths if os.path.exists(p)]

                    if result and len(result) > 0:
                        for p in result:
                            if p not in sys.path:
                                sys.path.append(p)
    
    except Exception as e:
        print(f"Error finding PySide6 in virtual environment: {e}")
    
    return 


if __name__ == "__main__":
    sys.exit(main())