<<<<<<< HEAD
"""
Maggie AI Assistant - Setup Verification Script
===========================================
Verifies that all required components are properly installed and configured.
"""

import os
import sys
import platform
import subprocess
from typing import List, Dict, Any, Tuple

def print_color(text: str, color: str) -> None:
    """
    Print colored text to the console.
    
    Parameters
    ----------
    text : str
        Text to print
    color : str
        Color to use (red, green, yellow, cyan, magenta)
    """
    colors = {
        "red": "\033[91m",
        "green": "\033[92m",
        "yellow": "\033[93m",
        "cyan": "\033[96m",
        "magenta": "\033[95m",
        "reset": "\033[0m"
    }
    
    # Windows requires special handling for colors
    if platform.system() == "Windows":
        os.system("")  # Enable VT100 escape sequences
    
    print(f"{colors.get(color, '')}{text}{colors['reset']}")

def check_python_version() -> bool:
    """
    Check if Python version is 3.10.x.
    
    Returns
    -------
    bool
        True if Python version is 3.10.x, False otherwise
    """
    version = platform.python_version_tuple()
    if int(version[0]) != 3 or int(version[1]) != 10:
        print_color(f"ERROR: Unsupported Python version: {platform.python_version()}", "red")
        print_color("Maggie requires Python 3.10.x", "red")
        return False
    
    print_color(f"✓ Python version: {platform.python_version()}", "green")
    return True

def check_pip_packages() -> bool:
    """
    Check if required pip packages are installed.
    
    Returns
    -------
    bool
        True if all required packages are installed, False otherwise
    """
    required_packages = [
        "pvporcupine",
        "SpeechRecognition",
        "PyAudio",
        "faster-whisper",
        "ctransformers",
        "PyQt6",
        "transitions",
        "python-docx",
        "numpy",
        "PyYAML",
        "loguru",
        "torch",
        "psutil"
    ]
    
    installed_packages = []
    try:
        # Get list of installed packages
        result = subprocess.run(
            [sys.executable, "-m", "pip", "list"],
            capture_output=True,
            text=True,
            check=True
        )
        
        for line in result.stdout.split("\n")[2:]:  # Skip header rows
            if line.strip():
                package_name = line.split()[0].lower()
                installed_packages.append(package_name)
    except subprocess.SubprocessError as e:
        print_color(f"ERROR: Failed to get installed packages: {e}", "red")
        return False
    
    missing_packages = []
    for package in required_packages:
        package_lower = package.lower()
        if package_lower not in installed_packages:
            # Special case for hyphenated packages
            alt_name = package_lower.replace("-", "_")
            if alt_name not in installed_packages:
                missing_packages.append(package)
    
    if missing_packages:
        print_color(f"ERROR: Missing required packages: {', '.join(missing_packages)}", "red")
        print_color("Install with: pip install -r requirements.txt", "red")
        return False
    
    print_color("✓ All required Python packages are installed", "green")
    return True

def check_gpu() -> Tuple[bool, dict]:
    """
    Check if GPU is available and get info.
    
    Returns
    -------
    Tuple[bool, dict]
        Success status and GPU information
    """
    gpu_info = {"available": False, "name": None, "memory": None, "cuda": None}
    
    try:
        import torch
        
        if torch.cuda.is_available():
            gpu_info["available"] = True
            gpu_info["name"] = torch.cuda.get_device_name(0)
            gpu_info["memory"] = torch.cuda.get_device_properties(0).total_memory / (1024**3)
            gpu_info["cuda"] = torch.version.cuda
            
            print_color(f"✓ GPU detected: {gpu_info['name']}", "green")
            print_color(f"✓ GPU memory: {gpu_info['memory']:.2f} GB", "green")
            print_color(f"✓ CUDA version: {gpu_info['cuda']}", "green")
            
            # Check for RTX 3080 compatibility
            if "3080" not in gpu_info["name"] and gpu_info["memory"] < 10.0:
                print_color(
                    f"WARNING: Your GPU ({gpu_info['name']}) has less memory than the "
                    f"recommended RTX 3080 (10GB). Performance may be reduced.",
                    "yellow"
                )
            
            return True, gpu_info
        else:
            print_color("WARNING: CUDA not available, GPU acceleration disabled", "yellow")
            print_color("Performance may be significantly reduced without GPU acceleration", "yellow")
            return False, gpu_info
    except ImportError:
        print_color("ERROR: PyTorch not installed, GPU acceleration unavailable", "red")
        return False, gpu_info

def check_directories() -> bool:
    """
    Check if required directories exist, create if necessary.
    
    Returns
    -------
    bool
        True if all directories exist or were created, False on error
    """
    required_dirs = [
        "logs",
        "models",
        "models/tts",
        "recipes",
        "templates"
    ]
    
    all_created = True
    for directory in required_dirs:
        if not os.path.exists(directory):
            try:
                os.makedirs(directory, exist_ok=True)
                print_color(f"✓ Created directory: {directory}", "green")
            except Exception as e:
                print_color(f"ERROR: Failed to create directory {directory}: {e}", "red")
                all_created = False
        else:
            print_color(f"✓ Directory exists: {directory}", "green")
    
    return all_created

def check_config_file() -> bool:
    """
    Check if config.yaml exists, create from example if necessary.
    
    Returns
    -------
    bool
        True if config exists or was created, False on error
    """
    config_path = "config.yaml"
    example_path = "config.yaml.example"
    
    if os.path.exists(config_path):
        print_color(f"✓ Configuration file exists: {config_path}", "green")
        return True
    
    if not os.path.exists(example_path):
        print_color(f"ERROR: Example configuration file not found: {example_path}", "red")
        return False
    
    try:
        import shutil
        shutil.copy(example_path, config_path)
        print_color(f"✓ Created configuration file from example: {config_path}", "green")
        print_color("WARNING: You need to edit config.yaml to add your Picovoice access key", "yellow")
        return True
    except Exception as e:
        print_color(f"ERROR: Failed to create configuration file: {e}", "red")
        return False

def create_template_files() -> bool:
    """
    Create template files if they don't exist.
    
    Returns
    -------
    bool
        True if templates exist or were created, False on error
    """
    template_path = "templates/recipe_template.docx"
    
    if os.path.exists(template_path):
        print_color(f"✓ Recipe template exists: {template_path}", "green")
        return True
    
    try:
        from docx import Document
        
        doc = Document()
        doc.add_heading("Recipe Name", level=1)
        doc.add_heading("Ingredients", level=2)
        doc.add_paragraph("• Ingredient 1", style='ListBullet')
        doc.add_paragraph("• Ingredient 2", style='ListBullet')
        doc.add_heading("Instructions", level=2)
        doc.add_paragraph("1. Step 1")
        doc.add_paragraph("2. Step 2")
        doc.add_heading("Notes", level=2)
        doc.add_paragraph("Add any additional notes here.")
        
        os.makedirs(os.path.dirname(template_path), exist_ok=True)
        doc.save(template_path)
        print_color(f"✓ Created recipe template: {template_path}", "green")
        return True
    except Exception as e:
        print_color(f"ERROR: Failed to create recipe template: {e}", "red")
        return False

def main() -> int:
    """
    Run all checks and report status.
    
    Returns
    -------
    int
        Exit code (0 for success, non-zero for error)
    """
    print_color("=== Maggie AI Assistant - Setup Verification ===", "cyan")
    print_color(f"System: {platform.system()} {platform.release()}", "cyan")
    
    checks = [
        ("Python Version", check_python_version()),
        ("Required Packages", check_pip_packages()),
        ("GPU Availability", check_gpu()[0]),
        ("Required Directories", check_directories()),
        ("Configuration File", check_config_file()),
        ("Template Files", create_template_files())
    ]
    
    print_color("\n=== Summary ===", "cyan")
    all_passed = True
    for name, status in checks:
        if status:
            print_color(f"✓ {name}: PASS", "green")
        else:
            print_color(f"✗ {name}: FAIL", "red")
            all_passed = False
    
    if all_passed:
        print_color("\nAll checks passed! Maggie AI Assistant is ready to run.", "green")
        print_color("Start with: python main.py", "green")
        return 0
    else:
        print_color("\nSome checks failed. Please fix the issues before running Maggie.", "red")
        return 1

if __name__ == "__main__":
=======
"""
Maggie AI Assistant - Setup Verification Script
===========================================
Verifies that all required components are properly installed and configured.
"""

import os
import sys
import platform
import subprocess
from typing import List, Dict, Any, Tuple

def print_color(text: str, color: str) -> None:
    """
    Print colored text to the console.
    
    Parameters
    ----------
    text : str
        Text to print
    color : str
        Color to use (red, green, yellow, cyan, magenta)
    """
    colors = {
        "red": "\033[91m",
        "green": "\033[92m",
        "yellow": "\033[93m",
        "cyan": "\033[96m",
        "magenta": "\033[95m",
        "reset": "\033[0m"
    }
    
    # Windows requires special handling for colors
    if platform.system() == "Windows":
        os.system("")  # Enable VT100 escape sequences
    
    print(f"{colors.get(color, '')}{text}{colors['reset']}")

def check_python_version() -> bool:
    """
    Check if Python version is 3.10.x.
    
    Returns
    -------
    bool
        True if Python version is 3.10.x, False otherwise
    """
    version = platform.python_version_tuple()
    if int(version[0]) != 3 or int(version[1]) != 10:
        print_color(f"ERROR: Unsupported Python version: {platform.python_version()}", "red")
        print_color("Maggie requires Python 3.10.x", "red")
        return False
    
    print_color(f"✓ Python version: {platform.python_version()}", "green")
    return True

def check_pip_packages() -> bool:
    """
    Check if required pip packages are installed.
    
    Returns
    -------
    bool
        True if all required packages are installed, False otherwise
    """
    required_packages = [
        "pvporcupine",
        "SpeechRecognition",
        "PyAudio",
        "faster-whisper",
        "ctransformers",
        "PyQt6",
        "transitions",
        "python-docx",
        "numpy",
        "PyYAML",
        "loguru",
        "torch",
        "psutil"
    ]
    
    installed_packages = []
    try:
        # Get list of installed packages
        result = subprocess.run(
            [sys.executable, "-m", "pip", "list"],
            capture_output=True,
            text=True,
            check=True
        )
        
        for line in result.stdout.split("\n")[2:]:  # Skip header rows
            if line.strip():
                package_name = line.split()[0].lower()
                installed_packages.append(package_name)
    except subprocess.SubprocessError as e:
        print_color(f"ERROR: Failed to get installed packages: {e}", "red")
        return False
    
    missing_packages = []
    for package in required_packages:
        package_lower = package.lower()
        if package_lower not in installed_packages:
            # Special case for hyphenated packages
            alt_name = package_lower.replace("-", "_")
            if alt_name not in installed_packages:
                missing_packages.append(package)
    
    if missing_packages:
        print_color(f"ERROR: Missing required packages: {', '.join(missing_packages)}", "red")
        print_color("Install with: pip install -r requirements.txt", "red")
        return False
    
    print_color("✓ All required Python packages are installed", "green")
    return True

def check_gpu() -> Tuple[bool, dict]:
    """
    Check if GPU is available and get info.
    
    Returns
    -------
    Tuple[bool, dict]
        Success status and GPU information
    """
    gpu_info = {"available": False, "name": None, "memory": None, "cuda": None}
    
    try:
        import torch
        
        if torch.cuda.is_available():
            gpu_info["available"] = True
            gpu_info["name"] = torch.cuda.get_device_name(0)
            gpu_info["memory"] = torch.cuda.get_device_properties(0).total_memory / (1024**3)
            gpu_info["cuda"] = torch.version.cuda
            
            print_color(f"✓ GPU detected: {gpu_info['name']}", "green")
            print_color(f"✓ GPU memory: {gpu_info['memory']:.2f} GB", "green")
            print_color(f"✓ CUDA version: {gpu_info['cuda']}", "green")
            
            # Check for RTX 3080 compatibility
            if "3080" not in gpu_info["name"] and gpu_info["memory"] < 10.0:
                print_color(
                    f"WARNING: Your GPU ({gpu_info['name']}) has less memory than the "
                    f"recommended RTX 3080 (10GB). Performance may be reduced.",
                    "yellow"
                )
            
            return True, gpu_info
        else:
            print_color("WARNING: CUDA not available, GPU acceleration disabled", "yellow")
            print_color("Performance may be significantly reduced without GPU acceleration", "yellow")
            return False, gpu_info
    except ImportError:
        print_color("ERROR: PyTorch not installed, GPU acceleration unavailable", "red")
        return False, gpu_info

def check_directories() -> bool:
    """
    Check if required directories exist, create if necessary.
    
    Returns
    -------
    bool
        True if all directories exist or were created, False on error
    """
    required_dirs = [
        "logs",
        "models",
        "models/tts",
        "recipes",
        "templates"
    ]
    
    all_created = True
    for directory in required_dirs:
        if not os.path.exists(directory):
            try:
                os.makedirs(directory, exist_ok=True)
                print_color(f"✓ Created directory: {directory}", "green")
            except Exception as e:
                print_color(f"ERROR: Failed to create directory {directory}: {e}", "red")
                all_created = False
        else:
            print_color(f"✓ Directory exists: {directory}", "green")
    
    return all_created

def check_config_file() -> bool:
    """
    Check if config.yaml exists, create from example if necessary.
    
    Returns
    -------
    bool
        True if config exists or was created, False on error
    """
    config_path = "config.yaml"
    example_path = "config.yaml.example"
    
    if os.path.exists(config_path):
        print_color(f"✓ Configuration file exists: {config_path}", "green")
        return True
    
    if not os.path.exists(example_path):
        print_color(f"ERROR: Example configuration file not found: {example_path}", "red")
        return False
    
    try:
        import shutil
        shutil.copy(example_path, config_path)
        print_color(f"✓ Created configuration file from example: {config_path}", "green")
        print_color("WARNING: You need to edit config.yaml to add your Picovoice access key", "yellow")
        return True
    except Exception as e:
        print_color(f"ERROR: Failed to create configuration file: {e}", "red")
        return False

def create_template_files() -> bool:
    """
    Create template files if they don't exist.
    
    Returns
    -------
    bool
        True if templates exist or were created, False on error
    """
    template_path = "templates/recipe_template.docx"
    
    if os.path.exists(template_path):
        print_color(f"✓ Recipe template exists: {template_path}", "green")
        return True
    
    try:
        from docx import Document
        
        doc = Document()
        doc.add_heading("Recipe Name", level=1)
        doc.add_heading("Ingredients", level=2)
        doc.add_paragraph("• Ingredient 1", style='ListBullet')
        doc.add_paragraph("• Ingredient 2", style='ListBullet')
        doc.add_heading("Instructions", level=2)
        doc.add_paragraph("1. Step 1")
        doc.add_paragraph("2. Step 2")
        doc.add_heading("Notes", level=2)
        doc.add_paragraph("Add any additional notes here.")
        
        os.makedirs(os.path.dirname(template_path), exist_ok=True)
        doc.save(template_path)
        print_color(f"✓ Created recipe template: {template_path}", "green")
        return True
    except Exception as e:
        print_color(f"ERROR: Failed to create recipe template: {e}", "red")
        return False

def main() -> int:
    """
    Run all checks and report status.
    
    Returns
    -------
    int
        Exit code (0 for success, non-zero for error)
    """
    print_color("=== Maggie AI Assistant - Setup Verification ===", "cyan")
    print_color(f"System: {platform.system()} {platform.release()}", "cyan")
    
    checks = [
        ("Python Version", check_python_version()),
        ("Required Packages", check_pip_packages()),
        ("GPU Availability", check_gpu()[0]),
        ("Required Directories", check_directories()),
        ("Configuration File", check_config_file()),
        ("Template Files", create_template_files())
    ]
    
    print_color("\n=== Summary ===", "cyan")
    all_passed = True
    for name, status in checks:
        if status:
            print_color(f"✓ {name}: PASS", "green")
        else:
            print_color(f"✗ {name}: FAIL", "red")
            all_passed = False
    
    if all_passed:
        print_color("\nAll checks passed! Maggie AI Assistant is ready to run.", "green")
        print_color("Start with: python main.py", "green")
        return 0
    else:
        print_color("\nSome checks failed. Please fix the issues before running Maggie.", "red")
        return 1

if __name__ == "__main__":
>>>>>>> 6062514b96de23fbf6dcdbfd4420d6e2f22903ff
    sys.exit(main())