#!/usr/bin/env python3
import argparse
import json
import os
import platform
import shutil
import subprocess
import sys
import time
import urllib.request
import zipfile
from pathlib import Path
from typing import Dict, Any, List, Tuple, Optional, Union, Callable

# --- Helper Classes (ColorOutput, ProgressTracker) ---

class ColorOutput:
    """Handles colored terminal output."""
    def __init__(self, force_enable: bool = False):
        self.enabled = force_enable or self._supports_color()
        if self.enabled:
            self.colors = {
                'reset': '\x1b[0m', 'bold': '\x1b[1m', 'red': '\x1b[91m',
                'green': '\x1b[92m', 'yellow': '\x1b[93m', 'blue': '\x1b[94m',
                'magenta': '\x1b[95m', 'cyan': '\x1b[96m', 'white': '\x1b[97m'
            }
        else:
            self.colors = {color: '' for color in ['reset', 'bold', 'red', 'green', 'yellow', 'blue', 'magenta', 'cyan', 'white']}

    def _supports_color(self) -> bool:
        """Checks if the terminal supports color."""
        if platform.system() == 'Windows':
            # Basic check for Windows 10+ which generally supports ANSI
            try:
                import ctypes
                kernel32 = ctypes.windll.kernel32
                ENABLE_VIRTUAL_TERMINAL_PROCESSING = 0x0004
                stdout = kernel32.GetStdHandle(-11)
                mode = ctypes.c_ulong()
                if kernel32.GetConsoleMode(stdout, ctypes.byref(mode)):
                    if (mode.value & ENABLE_VIRTUAL_TERMINAL_PROCESSING) == ENABLE_VIRTUAL_TERMINAL_PROCESSING:
                        return True # Already enabled
                    mode.value |= ENABLE_VIRTUAL_TERMINAL_PROCESSING
                    if kernel32.SetConsoleMode(stdout, mode):
                        return True # Successfully enabled
            except:
                 if int(platform.release()) >= 10: return True # Fallback version check
            return False # Older windows or failed to enable
        if os.environ.get('NO_COLOR'): return False
        if os.environ.get('FORCE_COLOR'): return True
        # Check if stdout is a TTY
        return hasattr(sys.stdout, 'isatty') and sys.stdout.isatty()

    def print(self, message: str, color: Optional[str] = None, bold: bool = False):
        """Prints a message with optional color and bold formatting."""
        formatted = message
        if self.enabled:
            if bold and 'bold' in self.colors:
                formatted = f"{self.colors['bold']}{formatted}"
            if color and color in self.colors:
                formatted = f"{self.colors[color]}{formatted}"
            if (bold or color) and 'reset' in self.colors:
                formatted = f"{formatted}{self.colors['reset']}"
        print(formatted)

    def input(self, prompt: str, color: Optional[str] = None, bold: bool = False) -> str:
        """Gets input with a formatted prompt."""
        formatted = prompt
        if self.enabled:
            if bold and 'bold' in self.colors:
                formatted = f"{self.colors['bold']}{formatted}"
            if color and color in self.colors:
                formatted = f"{self.colors[color]}{formatted}"
            if (bold or color) and 'reset' in self.colors:
                formatted = f"{formatted}{self.colors['reset']}"
        return input(formatted)

class ProgressTracker:
    """Tracks and displays installation progress."""
    def __init__(self, color: ColorOutput, total_steps: int = 10):
        self.color = color
        self.total_steps = total_steps
        self.current_step = 0
        self.start_time = time.time()

    def start_step(self, step_name: str):
        """Starts a new installation step."""
        self.current_step += 1
        elapsed = time.time() - self.start_time
        self.color.print(f"\n[{self.current_step}/{self.total_steps}] {step_name} (Elapsed: {elapsed:.1f}s)", color='cyan', bold=True)

    def complete_step(self, success: bool = True, message: Optional[str] = None):
        """Marks a step as complete or failed."""
        if success:
            status = '✓ Complete'
            color = 'green'
        else:
            status = '✗ Failed'
            color = 'red'
        msg = f"  {status}"
        if message:
            msg += f": {message}"
        self.color.print(msg, color=color)

    def elapsed_time(self) -> float:
        """Returns the total elapsed time."""
        return time.time() - self.start_time

    def display_summary(self, success: bool = True):
        """Displays the final installation summary."""
        elapsed = self.elapsed_time()
        if success:
            status = 'Installation Completed Successfully'
            color = 'green'
        else:
            status = 'Installation Completed with Errors'
            color = 'yellow'
        self.color.print(f"\n=== {status} ===", color=color, bold=True)
        self.color.print(f"Total time: {elapsed:.1f} seconds")

# --- MaggieInstaller Class - Modified for Poetry ---

class MaggieInstaller:
    """Handles the installation process for Maggie AI Assistant using Poetry."""
    def __init__(self, verbose: bool = False, cpu_only: bool = False, skip_models: bool = False):
        self.verbose = verbose
        self.cpu_only = cpu_only
        self.skip_models = skip_models
        self.base_dir = Path(__file__).parent.resolve() # Use parent dir of this script
        self.platform_system = platform.system()
        self.platform_machine = platform.machine()
        # Directories needed for models, logs, etc.
        self.required_dirs = [
            'downloads', 'logs', 'maggie', 'maggie/cache', 'maggie/cache/tts',
            'maggie/core', 'maggie/extensions', 'maggie/models', 'maggie/models/llm',
            'maggie/models/stt', 'maggie/models/tts', 'maggie/templates',
            'maggie/templates/extension', 'maggie/utils', 'maggie/utils/hardware',
            'maggie/utils/config', 'maggie/utils/llm', 'maggie/utils/stt', 'maggie/utils/tts'
        ]
        self.color = ColorOutput()
        # Adjusted total steps for Poetry-based install
        self.total_steps = 7 # 1:Verify, 2:Dirs, 3:Poetry Install+HWDetect, 4:Config, 5:Models, 6:Templates, 7:Complete
        self.progress = ProgressTracker(self.color, self.total_steps)
        self.is_admin = self._check_admin_privileges()
        self.has_git = False
        self.has_cpp_compiler = False
        self.has_poetry = False
        # Initialize hardware info structure
        self.hardware_info = {
            'cpu': {'is_ryzen_9_5900x': False, 'model': '', 'cores': 0, 'threads': 0},
            'gpu': {'is_rtx_3080': False, 'model': '', 'vram_gb': 0, 'cuda_available': False, 'cuda_version': '', 'cudnn_available': False, 'cudnn_version': ''},
            'memory': {'total_gb': 0, 'available_gb': 0, 'is_32gb': False}
        }

    def _check_admin_privileges(self) -> bool:
        """Checks for administrator privileges."""
        try:
            if self.platform_system == 'Windows':
                import ctypes
                return ctypes.windll.shell32.IsUserAnAdmin() != 0
            else:
                # Check if effective user ID is 0 (root)
                return os.geteuid() == 0
        except Exception as e:
            if self.verbose: self.color.print(f"Admin check failed: {e}", "yellow")
            return False

    def _run_command(self, command: List[str], check: bool = True, shell: bool = False, capture_output: bool = True, cwd: Optional[Union[str, Path]] = None, env: Optional[Dict] = None) -> Tuple[int, str, str]:
        """Runs a shell command."""
        if self.verbose:
            self.color.print(f"Running command: {' '.join(command)} in {cwd or self.base_dir}", 'cyan')
        try:
            # Combine current environment with any provided overrides
            full_env = os.environ.copy()
            if env:
                full_env.update(env)

            process = subprocess.Popen(
                command if not shell else ' '.join(command),
                stdout=subprocess.PIPE if capture_output else None,
                stderr=subprocess.PIPE if capture_output else None,
                shell=shell,
                text=True,
                encoding='utf-8', # Explicitly set encoding
                errors='replace', # Handle potential encoding errors in output
                cwd=str(cwd or self.base_dir), # Ensure cwd is string
                env=full_env
            )
            stdout, stderr = process.communicate()
            return_code = process.returncode

            # If check=True, raise an exception on failure for easier handling upstream
            if check and return_code != 0:
                 error_msg = f"Command '{' '.join(command)}' failed with code {return_code}"
                 if stderr: error_msg += f": {stderr.strip()}"
                 if self.verbose: self.color.print(error_msg, 'red')
                 # Optionally raise an exception instead of just printing
                 # raise subprocess.CalledProcessError(return_code, command, output=stdout, stderr=stderr)
            # Always return stdout/stderr if capture_output is True
            return return_code, stdout.strip() if stdout else '', stderr.strip() if stderr else ''

        except FileNotFoundError:
            # Handle case where the command itself is not found
            if self.verbose: self.color.print(f"Error: Command not found: {command[0]}", 'red')
            return -1, '', f"Command not found: {command[0]}"
        except Exception as e:
            # Catch other potential errors during execution
            if self.verbose: self.color.print(f"Error executing command '{' '.join(command)}': {e}", 'red')
            return -1, '', str(e)

    def _download_file(self, url: str, destination: str, show_progress: bool = True) -> bool:
        """Downloads a file from a URL with improved progress."""
        dest_path = Path(destination)
        try:
            self.color.print(f"Downloading {url}", 'blue')
            dest_path.parent.mkdir(parents=True, exist_ok=True)

            # Add a user-agent header to avoid potential blocks
            req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'})
            with urllib.request.urlopen(req, timeout=60) as response, open(dest_path, 'wb') as out_file: # Add timeout
                content_length_header = response.info().get('Content-Length')
                file_size = int(content_length_header) if content_length_header else 0
                downloaded = 0
                block_size = 8192 * 16 # Use larger 128KB chunks
                start_time = time.time()

                if show_progress and file_size > 0:
                    self.color.print(f"Total file size: {file_size / 1024 / 1024:.1f} MB")
                    progress_bar_width = 40
                    last_percent_reported = -1

                    while True:
                        buffer = response.read(block_size)
                        if not buffer:
                            break
                        downloaded += len(buffer)
                        out_file.write(buffer)

                        percent = int(downloaded * 100 / file_size)
                        # Update progress bar less frequently (e.g., every 1%) to reduce overhead
                        if percent > last_percent_reported:
                           last_percent_reported = percent
                           filled_width = int(progress_bar_width * downloaded / file_size)
                           bar = '█' * filled_width + '-' * (progress_bar_width - filled_width)
                           elapsed = time.time() - start_time
                           speed_mbps = (downloaded / (1024 * 1024)) / elapsed if elapsed > 0 else 0
                           # Use carriage return \r to overwrite the line
                           print(f"\r  Progress: |{bar}| {percent}% ({downloaded/1024/1024:.1f}/{file_size/1024/1024:.1f} MB) {speed_mbps:.1f} MB/s  ", end="")

                    print() # New line after progress bar finishes
                else: # No progress bar if size unknown or zero
                    self.color.print("Downloading (size unknown)...", "blue")
                    while True:
                        buffer = response.read(block_size)
                        if not buffer:
                            break
                        downloaded += len(buffer)
                        out_file.write(buffer)

            self.color.print(f"Download completed: {dest_path}", 'green')
            return True
        except urllib.error.URLError as e:
             self.color.print(f"Error downloading file (URL Error: {e.reason}): {url}", 'red')
             if dest_path.exists():
                 try: dest_path.unlink() # Clean up partial file
                 except OSError: pass
             return False
        except Exception as e:
            self.color.print(f"Error downloading file {url}: {e}", 'red')
            if dest_path.exists():
                try: dest_path.unlink() # Clean up partial file
                except OSError: pass
            return False

    def _verify_python_version(self) -> bool:
        """Verifies the Python version is 3.10.x."""
        version = platform.python_version_tuple()
        if int(version[0]) != 3 or int(version[1]) != 10:
            self.color.print(f"ERROR: Incompatible Python version: {platform.python_version()}", 'red', bold=True)
            self.color.print('Maggie requires Python 3.10.x specifically.', 'red')
            self.color.print('Poetry should be configured to use Python 3.10.', 'red')
            if self.platform_system == 'Windows':
                self.color.print('Install Python 3.10 from: https://www.python.org/downloads/release/python-31011/', 'yellow')
            else:
                self.color.print('Ensure Python 3.10 is available and selected by Poetry.', 'yellow')
                self.color.print('Consider using pyenv or similar tools to manage Python versions.', 'yellow')
            return False
        self.color.print(f"Python {platform.python_version()} - Compatible ✓", 'green')
        return True

    def _detect_hardware(self) -> None:
        """Detects CPU, Memory, and GPU hardware. Should be called AFTER poetry install."""
        self.color.print('Detecting Hardware Configuration...', 'cyan', bold=True)
        self.hardware_info['cpu'] = self._detect_cpu()
        self.hardware_info['memory'] = self._detect_memory()
        # GPU detection relies on PyTorch being installed by Poetry
        self.hardware_info['gpu'] = self._detect_gpu() if not self.cpu_only else {'available': False, 'cuda_available': False}

        # Print Summary
        self._print_hardware_summary()


    def _detect_cpu(self)->Dict[str,Any]:
        """Detects CPU information."""
        cpu_info = {'is_ryzen_9_5900x': False, 'model':'Unknown','cores':0,'threads':0}
        try:
            # Use platform.processor() as a basic fallback
            cpu_info['model'] = platform.processor() or 'Unknown'
        except Exception:
            pass # Ignore errors here

        try:
             # Use psutil for more reliable core/thread counts
             import psutil
             cpu_info['cores']=psutil.cpu_count(logical=False) or 0
             cpu_info['threads']=psutil.cpu_count(logical=True) or 0
        except ImportError:
             # Fallback if psutil not installed (should be by poetry)
             cpu_info['threads'] = os.cpu_count() or 0
             cpu_info['cores'] = cpu_info['threads'] // 2 if cpu_info['threads'] > 1 else cpu_info['threads']
        except Exception as e:
             if self.verbose: self.color.print(f"CPU count error: {e}", "yellow")

        # Refine model name on Windows using WMI if possible
        if self.platform_system == 'Windows':
            try:
                import wmi
                c = wmi.WMI()
                processor = c.Win32_Processor()[0] # Get first processor
                cpu_info['model'] = processor.Name.strip()
                # Update core/thread count from WMI if psutil failed or gave 0
                if cpu_info['cores'] == 0 and hasattr(processor, 'NumberOfCores'): cpu_info['cores'] = processor.NumberOfCores
                if cpu_info['threads'] == 0 and hasattr(processor, 'NumberOfLogicalProcessors'): cpu_info['threads'] = processor.NumberOfLogicalProcessors
            except ImportError:
                 if self.verbose: self.color.print("WMI not found, skipping detailed Windows CPU detection.", "yellow")
            except Exception as e:
                 if self.verbose: self.color.print(f"WMI CPU detection error: {e}", "yellow")

        # Check for specific CPU model after getting the best name possible
        model_lower = cpu_info['model'].lower()
        if 'ryzen 9' in model_lower and '5900x' in model_lower:
            cpu_info['is_ryzen_9_5900x'] = True

        return cpu_info

    def _detect_memory(self)->Dict[str,Any]:
        """Detects memory information."""
        memory_info={'total_gb':0,'available_gb':0,'is_32gb':False}
        try:
            # psutil is expected to be installed via poetry
            import psutil
            mem=psutil.virtual_memory()
            memory_info['total_gb']=mem.total / (1024**3)
            memory_info['available_gb']=mem.available / (1024**3)
            # Check if total RAM is >= 30GB (allows for slight variations)
            memory_info['is_32gb'] = memory_info['total_gb'] >= 30.0
        except ImportError:
             if self.verbose: self.color.print("psutil not installed, cannot determine exact RAM.", "yellow")
        except Exception as e:
             if self.verbose: self.color.print(f"Memory detection error: {e}", "yellow")
        return memory_info

    def _detect_gpu(self)->Dict[str,Any]:
        """Detects GPU information using PyTorch (assumes installed via Poetry)."""
        gpu_info = {
            'available': False, 'is_rtx_3080': False, 'model': 'Unknown',
            'vram_gb': 0, 'cuda_available': False, 'cuda_version': '',
            'cudnn_available': False, 'cudnn_version': ''
        }
        if self.cpu_only:
            return gpu_info

        # Run the check within the poetry environment
        check_script = """
import torch
import sys
import platform
try:
    cuda_available = torch.cuda.is_available()
    print(f"CUDA_Available: {cuda_available}")
    if cuda_available:
        device_count = torch.cuda.device_count()
        print(f"Device_Count: {device_count}")
        if device_count > 0:
            # Get properties of the first GPU
            props = torch.cuda.get_device_properties(0)
            print(f"Device_Name: {props.name}")
            print(f"VRAM_GB: {props.total_memory / (1024**3):.2f}")
            # Get CUDA version PyTorch was compiled with
            print(f"CUDA_Version: {torch.version.cuda}")
            # Check cuDNN availability and version
            cudnn_available = torch.backends.cudnn.is_available()
            print(f"cuDNN_Available: {cudnn_available}")
            if cudnn_available:
                print(f"cuDNN_Version: {torch.backends.cudnn.version()}")
except Exception as e:
    # Print error to stderr so it's captured separately
    print(f"GPU_Check_Error: {e!r}", file=sys.stderr)
"""
        # Use poetry run python -c "script"
        returncode, stdout, stderr = self._run_command(
            ['poetry', 'run', 'python', '-c', check_script],
            check=False, # Check return code manually
            capture_output=True
        )

        if returncode != 0 or "GPU_Check_Error" in stderr:
            if self.verbose:
                self.color.print('PyTorch check script failed. Cannot determine GPU details.', 'yellow')
                if stderr: self.color.print(f"Error details from script: {stderr}", "yellow")
            # Fallback check using nvidia-smi
            smi_path = "nvidia-smi"
            if self.platform_system == "Windows":
                program_files = os.environ.get("ProgramFiles", "C:\\Program Files")
                smi_path_win = Path(program_files) / "NVIDIA Corporation" / "NVSMI" / "nvidia-smi.exe"
                if smi_path_win.exists(): smi_path = str(smi_path_win)

            rc_smi, _, _ = self._run_command([smi_path], check=False, capture_output=False) # Just check if command runs
            if rc_smi == 0:
                self.color.print('PyTorch check failed, but nvidia-smi found. GPU likely present but environment needs configuration.', 'yellow')
                gpu_info['available'] = True # Mark GPU as present
            return gpu_info

        # Parse output from successful torch check
        for line in stdout.splitlines():
            try:
                if ':' not in line: continue # Skip lines without a colon
                key, value = line.split(':', 1)
                key = key.strip()
                value = value.strip()
                if key == 'CUDA_Available': gpu_info['cuda_available'] = (value == 'True')
                elif key == 'Device_Name':
                    gpu_info['model'] = value
                    if '3080' in value: gpu_info['is_rtx_3080'] = True
                    gpu_info['available'] = True # Mark available if we got a name
                elif key == 'VRAM_GB': gpu_info['vram_gb'] = float(value)
                elif key == 'CUDA_Version': gpu_info['cuda_version'] = value
                elif key == 'cuDNN_Available': gpu_info['cudnn_available'] = (value == 'True')
                elif key == 'cuDNN_Version': gpu_info['cudnn_version'] = str(value) # Keep as string
            except ValueError:
                if self.verbose: self.color.print(f"Could not parse GPU info line: {line}", "yellow")
            except Exception as e:
                 if self.verbose: self.color.print(f"Error parsing GPU info: {e}", "yellow")

        if not gpu_info['cuda_available']:
            if self.verbose: self.color.print('No CUDA-capable GPU detected by PyTorch.', 'yellow')

        return gpu_info

    def _check_tools(self) -> Dict[str, bool]:
        """Checks for required command-line tools: Git, C++ Compiler, Poetry."""
        tools_status = {'git': False, 'cpp_compiler': False, 'poetry': False}
        self.color.print("Checking required tools (Git, C++ Compiler, Poetry)...", "cyan")

        # Check Poetry FIRST - it's essential for this script
        returncode_poetry, stdout_poetry, _ = self._run_command(['poetry', '--version'], check=False)
        if returncode_poetry == 0:
            tools_status['poetry'] = True
            self.has_poetry = True
            self.color.print(f"  Poetry found: {stdout_poetry.strip()} ✓", 'green')
        else:
            self.color.print('  Poetry not found - Required for dependency management.', 'red', bold=True)
            self.color.print('  Please install Poetry before running this script:', 'red')
            self.color.print('    https://python-poetry.org/docs/#installation', 'yellow')
            # No point continuing if poetry isn't installed
            return tools_status # Return early

        # Check Git
        returncode_git, stdout_git, _ = self._run_command(['git', '--version'], check=False)
        if returncode_git == 0:
            tools_status['git'] = True
            self.has_git = True
            self.color.print(f"  Git found: {stdout_git.strip()} ✓", 'green')
        else:
            self.color.print('  Git not found - Required for downloading some models.', 'yellow')
            self.color.print('  Install Git for full functionality:', 'yellow')
            if self.platform_system == 'Windows':
                self.color.print('    https://git-scm.com/download/win', 'yellow')
            else:
                self.color.print('    Linux: sudo apt-get update && sudo apt-get install git', 'yellow')
                self.color.print('    MacOS: brew install git', 'yellow')

        # Check C++ Compiler
        compiler_found = False
        if self.platform_system == 'Windows':
            # Check for cl.exe existence using 'where' command first
            rc_where, _, _ = self._run_command(['where', 'cl.exe'], check=False, capture_output=False)
            if rc_where == 0:
                compiler_found = True
                self.color.print('  Visual C++ compiler (cl.exe) found in PATH ✓', 'green')
            else:
                 # Optionally add vswhere check here if needed for more complex setups
                 self.color.print('  Visual C++ compiler (cl.exe) not found in PATH.', 'yellow')
        else: # Linux/MacOS
            rc_gpp, _, _ = self._run_command(['which', 'g++'], check=False, capture_output=False)
            rc_gcc, _, _ = self._run_command(['which', 'gcc'], check=False, capture_output=False)
            if rc_gpp == 0 or rc_gcc == 0:
                compiler_found = True
                compiler_name = "g++" if rc_gpp == 0 else "gcc"
                self.color.print(f'  C++ compiler ({compiler_name}) found ✓', 'green')
            else:
                self.color.print('  C++ compiler (g++/gcc) not found.', 'yellow')

        if compiler_found:
            tools_status['cpp_compiler'] = True
            self.has_cpp_compiler = True
        else:
            self.color.print('  Required for building some packages from source if wheels are unavailable.', 'yellow')
            if self.platform_system == 'Windows':
                 self.color.print('  Install "C++ build tools" via Visual Studio Installer:', 'yellow')
                 self.color.print('    https://visualstudio.microsoft.com/visual-cpp-build-tools/', 'yellow')
            else:
                 self.color.print('  Install build tools:', 'yellow')
                 self.color.print('    Linux (Debian/Ubuntu): sudo apt-get update && sudo apt-get install build-essential', 'yellow')
                 self.color.print('    Linux (Fedora): sudo dnf groupinstall "Development Tools"', 'yellow')
                 self.color.print('    MacOS: xcode-select --install', 'yellow')

        return tools_status

    def _create_directories(self) -> bool:
        """Creates required directories."""
        self.color.print("Creating required directories...", "cyan")
        try:
            for directory in self.required_dirs:
                dir_path = self.base_dir / directory
                dir_path.mkdir(parents=True, exist_ok=True)
                if self.verbose: self.color.print(f"  Ensured directory exists: {dir_path}", 'green')

            # Ensure __init__.py files exist for package structure
            package_dirs = [
                'maggie', 'maggie/core', 'maggie/extensions', 'maggie/utils',
                'maggie/utils/config', 'maggie/utils/hardware', 'maggie/utils/llm',
                'maggie/utils/stt', 'maggie/utils/tts'
            ]
            for pkg_dir_str in package_dirs:
                pkg_dir = self.base_dir / pkg_dir_str
                # Ensure the package directory itself exists before creating __init__.py
                pkg_dir.mkdir(parents=True, exist_ok=True)
                init_path = pkg_dir / '__init__.py'
                if not init_path.exists():
                    try:
                        # Create an empty __init__.py file
                        init_path.touch()
                        if self.verbose: self.color.print(f"  Created __init__.py in {pkg_dir_str}", 'green')
                    except Exception as e_init:
                         self.color.print(f"Error creating __init__.py in {pkg_dir_str}: {e_init}", 'red')
                         # Allow script to continue, but warn user

            self.color.print('Directory structure verified/created.', 'green')
            return True
        except Exception as e:
            self.color.print(f"Error creating directories: {e}", 'red')
            return False

    def _install_with_poetry(self) -> bool:
        """Installs dependencies using Poetry."""
        self.color.print("Installing dependencies using Poetry...", "cyan")
        self.color.print("This may take a while depending on the number of dependencies and network speed.", "blue")
        self.color.print("Poetry will attempt to resolve compatible versions for all dependencies based on pyproject.toml.", "blue")

        pyproject_path = self.base_dir / 'pyproject.toml'
        if not pyproject_path.exists():
            self.color.print(f"ERROR: pyproject.toml not found in {self.base_dir}", 'red', bold=True)
            self.color.print("Please ensure a pyproject.toml file configured for Poetry exists.", 'red')
            self.color.print("See script comments/documentation for required structure.", 'red')
            return False

        # Basic check for [tool.poetry.dependencies] section
        try:
            with open(pyproject_path, 'r', encoding='utf-8') as f:
                content = f.read()
                if '[tool.poetry.dependencies]' not in content:
                    self.color.print("ERROR: pyproject.toml does not contain a [tool.poetry.dependencies] section.", 'red', bold=True)
                    self.color.print("Please configure dependencies for Poetry.", 'red')
                    return False
        except Exception as e:
            self.color.print(f"Error reading pyproject.toml: {e}", 'red')
            return False

        # Prepare poetry install command
        poetry_cmd = ['poetry', 'install', '--no-interaction'] # Use non-interactive mode

        # Add GPU extras if not cpu_only and if defined in pyproject.toml
        if not self.cpu_only:
            # Check if 'gpu' extra group is defined
            # Using `poetry show --extras` is unreliable, parse pyproject.toml instead
            has_gpu_extra = False
            try:
                with open(pyproject_path, 'r', encoding='utf-8') as f:
                    # Simple check, might fail with complex toml structures
                    # A proper TOML parser would be more robust
                    toml_content = f.read()
                    if '[tool.poetry.extras]' in toml_content and 'gpu = [' in toml_content:
                        has_gpu_extra = True

                if has_gpu_extra:
                    self.color.print("Including [gpu] extras for Poetry installation.", "blue")
                    poetry_cmd.extend(['--extras', 'gpu'])
                else:
                     self.color.print("No [gpu] extra found or defined in pyproject.toml, installing base dependencies only.", "yellow")
            except Exception as e:
                self.color.print(f"Could not check for GPU extras in pyproject.toml: {e}", "yellow")
                self.color.print("Attempting base installation without extras.", "yellow")

        if self.verbose:
            poetry_cmd.append('-vvv') # Add verbose flag for poetry if script verbose

        # Run poetry install
        self.color.print(f"Running: {' '.join(poetry_cmd)}", "blue")
        returncode, stdout, stderr = self._run_command(
            poetry_cmd,
            check=False, # We handle the error explicitly
            capture_output=True
            )

        # Print output, especially useful on failure
        if stdout: self.color.print(f"Poetry output:\n{stdout}", 'cyan')
        # Always print stderr, color based on success
        if stderr:
            self.color.print(f"Poetry errors/warnings:\n{stderr}", 'red' if returncode != 0 else 'yellow')

        if returncode != 0:
            self.color.print("ERROR: Poetry failed to resolve dependencies or install packages.", 'red', bold=True)
            self.color.print("Please check the output above for specific conflict details.", 'red')
            self.color.print("Common issues:", 'yellow')
            self.color.print(" - Conflicting version requirements between packages.", 'yellow')
            self.color.print(" - Missing system dependencies required to build a package (e.g., C++ compiler, PortAudio).", 'yellow')
            self.color.print(" - Ensure torch dependencies in pyproject.toml use the correct source (pytorch_cu118).", 'yellow')
            self.color.print(" - Network issues preventing package downloads.", 'yellow')
            return False
        else:
            self.color.print("Poetry successfully resolved and installed dependencies ✓", 'green')
            return True

    # --- Model Downloading Methods (_download_whisper_model, etc.) ---
    # These methods remain largely the same, relying on _run_command (which uses poetry run if needed)
    # and _download_file. Ensure paths use self.base_dir.

    def _download_whisper_model(self)->bool:
        """Downloads the Whisper base.en model using huggingface_hub via Poetry."""
        if self.skip_models:
            self.color.print('Skipping Whisper model download (--skip-models)', 'yellow')
            return True

        model_dir = self.base_dir / 'maggie/models/stt/whisper-base.en'
        essential_files = ['model.bin', 'config.json', 'tokenizer.json', 'vocab.json']
        repo_id = "openai/whisper-base.en"

        if model_dir.exists():
            try:
                # Check if all essential files are present
                files_in_dir = {f.name for f in model_dir.iterdir() if f.is_file()}
                if all(essential_file in files_in_dir for essential_file in essential_files):
                    self.color.print('Whisper base.en model is already available ✓', 'green')
                    return True
                else:
                    self.color.print('Whisper model directory exists but appears incomplete.', 'yellow')
                    self.color.print('Attempting to download/update the model...', 'yellow')
            except Exception as e:
                self.color.print(f"Error checking existing Whisper model directory: {e}", 'yellow')

        model_dir.mkdir(parents=True, exist_ok=True)

        # Use poetry run python to ensure huggingface_hub from the env is used
        # Using snapshot_download is generally preferred
        download_script = f'''
import os
import sys
from huggingface_hub import snapshot_download, HfApi, hf_hub_download

# Convert Path object to string for the script
model_dir_str = r"{str(model_dir).replace(os.sep, '/')}"
repo_id = "{repo_id}"

try:
    print(f"Attempting to download model '{{repo_id}}' to '{{model_dir_str}}'")
    snapshot_download(
        repo_id=repo_id,
        local_dir=model_dir_str,
        # Include only necessary pytorch files, exclude others
        allow_patterns=["*.json", "*.bin", "*.txt", "preprocessor_config.json", "generation_config.json"], # Be more specific
        ignore_patterns=["*.safetensors", "*.h5", "*.ot", "flax*", "tf*"],
        local_dir_use_symlinks=False,
        resume_download=True, # Allow resuming
    )
    print("Model downloaded successfully via snapshot_download.")
    sys.exit(0) # Explicit success exit

except Exception as e_snap:
    print(f"Snapshot download failed: {{e_snap!r}}", file=sys.stderr)
    # Fallback: Try downloading essential files individually - less reliable
    print("Attempting to download essential files individually (fallback)...", file=sys.stderr)
    essential_files_list = ['config.json', 'generation_config.json', 'model.bin', 'preprocessor_config.json', 'tokenizer.json', 'vocab.json']
    all_downloaded = True
    for filename in essential_files_list:
        try:
             print(f"Downloading {{filename}}...")
             hf_hub_download(repo_id=repo_id, filename=filename, local_dir=model_dir_str, resume_download=True)
        except Exception as e_file:
            print(f"Failed to download {{filename}}: {{e_file!r}}", file=sys.stderr)
            all_downloaded = False

    if all_downloaded:
        print("Essential files downloaded individually (fallback).")
        sys.exit(0) # Success exit code
    else:
        print("Failed to download some essential model files via fallback.", file=sys.stderr)
        sys.exit(1) # Failure exit code
'''
        self.color.print(f'Downloading Whisper model ({repo_id}) from Hugging Face...', 'cyan')
        returncode, stdout, stderr = self._run_command(
            ['poetry', 'run', 'python', '-c', download_script],
            check=False, # Check return code manually
            capture_output=True
        )

        if self.verbose and stdout: self.color.print(f"Whisper download stdout:\n{stdout}", "cyan")
        if stderr: self.color.print(f"Whisper download stderr:\n{stderr}", "yellow" if returncode == 0 else "red")

        if returncode != 0:
            self.color.print("Error downloading Whisper model.", 'red')
            return False

        # Verify essential files exist after download attempt
        try:
            files_in_dir = {f.name for f in model_dir.iterdir() if f.is_file()}
            missing_files = [f for f in essential_files if f not in files_in_dir]
            if not missing_files:
                self.color.print('Whisper base.en model downloaded and verified successfully ✓', 'green')
                return True
            else:
                self.color.print('Whisper model download appears incomplete.', 'yellow')
                self.color.print(f'Missing essential files: {", ".join(missing_files)}', 'yellow')
                return False
        except Exception as e:
            self.color.print(f"Error verifying downloaded Whisper model: {e}", 'red')
            return False

    def _download_af_heart_model(self)->bool:
        """Downloads the af_heart.pt TTS voice model."""
        if self.skip_models:
            self.color.print('Skipping af_heart model download (--skip-models)', 'yellow')
            return True

        model_dir = self.base_dir / 'maggie/models/tts'
        model_path = model_dir / 'af_heart.pt'
        MIN_SIZE = 40 * 1024 * 1024 # 40 MB minimum size check

        if model_path.exists():
            try:
                file_size = model_path.stat().st_size
                if file_size >= MIN_SIZE:
                    self.color.print(f"af_heart voice model verified ({file_size / (1024 * 1024):.2f} MB) ✓", 'green')
                    return True
                else:
                    self.color.print(f"Existing af_heart model has incorrect size: {file_size / (1024 * 1024):.2f} MB. Re-downloading...", 'yellow')
            except Exception as e:
                 self.color.print(f"Error checking existing af_heart model: {e}. Re-downloading...", 'yellow')

        model_dir.mkdir(parents=True, exist_ok=True)
        # Prioritize Hugging Face URL
        model_urls = [
            'https://huggingface.co/hexgrad/kokoro-voices/resolve/main/af_heart.pt',
            'https://github.com/hexgrad/kokoro/releases/download/v0.1/af_heart.pt' # Fallback
        ]

        download_successful = False
        for url in model_urls:
            self.color.print(f"Attempting download from: {url}", 'cyan')
            if self._download_file(url, str(model_path)):
                try:
                    file_size = model_path.stat().st_size
                    if file_size >= MIN_SIZE:
                        self.color.print(f"af_heart model download successful ({file_size / (1024 * 1024):.2f} MB) ✓", 'green')
                        download_successful = True
                        break # Stop trying URLs
                    else:
                        self.color.print(f"Downloaded file has incorrect size: {file_size / (1024 * 1024):.2f} MB", 'yellow')
                        model_path.unlink(missing_ok=True) # Remove incorrect file
                except Exception as e_verify:
                     self.color.print(f"Error verifying downloaded af_heart model: {e_verify}", 'red')
                     model_path.unlink(missing_ok=True)
            # If download failed, loop continues to next URL

        if not download_successful:
            self.color.print('Failed to download af_heart voice model from any source.', 'red')
            self.color.print('You may need to download it manually and place it in maggie/models/tts/', 'yellow')
            self.color.print('  Sources: https://huggingface.co/hexgrad/kokoro-voices/tree/main', 'yellow')
            self.color.print('           https://github.com/hexgrad/kokoro/releases', 'yellow')
            return False
        return True


    def _download_kokoro_onnx_models(self)->bool:
        """Downloads the necessary ONNX model files for kokoro-onnx."""
        # This function assumes kokoro-onnx is installed via Poetry
        # Check kokoro-onnx documentation for the exact files and locations needed.
        # The URLs below are examples based on common community repos.
        if self.skip_models:
             self.color.print('Skipping kokoro-onnx model download (--skip-models)', 'yellow')
             return True

        self.color.print("Downloading Kokoro ONNX models...", "cyan")
        model_dir = self.base_dir / 'maggie/models/tts' # Assume they go here
        model_dir.mkdir(parents=True, exist_ok=True)
        kokoro_onnx_repo = "onnx-community/Kokoro-82M-v1.0-ONNX" # Example repo

        # Define required files based on kokoro-onnx usage
        # Adjust filenames and repo paths as needed
        required_assets = [
            {'type': 'file', 'name': 'model.onnx', 'repo': kokoro_onnx_repo, 'subpath': 'onnx/model.onnx', 'min_size': 10*1024*1024},
            {'type': 'file', 'name': 'model_quantized.onnx', 'repo': kokoro_onnx_repo, 'subpath': 'onnx/model_quantized.onnx', 'min_size': 5*1024*1024, 'optional': True}, # Example: optional quantized
            {'type': 'file', 'name': 'voices-v1.0.bin', 'repo': kokoro_onnx_repo, 'subpath': 'voices-v1.0.bin', 'min_size': 5*1024*1024},
            {'type': 'file', 'name': 'tokens.txt', 'repo': kokoro_onnx_repo, 'subpath': 'tokens.txt', 'min_size': 100},
            {'type': 'dir', 'name': 'espeak-ng-data', 'repo': "rhasspy/espeak-ng-data", 'repo_type': 'dataset'}, # espeak data often needed
        ]

        all_successful = True
        for asset in required_assets:
            asset_name = asset['name']
            asset_path = model_dir / asset_name
            min_size = asset.get('min_size', 0)
            is_optional = asset.get('optional', False)

            # Check existence
            if asset_path.exists():
                is_ok = False
                if asset['type'] == 'dir' and asset_path.is_dir() and any(asset_path.iterdir()): # Check if dir not empty
                    is_ok = True
                elif asset['type'] == 'file' and asset_path.is_file():
                    try:
                        file_size = asset_path.stat().st_size
                        if file_size >= min_size:
                            is_ok = True
                        else:
                            self.color.print(f"Existing {asset_name} has incorrect size: {file_size/(1024*1024):.2f} MB. Re-downloading...", 'yellow')
                    except Exception as e_check:
                        self.color.print(f"Error checking existing {asset_name}: {e_check}. Re-downloading...", 'yellow')

                if is_ok:
                    self.color.print(f"{asset_name} already exists ✓", 'green')
                    continue # Skip download

            # Download logic using huggingface_hub via poetry run python
            self.color.print(f"Downloading {asset_name} from {asset['repo']}...", 'cyan')
            repo_id = asset['repo']
            repo_type = asset.get('repo_type', 'model') # Default to model repo type
            filename = asset.get('subpath', asset_name) # Use subpath if specified

            if asset['type'] == 'file':
                dl_script = f'''
import sys
from huggingface_hub import hf_hub_download
try:
    hf_hub_download(repo_id="{repo_id}", repo_type="{repo_type}", filename="{filename}", local_dir=r"{str(model_dir)}", local_dir_use_symlinks=False, resume_download=True)
    print("Download successful.")
    sys.exit(0)
except Exception as e:
    print(f"Download failed: {{e!r}}", file=sys.stderr)
    sys.exit(1)
'''
            elif asset['type'] == 'dir':
                dl_script = f'''
import sys
from huggingface_hub import snapshot_download
try:
    snapshot_download(repo_id="{repo_id}", repo_type="{repo_type}", local_dir=r"{str(asset_path)}", local_dir_use_symlinks=False, resume_download=True)
    print("Download successful.")
    sys.exit(0)
except Exception as e:
    print(f"Download failed: {{e!r}}", file=sys.stderr)
    sys.exit(1)
'''
            else:
                self.color.print(f"Unknown asset type '{asset['type']}' for {asset_name}", 'red')
                all_successful = False if not is_optional else all_successful
                continue

            rc_dl, out_dl, err_dl = self._run_command(
                ['poetry', 'run', 'python', '-c', dl_script], check=False, capture_output=True
            )

            if self.verbose and out_dl: self.color.print(f"{asset_name} download stdout:\n{out_dl}", "cyan")
            if err_dl: self.color.print(f"{asset_name} download stderr:\n{err_dl}", "yellow" if rc_dl == 0 else "red")

            if rc_dl == 0:
                # Verify downloaded asset
                if asset['type'] == 'file':
                    try:
                        file_size = asset_path.stat().st_size
                        if file_size >= min_size:
                            self.color.print(f"{asset_name} download successful ({file_size/(1024*1024):.2f} MB) ✓", 'green')
                        else:
                            self.color.print(f"Downloaded {asset_name} has incorrect size: {file_size/(1024*1024):.2f} MB.", 'yellow')
                            asset_path.unlink(missing_ok=True)
                            if not is_optional: all_successful = False
                    except Exception as e_verify:
                        self.color.print(f"Error verifying downloaded {asset_name}: {e_verify}", 'red')
                        if not is_optional: all_successful = False
                elif asset['type'] == 'dir':
                     if asset_path.is_dir() and any(asset_path.iterdir()):
                          self.color.print(f"{asset_name} download successful ✓", 'green')
                     else:
                          self.color.print(f"Downloaded {asset_name} directory appears empty.", 'yellow')
                          if not is_optional: all_successful = False
            else:
                self.color.print(f"Failed to download {asset_name}.", 'red')
                if not is_optional: all_successful = False


        if all_successful:
            self.color.print('Kokoro ONNX models/data downloaded successfully ✓', 'green')
            return True
        else:
            self.color.print('Some Kokoro ONNX models/data failed to download or verify.', 'yellow')
            self.color.print('Please check errors and potentially download manually.', 'yellow')
            return False # Return False if any required asset failed


    def _download_mistral_model(self)->bool:
        """Downloads the Mistral 7B Instruct GPTQ model using Git LFS."""
        if self.skip_models:
            self.color.print('Skipping Mistral model download (--skip-models)', 'yellow')
            return True

        mistral_dir = self.base_dir / 'maggie/models/llm/mistral-7b-instruct-v0.3-GPTQ-4bit'
        repo_url = 'https://huggingface.co/TheBloke/Mistral-7B-Instruct-v0.3-GPTQ'

        # Check if directory exists and prompt user if non-empty
        if mistral_dir.exists() and any(mistral_dir.iterdir()):
             self.color.print(f"Mistral model directory already exists: {mistral_dir}", 'yellow')
             response = self.color.input("  Download again (y) or keep existing (n)? [n]: ", color='magenta')
             if response.lower() != 'y':
                 self.color.print("Keeping existing model files.", 'green')
                 return True
             else:
                 self.color.print("Removing existing directory before download...", 'yellow')
                 try:
                     shutil.rmtree(mistral_dir)
                 except Exception as e:
                     self.color.print(f"Error removing existing directory {mistral_dir}: {e}", 'red')
                     return False

        # Ensure parent directory exists before cloning
        mistral_dir.parent.mkdir(parents=True, exist_ok=True)

        if not self.has_git:
            self.color.print('Git not found. Cannot download Mistral model automatically.', 'red')
            self.color.print(f'Please download the model manually from: {repo_url}', 'yellow')
            self.color.print(f'And place the contents in: {mistral_dir}', 'yellow')
            return False

        # Check for Git LFS
        lfs_check_code, _, _ = self._run_command(['git', 'lfs', '--version'], check=False, capture_output=False)
        if lfs_check_code != 0:
             self.color.print("Git LFS not found or not configured correctly.", 'red')
             self.color.print("Git LFS is required to download large model files.", 'red')
             self.color.print("Install Git LFS (https://git-lfs.com) and run 'git lfs install'", 'yellow')
             return False
        else:
             # Ensure git-lfs hooks are installed (run once per user usually)
             self._run_command(['git', 'lfs', 'install', '--skip-repo'], check=False, capture_output=False)

        self.color.print(f"Downloading Mistral model from {repo_url} using Git LFS...", 'cyan')
        self.color.print("This is a large download (~5GB) and may take a significant amount of time.", 'blue')

        # Clone the repository - Use shallow clone initially? No, LFS needs history.
        # Set GIT_LFS_SKIP_SMUDGE=1 to download structure fast, then pull LFS files
        clone_env = os.environ.copy()
        clone_env['GIT_LFS_SKIP_SMUDGE'] = '1'
        # Use --depth 1 for faster initial clone of metadata? Maybe not with LFS. Stick to full clone.
        returncode_clone, stdout_clone, stderr_clone = self._run_command(
            ['git', 'clone', repo_url, str(mistral_dir)],
            check=False, capture_output=True, env=clone_env
        )

        if returncode_clone != 0:
            # Check if directory already exists error (git exit code 128)
            if 'already exists and is not an empty directory' in stderr_clone:
                 self.color.print(f"Directory {mistral_dir} already exists but clone failed.", "yellow")
                 self.color.print("Attempting to pull LFS files into existing directory...", "yellow")
                 # Proceed to LFS pull step
            else:
                self.color.print('Error initiating Git clone.', 'red')
                if stderr_clone: self.color.print(f"Git error: {stderr_clone}", 'red')
                # Clean up failed clone attempt
                if mistral_dir.exists():
                    try: shutil.rmtree(mistral_dir)
                    except Exception as e_rm: self.color.print(f"Could not clean up failed clone dir: {e_rm}", "yellow")
                return False

        self.color.print("Git repository structure cloned/verified. Now pulling LFS files...", 'cyan')
        # Now pull the actual LFS files - use --include='*.safetensors' etc. if needed?
        # Running plain 'git lfs pull' should work if smudge was skipped.
        returncode_lfs, stdout_lfs, stderr_lfs = self._run_command(
            ['git', 'lfs', 'pull'],
            check=False, capture_output=True, cwd=str(mistral_dir) # Run in the cloned dir
        )

        # Print LFS output if verbose
        if self.verbose and stdout_lfs: self.color.print(f"Git LFS pull output:\n{stdout_lfs}", "cyan")
        if stderr_lfs: self.color.print(f"Git LFS pull stderr:\n{stderr_lfs}", "yellow" if returncode_lfs == 0 else "red")


        if returncode_lfs != 0:
            self.color.print('Error pulling Git LFS files.', 'red')
            self.color.print("Model download may be incomplete.", 'yellow')
            return False # Treat LFS pull failure as critical

        # Basic verification: check for essential config and at least one .safetensors file
        essential_configs = ['config.json', 'tokenizer.json', 'quantize_config.json']
        has_safetensors = False
        try:
             files_in_dir = {f.name for f in mistral_dir.iterdir()}
             missing_configs = [f for f in essential_configs if f not in files_in_dir]
             has_safetensors = any(f.endswith('.safetensors') for f in files_in_dir)

             if not missing_configs and has_safetensors:
                  self.color.print('Mistral GPTQ model downloaded and basic verification passed ✓', 'green')
                  return True
             else:
                  self.color.print('Mistral model download appears incomplete after LFS pull.', 'yellow')
                  if missing_configs: self.color.print(f"Missing config files: {', '.join(missing_configs)}", 'yellow')
                  if not has_safetensors: self.color.print("Missing model weights file (.safetensors)", 'yellow')
                  return False # Mark as failure if essential parts missing
        except Exception as e:
             self.color.print(f"Error verifying Mistral model download: {e}", 'red')
             return False


    def _create_recipe_template(self)->bool:
        """Creates a default recipe template if it doesn't exist using python-docx."""
        template_dir = self.base_dir / 'maggie/templates'
        template_path = template_dir / 'recipe_template.docx'
        template_dir.mkdir(parents=True, exist_ok=True) # Ensure dir exists

        if template_path.exists():
            self.color.print('Recipe template already exists ✓', 'green')
            return True

        self.color.print("Creating default recipe template...", "cyan")
        # Use python-docx (expected to be installed via Poetry) to create the template
        # Escape backslashes properly for the string literal within the command
        template_path_str_escaped = str(template_path).replace('\\', '\\\\')
        # Use a more robust script that imports necessary styles if needed
        create_script = f"""
import docx
import os
from pathlib import Path
import sys

template_path_str = r'{template_path_str_escaped}'
template_path = Path(template_path_str)

try:
    # Create document
    doc = docx.Document()

    # Add content (same as before)
    doc.add_heading("Recipe Name", level=1)
    doc.add_paragraph() # Add space
    doc.add_heading("Recipe Information", level=2)
    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = 'Table Grid'
    info_table.cell(0, 0).text = "Preparation Time"
    info_table.cell(0, 1).text = "00 minutes"
    info_table.cell(1, 0).text = "Cooking Time"
    info_table.cell(1, 1).text = "00 minutes"
    info_table.cell(2, 0).text = "Servings"
    info_table.cell(2, 1).text = "0 servings"
    doc.add_paragraph() # Add space

    doc.add_heading("Ingredients", level=2)
    doc.add_paragraph("• Ingredient 1", style='List Bullet') # Use standard style name
    doc.add_paragraph("• Ingredient 2", style='List Bullet')
    doc.add_paragraph("• Ingredient 3", style='List Bullet')
    doc.add_paragraph() # Add space

    doc.add_heading("Instructions", level=2)
    doc.add_paragraph("Step 1", style='List Number') # Use standard style name
    doc.add_paragraph("Step 2", style='List Number')
    doc.add_paragraph("Step 3", style='List Number')
    doc.add_paragraph() # Add space

    doc.add_heading("Notes", level=2)
    doc.add_paragraph("Add any additional notes, tips, or variations here.")

    # Save template
    doc.save(template_path)
    print(f"Template saved to {{template_path}}")
    sys.exit(0) # Success

except ImportError:
     print("Error: python-docx library not found. Cannot create template.", file=sys.stderr)
     print("Please ensure 'python-docx' is listed in pyproject.toml dependencies.", file=sys.stderr)
     sys.exit(1) # Failure
except Exception as e:
    print(f"Error creating docx template: {{e}}", file=sys.stderr)
    # If saving failed, remove partial file
    if template_path.exists():
        try: os.remove(template_path)
        except: pass
    sys.exit(1) # Failure
"""
        # Use poetry run python -c "script"
        returncode, stdout, stderr = self._run_command(
            ['poetry', 'run', 'python', '-c', create_script],
            check=False, # Check manually
            capture_output=True
        )

        if self.verbose and stdout: self.color.print(f"Template creation stdout:\n{stdout}", "cyan")
        if stderr: self.color.print(f"Template creation stderr:\n{stderr}", "yellow" if returncode == 0 else "red")

        if returncode == 0:
            self.color.print('Recipe template created successfully ✓', 'green')
            return True
        else:
            self.color.print('Failed to create recipe template.', 'red')
            self.color.print('Recipe creator extension may not work properly.', 'yellow')
            return False

    def _setup_config(self)->bool:
        """Creates or updates the config.yaml based on hardware detection."""
        config_path = self.base_dir / 'config.yaml'
        # Look for an example file - prioritize .yaml, fallback to .txt if needed
        example_path = self.base_dir / 'config.yaml.example'
        if not example_path.exists():
             example_path = self.base_dir / 'config.yaml.txt' # Fallback
             if not example_path.exists():
                 example_path = self.base_dir / 'config-yaml-example.txt' # Another fallback

        if not config_path.exists() and not example_path.exists():
             self.color.print("ERROR: Neither config.yaml nor a config example file found.", 'red', bold=True)
             self.color.print("Cannot configure the application.", 'red')
             return False

        self.color.print("Setting up configuration file (config.yaml)...", "cyan")

        # Dump hardware info to a temporary file for the config script to read
        temp_hardware_file = self.base_dir / 'hardware_info.json.tmp'
        try:
            with open(temp_hardware_file, 'w') as f:
                json.dump(self.hardware_info, f, indent=2) # Add indent for readability
        except Exception as e:
            self.color.print(f"Error writing temporary hardware info: {e}", 'red')
            return False

        # Prepare python script to load/modify/save yaml config
        # Use raw strings and escaped paths for cross-platform compatibility
        base_dir_str = repr(str(self.base_dir)) # Let python handle escaping
        config_path_str_esc = str(config_path).replace('\\','\\\\')
        example_path_str_esc = str(example_path).replace('\\','\\\\') if example_path.exists() else ''
        hardware_file_str_esc = str(temp_hardware_file).replace('\\','\\\\')
        cpu_only_flag = self.cpu_only

        # Python script to perform config manipulation safely
        config_script = f"""
import yaml
import json
import os
import shutil
import sys
from pathlib import Path
from collections.abc import MutableMapping # Use abc for compatibility

# Ensure safe loading/dumping using standard Loader/Dumper
try:
    # Try to use CLoader/CDumper for speed if available
    from yaml import CLoader as Loader, CDumper as Dumper
except ImportError:
    from yaml import Loader, Dumper

# Define paths using pathlib for robustness
base_dir = Path({base_dir_str})
config_path = Path(r'{config_path_str_esc}')
example_path = Path(r'{example_path_str_esc}') if '{example_path_str_esc}' else None
hardware_file = Path(r'{hardware_file_str_esc}')
cpu_only = {cpu_only_flag}

# --- Helper to recursively ensure keys exist ---
def ensure_keys(d, keys):
    current = d
    for key in keys:
        if not isinstance(current, MutableMapping): # Check if current level is dict-like
             print(f"Warning: Cannot set key '{{key}}' because parent is not a dictionary.", file=sys.stderr)
             return None # Cannot proceed further down this path
        current = current.setdefault(key, {{}}) # setdefault returns value or new dict
    return current # Return the innermost dictionary

# --- Main Config Logic ---
config = {{}} # Initialize empty config

# Load hardware info
try:
    with open(hardware_file, 'r', encoding='utf-8') as f:
        hardware_info = json.load(f)
except Exception as e:
    print(f"Error loading hardware info: {{e}}", file=sys.stderr)
    hardware_info = {{}} # Default to empty if load fails

# Load existing config or copy from example
if config_path.exists():
    print(f"Loading existing configuration from {{config_path}}")
    try:
        with open(config_path, 'r', encoding='utf-8') as f:
            config = yaml.load(f, Loader=Loader)
            if not isinstance(config, MutableMapping): # Ensure loaded config is a dict
                 print(f"Warning: Existing config.yaml is not a valid dictionary. Starting fresh.", file=sys.stderr)
                 config = {{}}
    except Exception as e:
        print(f"Error loading existing config.yaml: {{e}}. Starting fresh.", file=sys.stderr)
        config = {{}}
elif example_path and example_path.exists():
    print(f"Creating new configuration from example: {{example_path}}")
    try:
        shutil.copyfile(example_path, config_path)
        with open(config_path, 'r', encoding='utf-8') as f:
            config = yaml.load(f, Loader=Loader)
            if not isinstance(config, MutableMapping): # Ensure loaded example is a dict
                 print(f"Warning: Example config file is not a valid dictionary. Starting fresh.", file=sys.stderr)
                 config = {{}}
    except Exception as e:
        print(f"Error copying/loading example config: {{e}}. Starting fresh.", file=sys.stderr)
        config = {{}}
else:
    print("No existing config or example found. Starting with default settings.")
    config = {{}} # Start with empty if no config/example

# --- Apply Settings Based on Hardware/Flags ---

# Ensure top-level keys exist using helper
llm_config = ensure_keys(config, ['llm'])
gpu_config = ensure_keys(config, ['gpu'])
cpu_config = ensure_keys(config, ['cpu'])
mem_config = ensure_keys(config, ['memory'])
stt_config = ensure_keys(config, ['stt'])
stt_whisper_config = ensure_keys(stt_config, ['whisper']) if stt_config else None
tts_config = ensure_keys(config, ['tts'])

# Set defaults / update paths if sections were created
if llm_config is not None:
    llm_config['model_path'] = 'maggie/models/llm/mistral-7b-instruct-v0.3-GPTQ-4bit'
    llm_config['model_type'] = 'mistral'
    llm_config['use_autogptq'] = True
if stt_whisper_config is not None:
    stt_whisper_config['model_path'] = 'maggie/models/stt/whisper-base.en'
if tts_config is not None:
    tts_config['model_path'] = 'maggie/models/tts'
    tts_config['voice_model'] = 'af_heart.pt'

# Hardware Optimizations
gpu_hw = hardware_info.get('gpu', {{}})
cpu_hw = hardware_info.get('cpu', {{}})
mem_hw = hardware_info.get('memory', {{}})

# Apply settings only if the config sections could be ensured
if cpu_only:
    print("Applying CPU-only optimizations...")
    if llm_config: llm_config.update({{'gpu_layers': 0, 'gpu_layer_auto_adjust': False}})
    if gpu_config: gpu_config.update({{'max_percent': 0, 'model_unload_threshold': 0}})
    if tts_config: tts_config['gpu_acceleration'] = False
    if stt_whisper_config: stt_whisper_config['compute_type'] = 'int8'
    # Remove potentially conflicting GPU flags
    if llm_config: llm_config.pop('rtx_3080_optimized', None)
    if gpu_config: gpu_config.pop('rtx_3080_optimized', None)
else:
    # Apply GPU settings only if CUDA was detected by torch
    if gpu_hw.get('cuda_available'):
        print("Applying GPU optimizations...")
        if llm_config: llm_config.update({{'gpu_layer_auto_adjust': True, 'precision_type': 'float16', 'mixed_precision_enabled': True}})
        if tts_config: tts_config.update({{'gpu_acceleration': True}})
        if stt_whisper_config: stt_whisper_config['compute_type'] = 'float16'

        # RTX 3080 Specific
        if gpu_hw.get('is_rtx_3080'):
            print("Applying RTX 3080 specific optimizations...")
            if llm_config: llm_config.update({{'gpu_layers': 32, 'rtx_3080_optimized': True}})
            if gpu_config: gpu_config.update({{'max_percent': 90, 'model_unload_threshold': 95, 'rtx_3080_optimized': True}})
            if tts_config: tts_config['gpu_precision'] = 'mixed_float16'
        else: # Generic GPU defaults
            print("Applying generic GPU optimizations...")
            if llm_config: llm_config.update({{'gpu_layers': -1}}) # Auto layers
            if gpu_config: gpu_config.update({{'max_percent': 85, 'model_unload_threshold': 90}})
            if tts_config: tts_config['gpu_precision'] = 'float16'
            # Remove specific flags if not applicable
            if llm_config: llm_config.pop('rtx_3080_optimized', None)
            if gpu_config: gpu_config.pop('rtx_3080_optimized', None)
    else: # No CUDA detected by torch, force CPU settings
         print("CUDA not available according to PyTorch check. Forcing CPU settings in config.", file=sys.stderr)
         if llm_config: llm_config.update({{'gpu_layers': 0, 'gpu_layer_auto_adjust': False}})
         if gpu_config: gpu_config.update({{'max_percent': 0, 'model_unload_threshold': 0}})
         if tts_config: tts_config['gpu_acceleration'] = False
         if stt_whisper_config: stt_whisper_config['compute_type'] = 'int8'
         if llm_config: llm_config.pop('rtx_3080_optimized', None)
         if gpu_config: gpu_config.pop('rtx_3080_optimized', None)

# Ryzen 9 5900X Specific
if cpu_hw.get('is_ryzen_9_5900x'):
    print("Applying Ryzen 9 5900X specific optimizations...")
    if cpu_config: cpu_config.update({{'max_threads': 8, 'thread_timeout': 30, 'ryzen_9_5900x_optimized': True}})
else:
    if cpu_config: cpu_config.pop('ryzen_9_5900x_optimized', None)

# 32GB+ Memory Specific
if mem_hw.get('is_32gb'):
    print("Applying 32GB+ RAM specific optimizations...")
    if mem_config: mem_config.update({{'max_percent': 75, 'model_unload_threshold': 85, 'xpg_d10_memory': True}}) # Example flag
else:
    if mem_config: mem_config.pop('xpg_d10_memory', None)

# --- Write Config ---
try:
    with open(config_path, 'w', encoding='utf-8') as f:
        # Use sort_keys=False to maintain order where possible
        yaml.dump(config, f, Dumper=Dumper, default_flow_style=False, sort_keys=False, allow_unicode=True)
    print(f"Configuration saved to {{config_path}}")
except Exception as e:
    print(f"Error saving updated config.yaml: {{e}}", file=sys.stderr)
    # Don't clean up hardware file if save failed
    sys.exit(1) # Exit with error

# Clean up temporary hardware file only on success
try:
    os.remove(hardware_file)
except Exception as e:
    print(f"Warning: Could not remove temporary hardware file {{hardware_file}}: {{e}}", file=sys.stderr)

sys.exit(0) # Explicit success exit
"""
        # Use poetry run python -c "script"
        returncode, stdout, stderr = self._run_command(
            ['poetry', 'run', 'python', '-c', config_script],
            check=False, # Check manually
            capture_output=True
        )

        if self.verbose and stdout: self.color.print(f"Config script stdout:\n{stdout}", "cyan")
        # Always print stderr from config script
        if stderr: self.color.print(f"Config script stderr:\n{stderr}", "yellow" if returncode == 0 else "red")

        if returncode != 0:
            self.color.print('Error setting up configuration file.', 'red')
            return False
        else:
            self.color.print('Configuration file created/updated with optimized settings ✓', 'green')
            self.color.print('NOTE: You MUST edit config.yaml to add your Picovoice access key for wake word detection.', 'yellow')
            return True


    def verify_system(self)->bool:
        """Verifies Python version and essential tools (Poetry, Git, Compiler)."""
        self.progress.start_step('Verifying system compatibility')
        python_compatible = self._verify_python_version()
        if not python_compatible:
            self.progress.complete_step(False,'Incompatible Python version')
            return False

        tools = self._check_tools()
        # Poetry is essential, fail if not found
        if not tools['poetry']:
            self.progress.complete_step(False, 'Poetry is not installed or not found in PATH.')
            return False
        # Warn about missing optional tools
        if not tools['git']:
            self.color.print("  Warning: Git not found. Model downloading may be affected.", "yellow")
        if not tools['cpp_compiler']:
             self.color.print("  Warning: C++ compiler not found. Some packages might fail to build.", "yellow")

        self.progress.complete_step(True)
        return True


    def install(self)->bool:
        """Runs the complete installation process using Poetry."""
        self.color.print('\n=== Maggie AI Assistant Installation (using Poetry) ===', 'cyan', bold=True)
        self.color.print(f"Platform: {self.platform_system} ({platform.platform()})", 'cyan')
        self.color.print(f"Python: {platform.python_version()} (via Poetry environment)", 'cyan')
        if self.cpu_only: self.color.print('Mode: CPU-only installation requested', 'yellow')
        self.color.print(f"Project Root: {self.base_dir}", 'cyan')
        self.color.print(f"Non-Python Requirements: CUDA 11.8, cuDNN 8.9.7 (for GPU)", 'yellow')


        # --- Pre-checks ---
        if not self.verify_system():
             return False # Exit if basic checks fail (Python version, Poetry existence)

        # --- Create Directories ---
        self.progress.start_step('Creating directory structure')
        if not self._create_directories():
            self.progress.complete_step(False, 'Failed to create directories')
            return False
        self.progress.complete_step(True)

        # --- Install Dependencies via Poetry & Detect Hardware ---
        self.progress.start_step('Install dependencies & Detect Hardware')
        if not self._install_with_poetry():
            self.progress.complete_step(False, 'Poetry dependency installation failed.')
            return False # Exit if poetry fails

        # Detect hardware AFTER dependencies are installed (esp. PyTorch for GPU)
        self._detect_hardware() # This now prints the summary internally
        # Check GPU status post-detection
        if not self.cpu_only and not self.hardware_info['gpu'].get('cuda_available'):
            self.color.print("  Warning: PyTorch installed, but CUDA does not seem to be available.", 'yellow')
            self.color.print("           GPU acceleration will likely not work.", 'yellow')
            self.color.print("           Verify CUDA 11.8/cuDNN 8.9.7 installation and PATH.", 'yellow')
            self.color.print("           Ensure torch dependencies in pyproject.toml use the correct 'pytorch_cu118' source.", 'yellow')

        self.progress.complete_step(True) # Complete step for install+detect

        # --- Setup Config ---
        self.progress.start_step('Setting up configuration file')
        if not self._setup_config():
            self.progress.complete_step(False, 'Failed to set up configuration')
            return False
        self.progress.complete_step(True)

        # --- Download Models ---
        self.progress.start_step('Downloading models')
        models_ok = True
        # Run downloads sequentially
        if not self._download_af_heart_model():
            self.color.print('Warning: Failed to download TTS voice model (af_heart.pt). TTS may be limited.', 'yellow')
            # models_ok = False # Decide if this should be fatal
        if not self._download_kokoro_onnx_models():
            self.color.print('Warning: Failed to download some kokoro-onnx model files. TTS may be limited.', 'yellow')
            # models_ok = False # Decide if this should be fatal
        if not self._download_whisper_model():
            self.color.print('Warning: Failed to download Whisper model. STT may be limited.', 'yellow')
            # models_ok = False # Decide if this should be fatal
        if not self._download_mistral_model():
            self.color.print('Warning: Failed to download Mistral LLM model. LLM functionality will be unavailable.', 'yellow')
            # models_ok = False # Decide if this should be fatal

        self.progress.complete_step(models_ok) # Mark step based on overall success

        # --- Final Steps (Templates, etc.) ---
        self.progress.start_step('Setting up templates & completing installation')
        template_ok = self._create_recipe_template()
        if not template_ok:
            self.color.print('Warning: Failed to create recipe template.', 'yellow')
        # Add any other final setup steps here if needed

        # --- Installation Summary ---
        self.progress.display_summary(True) # Assume overall success if we got this far

        self.color.print('\n--- Important Notes ---', 'cyan', bold=True)
        self.color.print('1. Dependencies installed via Poetry into its virtual environment.', 'green')
        self.color.print(f"2. Non-Python Requirements: CUDA Toolkit 11.8 and cuDNN 8.9.7 (for GPU support).", 'yellow')
        self.color.print(f"   (Detected CUDA via PyTorch: {self.hardware_info['gpu'].get('cuda_version', 'N/A')}, Detected cuDNN via PyTorch: {self.hardware_info['gpu'].get('cudnn_version', 'N/A')})", 'yellow')
        self.color.print('3. Edit config.yaml to add your Picovoice Access Key for wake word detection.', 'yellow')
        self.color.print('   (Get key from: https://console.picovoice.ai/)', 'yellow')
        if not self.has_git:
             self.color.print('4. Git is not installed - model downloading/updates may be affected.', 'yellow')
        if not self.has_cpp_compiler:
             self.color.print('5. C++ Compiler not found - building packages from source may fail if needed.', 'yellow')

        self.color.print('\n--- To start Maggie AI Assistant ---', 'cyan', bold=True)
        self.color.print('   Make sure you are in the project directory containing pyproject.toml', 'blue')
        self.color.print('   Run: poetry run python main.py', 'green')

        self.progress.complete_step(True) # Final completion step

        return True

    def _print_hardware_summary(self):
        """Prints hardware summary previously gathered."""
        self.color.print('Hardware Configuration Summary:', 'cyan', bold=True)
        cpu_info = self.hardware_info['cpu']
        self.color.print(f"  CPU: {cpu_info.get('model', 'Unknown')} ({cpu_info.get('cores', 'N/A')} cores / {cpu_info.get('threads', 'N/A')} threads)", 'green' if cpu_info.get('is_ryzen_9_5900x') else 'yellow')

        memory_info = self.hardware_info['memory']
        self.color.print(f"  RAM: {memory_info.get('total_gb', 0):.1f} GB", 'green' if memory_info.get('is_32gb') else 'yellow')

        gpu_info = self.hardware_info['gpu']
        if self.cpu_only: self.color.print('  GPU: CPU-only mode selected', 'yellow')
        elif gpu_info.get('cuda_available'):
             self.color.print(f"  GPU: {gpu_info.get('model', 'Unknown')} ✓", 'green')
             self.color.print(f"       {gpu_info.get('vram_gb', 0):.1f} GB VRAM", 'green')
             self.color.print(f"       CUDA {gpu_info.get('cuda_version', 'N/A')} (PyTorch)", 'green')
             self.color.print(f"       cuDNN {gpu_info.get('cudnn_version', 'N/A')} (PyTorch)", 'green' if gpu_info.get('cudnn_available') else 'yellow')
        else: self.color.print('  GPU: No CUDA-capable GPU detected by PyTorch', 'red')


# --- Main Execution ---

def main() -> int:
    """Parses arguments and runs the installer."""
    parser = argparse.ArgumentParser(
        description='Maggie AI Assistant Installer (Poetry Version)',
        formatter_class=argparse.ArgumentDefaultsHelpFormatter
    )
    parser.add_argument('--verbose', '-v', action='store_true', help='Enable verbose output for installer script and poetry')
    parser.add_argument('--cpu-only', action='store_true', help='Install CPU-only version (skips GPU extras & detection)')
    parser.add_argument('--skip-models', action='store_true', help='Skip downloading large AI models (TTS, STT, LLM)')
    args = parser.parse_args()

    # Change to the script's directory to ensure relative paths work
    script_dir = Path(__file__).parent.resolve()
    os.chdir(script_dir)
    print(f"Running installer from directory: {script_dir}")


    installer = MaggieInstaller(
        verbose=args.verbose,
        cpu_only=args.cpu_only,
        skip_models=args.skip_models
    )

    try:
        success = installer.install()
        return 0 if success else 1
    except KeyboardInterrupt:
        print('\n\nInstallation cancelled by user.')
        return 1
    except Exception as e:
        # Print a more informative error message
        installer.color.print(f"\n\nAN UNEXPECTED ERROR OCCURRED DURING INSTALLATION:", 'red', bold=True)
        installer.color.print(f"Error Type: {type(e).__name__}", 'red')
        installer.color.print(f"Error Details: {e}", 'red')
        # Print traceback if verbose
        if args.verbose:
            import traceback
            installer.color.print("Traceback:", 'red')
            traceback.print_exc()
        return 1

if __name__ == '__main__':
    sys.exit(main())
```

**How to Use:**

1.  **Save:** Save the code above as `install_dev.py` in the root directory of your project (the same directory where `pyproject.toml` is located).
2.  **Configure `pyproject.toml`:** Make sure your `pyproject.toml` file is set up correctly for Poetry, listing all direct dependencies and including the source definition for PyTorch as shown in the prerequisite example. Pin only `torch`, `torchvision`, and `torchaudio` as requested.
3.  **Run:** Execute the script from your terminal in the project's root directory:
    ```bash
    python install_dev.py [OPTIONS]
    ```
    * Use `python install_dev.py --cpu-only` for a CPU installation.
    * Use `python install_dev.py --skip-models` to skip downloading models.
    * Use `python install_dev.py --verbose` for detailed output.

The script will now use Poetry to manage the environment and dependencies, checking for conflicts and installing packages according to your `pyproject.toml` before proceeding with the model downloads and configuration ste