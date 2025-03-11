"""
Maggie AI Assistant - Recipe Creator Utility
=========================================
Speech-to-document recipe creation utility.

This module provides a streamlined workflow for creating recipe documents
from speech input, with specific optimizations for AMD Ryzen 9 5900X
and NVIDIA RTX 3080 hardware.
"""

import os
import time
import threading
from enum import Enum, auto
from dataclasses import dataclass, field
from typing import Dict, Any, Optional, List, Tuple
import docx
from loguru import logger

from utils.utility_base import UtilityBase

class RecipeState(Enum):
    """
    States for the recipe creation process.
    
    Simplified state machine for the recipe creation workflow.
    """
    INITIAL = auto()      # Initial state
    NAME_INPUT = auto()   # Getting recipe name
    DESCRIPTION = auto()  # Getting recipe description
    PROCESSING = auto()   # Processing with LLM
    CREATING = auto()     # Creating document
    COMPLETED = auto()    # Process completed
    CANCELLED = auto()    # Process cancelled

@dataclass
class RecipeData:
    """
    Data structure for recipe information.
    
    Parameters
    ----------
    name : str
        Recipe name
    description : str
        Raw recipe description
    ingredients : List[str]
        Parsed list of ingredients
    steps : List[str]
        Parsed list of preparation steps
    notes : str
        Additional notes
    """
    name: str = ""
    description: str = ""
    ingredients: List[str] = field(default_factory=list)
    steps: List[str] = field(default_factory=list)
    notes: str = ""

class RecipeCreator(UtilityBase):
    """
    Recipe Creator utility for creating and formatting recipes from speech.
    
    A streamlined utility for creating recipe documents from speech input,
    with specific optimizations for processing speed and document formatting.
    
    Parameters
    ----------
    event_bus : EventBus
        Reference to the central event bus
    config : Dict[str, Any]
        Configuration dictionary
    """
    
    def __init__(self, event_bus, config: Dict[str, Any]):
        """
        Initialize the Recipe Creator utility.
        
        Parameters
        ----------
        event_bus : EventBus
            Reference to the central event bus
        config : Dict[str, Any]
            Configuration dictionary
        """
        super().__init__(event_bus, config)
        
        # Initialize attributes
        self.state = RecipeState.INITIAL
        self.recipe_data = RecipeData()
        self.output_dir = config.get("output_dir", "recipes")
        self.template_path = config.get("template_path", "templates/recipe_template.docx")
        
        # Processing thread
        self._workflow_thread = None
        
        # Component references - will be set during initialize()
        self.speech_processor = None
        self.llm_processor = None
        
        # Ensure directories exist
        os.makedirs(self.output_dir, exist_ok=True)
        os.makedirs(os.path.dirname(self.template_path), exist_ok=True)
        
        # Create template if it doesn't exist
        if not os.path.exists(self.template_path):
            self._create_template()
    
    def get_trigger(self) -> str:
        """
        Get the trigger phrase for this utility.
        
        Returns
        -------
        str
            Trigger phrase that activates this utility
        """
        return "new recipe"
    
    def initialize(self) -> bool:
        """
        Initialize the Recipe Creator.
        
        Acquires references to required components.
        
        Returns
        -------
        bool
            True if initialization successful, False otherwise
        """
        try:
            # Find the main app with speech and LLM processors
            for component in self.event_bus.subscribers.get("state_changed", []):
                if hasattr(component, "speech_processor") and hasattr(component, "llm_processor"):
                    self.speech_processor = component.speech_processor
                    self.llm_processor = component.llm_processor
                    break
            
            # Check if components were found
            if not self.speech_processor or not self.llm_processor:
                logger.error("Failed to acquire speech or LLM processor references")
                return False
            
            return True
            
        except Exception as e:
            logger.error(f"Error initializing Recipe Creator: {e}")
            return False
    
    def start(self) -> bool:
        """
        Start the Recipe Creator workflow.
        
        Returns
        -------
        bool
            True if started successfully, False otherwise
        """
        try:
            # Reset state and data
            self.state = RecipeState.INITIAL
            self.recipe_data = RecipeData()
            
            # Check if already running
            if self.running:
                logger.warning("Recipe Creator already running")
                return False
            
            # Initialize components if needed
            if not self.initialized and not self.initialize():
                logger.error("Failed to initialize Recipe Creator")
                return False
            
            # Start workflow thread
            self._workflow_thread = threading.Thread(
                target=self._workflow,
                name="RecipeWorkflow"
            )
            self._workflow_thread.daemon = True
            self._workflow_thread.start()
            
            self.running = True
            logger.info("Recipe Creator started")
            return True
            
        except Exception as e:
            logger.error(f"Error starting Recipe Creator: {e}")
            return False
    
    def stop(self) -> bool:
        """
        Stop the Recipe Creator.
        
        Returns
        -------
        bool
            True if stopped successfully, False otherwise
        """
        if not self.running:
            return True
            
        try:
            # Set state to CANCELLED to stop workflow
            self.state = RecipeState.CANCELLED
            self.running = False
            
            # Wait for thread to finish
            if self._workflow_thread and self._workflow_thread.is_alive():
                self._workflow_thread.join(timeout=2.0)
            
            logger.info("Recipe Creator stopped")
            return True
            
        except Exception as e:
            logger.error(f"Error stopping Recipe Creator: {e}")
            return False
    
    def process_command(self, command: str) -> bool:
        """
        Process a command directed to this utility.
        
        Parameters
        ----------
        command : str
            Command string to process
            
        Returns
        -------
        bool
            True if command processed, False otherwise
        """
        if not self.running:
            return False
            
        command = command.lower().strip()
        
        # Handle cancel command
        if "cancel" in command:
            logger.info("Recipe creation cancelled by user")
            self.state = RecipeState.CANCELLED
            return True
            
        # Handle name confirmation
        if self.state == RecipeState.NAME_INPUT:
            if command in ["yes", "correct", "right", "yeah"]:
                logger.info(f"Recipe name confirmed: {self.recipe_data.name}")
                self.state = RecipeState.DESCRIPTION
                return True
            elif command in ["no", "wrong", "incorrect", "nope"]:
                logger.info("Recipe name rejected, asking again")
                self.recipe_data.name = ""
                return True
                
        return False
    
    def _workflow(self) -> None:
        """
        Main recipe creation workflow.
        
        Manages the state machine for the recipe creation process.
        """
        try:
            # Welcome message
            self.speech_processor.speak("Starting recipe creator. Let's create a new recipe.")
            
            # State machine
            while self.running and self.state != RecipeState.COMPLETED and self.state != RecipeState.CANCELLED:
                # Process current state
                if self.state == RecipeState.INITIAL:
                    # Start getting the recipe name
                    self.state = RecipeState.NAME_INPUT
                    
                elif self.state == RecipeState.NAME_INPUT:
                    # Get recipe name
                    if not self.recipe_data.name:
                        self.speech_processor.speak("What would you like to name this recipe?")
                        success, name = self.speech_processor.recognize_speech(timeout=10.0)
                        
                        if success and name:
                            self.recipe_data.name = name
                            self.speech_processor.speak(f"I heard {name}. Is that correct?")
                        else:
                            self.speech_processor.speak("I didn't catch that. Let's try again.")
                    
                    # Name confirmation is handled in process_command()
                    time.sleep(0.5)
                    
                elif self.state == RecipeState.DESCRIPTION:
                    # Get recipe description
                    self.speech_processor.speak("Please describe the recipe, including ingredients and steps.")
                    success, description = self.speech_processor.recognize_speech(timeout=30.0)
                    
                    if success and description:
                        self.recipe_data.description = description
                        self.speech_processor.speak("Got it. Processing your recipe.")
                        self.state = RecipeState.PROCESSING
                    else:
                        self.speech_processor.speak("I didn't catch that. Let's try again.")
                    
                elif self.state == RecipeState.PROCESSING:
                    # Process with LLM
                    success = self._process_with_llm()
                    
                    if success:
                        self.state = RecipeState.CREATING
                    else:
                        self.speech_processor.speak("There was an error processing your recipe. Let's try describing it again.")
                        self.state = RecipeState.DESCRIPTION
                    
                elif self.state == RecipeState.CREATING:
                    # Create document
                    success = self._create_document()
                    
                    if success:
                        self.speech_processor.speak(f"Recipe '{self.recipe_data.name}' has been created and saved.")
                        self.state = RecipeState.COMPLETED
                    else:
                        self.speech_processor.speak("There was an error creating the document.")
                        self.state = RecipeState.CANCELLED
                
                # Brief pause to prevent tight loop
                time.sleep(0.1)
            
            # Finalize
            if self.state == RecipeState.CANCELLED:
                self.speech_processor.speak("Recipe creation cancelled.")
                
            # Publish completion event
            self.event_bus.publish("utility_completed", "recipe_creator")
            
        except Exception as e:
            logger.error(f"Error in recipe workflow: {e}")
            self.speech_processor.speak("An error occurred while creating the recipe.")
            self.event_bus.publish("utility_error", "recipe_creator")
            
        finally:
            self.running = False
    
    def _process_with_llm(self) -> bool:
        """
        Process recipe description with LLM.
        
        Uses the language model to extract structured recipe information.
        
        Returns
        -------
        bool
            True if processing successful, False otherwise
        """
        try:
            # Create prompt for LLM
            prompt = f"""
            Parse the following recipe description into structured format:
            
            Recipe: {self.recipe_data.description}
            
            Extract and format as follows:
            
            INGREDIENTS:
            - [ingredient with quantity]
            - [ingredient with quantity]
            ...
            
            STEPS:
            1. [step 1]
            2. [step 2]
            ...
            
            NOTES:
            [any additional notes or tips]
            """
            
            # Process with LLM
            response = self.llm_processor.generate_text(prompt, max_tokens=1024)
            
            if not response:
                logger.error("Empty response from LLM")
                return False
                
            # Parse response
            current_section = None
            ingredients = []
            steps = []
            notes = []
            
            for line in response.strip().split('\n'):
                line = line.strip()
                
                if not line:
                    continue
                    
                if line == "INGREDIENTS:":
                    current_section = "ingredients"
                elif line == "STEPS:":
                    current_section = "steps"
                elif line == "NOTES:":
                    current_section = "notes"
                elif current_section == "ingredients" and line.startswith("-"):
                    ingredients.append(line[1:].strip())
                elif current_section == "steps" and (line[0].isdigit() and "." in line[:3]):
                    steps.append(line[line.find(".")+1:].strip())
                elif current_section == "notes":
                    notes.append(line)
            
            # Update recipe data
            self.recipe_data.ingredients = ingredients
            self.recipe_data.steps = steps
            self.recipe_data.notes = "\n".join(notes)
            
            logger.info(f"Recipe processed: {len(ingredients)} ingredients, {len(steps)} steps")
            return True
            
        except Exception as e:
            logger.error(f"Error processing recipe with LLM: {e}")
            return False
    
    def _create_document(self) -> bool:
        """
        Create recipe document.
        
        Generates a formatted document using python-docx.
        
        Returns
        -------
        bool
            True if document created successfully, False otherwise
        """
        try:
            # Load template or create new document
            if os.path.exists(self.template_path):
                doc = docx.Document(self.template_path)
            else:
                doc = docx.Document()
                self._create_template()
                doc = docx.Document(self.template_path)
            
            # Clear template content
            for paragraph in doc.paragraphs:
                if paragraph.text and paragraph.text.strip():
                    paragraph.text = ""
            
            # Add recipe title
            doc.add_heading(self.recipe_data.name, level=1)
            
            # Add ingredients section
            doc.add_heading("Ingredients", level=2)
            for ingredient in self.recipe_data.ingredients:
                doc.add_paragraph(f"• {ingredient}", style='ListBullet')
            
            # Add steps section
            doc.add_heading("Instructions", level=2)
            for i, step in enumerate(self.recipe_data.steps, 1):
                doc.add_paragraph(f"{i}. {step}", style='ListNumber')
            
            # Add notes section if available
            if self.recipe_data.notes:
                doc.add_heading("Notes", level=2)
                doc.add_paragraph(self.recipe_data.notes)
            
            # Create safe filename
            safe_name = "".join(c if c.isalnum() or c in " -_" else "_" for c in self.recipe_data.name)
            filename = f"{safe_name}_{int(time.time())}.docx"
            filepath = os.path.join(self.output_dir, filename)
            
            # Save document
            doc.save(filepath)
            
            logger.info(f"Recipe document saved to {filepath}")
            return True
            
        except Exception as e:
            logger.error(f"Error creating recipe document: {e}")
            return False
    
    def _create_template(self) -> bool:
        """
        Create recipe template document.
        
        Generates a template document for recipes if it doesn't exist.
        
        Returns
        -------
        bool
            True if template created successfully, False otherwise
        """
        try:
            # Create template directory if it doesn't exist
            os.makedirs(os.path.dirname(self.template_path), exist_ok=True)
            
            # Create template document
            doc = docx.Document()
            
            # Add title placeholder
            doc.add_heading("Recipe Name", level=1)
            
            # Add metadata section
            doc.add_heading("Recipe Information", level=2)
            table = doc.add_table(rows=3, cols=2)
            table.style = 'Table Grid'
            
            # Set table headers and defaults
            cells = table.rows[0].cells
            cells[0].text = "Preparation Time"
            cells[1].text = "00 minutes"
            
            cells = table.rows[1].cells
            cells[0].text = "Cooking Time"
            cells[1].text = "00 minutes"
            
            cells = table.rows[2].cells
            cells[0].text = "Servings"
            cells[1].text = "0 servings"
            
            # Add ingredients section
            doc.add_heading("Ingredients", level=2)
            doc.add_paragraph("• Ingredient 1", style='ListBullet')
            doc.add_paragraph("• Ingredient 2", style='ListBullet')
            doc.add_paragraph("• Ingredient 3", style='ListBullet')
            
            # Add instructions section
            doc.add_heading("Instructions", level=2)
            doc.add_paragraph("1. Step 1", style='ListNumber')
            doc.add_paragraph("2. Step 2", style='ListNumber')
            doc.add_paragraph("3. Step 3", style='ListNumber')
            
            # Add notes section
            doc.add_heading("Notes", level=2)
            doc.add_paragraph("Add any additional notes, tips, or variations here.")
            
            # Save template
            doc.save(self.template_path)
            
            logger.info(f"Recipe template created at {self.template_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error creating recipe template: {e}")
            return False