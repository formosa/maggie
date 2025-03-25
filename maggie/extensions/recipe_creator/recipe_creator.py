import os,time,threading
from enum import Enum,auto
from dataclasses import dataclass,field
from typing import Dict,Any,Optional,List,Tuple,Union
import docx
from maggie.extensions.base import ExtensionBase
from maggie.core.state import State,StateTransition,StateAwareComponent
from maggie.core.event import EventListener,EventEmitter,EventPriority
from maggie.utils.error_handling import safe_execute,retry_operation,ErrorCategory,ErrorSeverity,with_error_handling,record_error,ExtensionError
from maggie.utils.logging import ComponentLogger,log_operation,logging_context
from maggie.service.locator import ServiceLocator
__all__=['RecipeState','RecipeData','RecipeCreator']
class RecipeState(Enum):INITIAL=auto();NAME_INPUT=auto();DESCRIPTION=auto();PROCESSING=auto();CREATING=auto();COMPLETED=auto();CANCELLED=auto();ERROR=auto()
@dataclass
class RecipeData:name:str='';description:str='';ingredients:List[str]=field(default_factory=list);steps:List[str]=field(default_factory=list);notes:str=''
class RecipeCreator(ExtensionBase,StateAwareComponent,EventListener):
	def __init__(self,event_bus,config:Dict[str,Any]):
		ExtensionBase.__init__(self,event_bus,config);self.state_manager=ServiceLocator.get('state_manager')
		if self.state_manager:StateAwareComponent.__init__(self,self.state_manager)
		EventListener.__init__(self,event_bus);self.state=RecipeState.INITIAL;self.recipe_data=RecipeData();self.output_dir=config.get('output_dir','recipes');self.template_path=config.get('template_path','templates/recipe_template.docx');self._retry_count=0;self._max_retries=config.get('max_retries',3);self.speech_timeout=config.get('speech_timeout',3e1);self._workflow_thread=None;self.stt_processor=None;self.llm_processor=None;self.tts_processor=None;self.logger=ComponentLogger('RecipeCreator');self._register_event_handlers();self._ensure_directories()
	def _register_event_handlers(self)->None:
		event_handlers=[('error_logged',self._handle_error,EventPriority.HIGH),('low_memory_warning',self._handle_memory_warning,EventPriority.NORMAL)]
		for(event_type,handler,priority)in event_handlers:self.listen(event_type,handler,priority=priority)
		self.logger.debug(f"Registered {len(event_handlers)} event handlers")
	def _handle_error(self,error_data:Dict[str,Any])->None:
		if isinstance(error_data,dict)and error_data.get('source')=='RecipeCreator':
			self.logger.error(f"Error in recipe creator: {error_data.get('message')}")
			if self.state not in[RecipeState.ERROR,RecipeState.CANCELLED]:self.state=RecipeState.ERROR
	def _handle_memory_warning(self,event_data:Dict[str,Any])->None:self.logger.warning('Low memory warning received - optimizing resource usage');self._cleanup_resources()
	def _ensure_directories(self)->None:
		try:
			os.makedirs(self.output_dir,exist_ok=True);directory=os.path.dirname(self.template_path)
			if directory:os.makedirs(directory,exist_ok=True)
			if not os.path.exists(self.template_path):self._create_template()
		except Exception as e:record_error(message=f"Error creating directories: {e}",exception=e,category=ErrorCategory.SYSTEM,severity=ErrorSeverity.ERROR,source='RecipeCreator._ensure_directories')
	def get_trigger(self)->str:return'new recipe'
	@log_operation(component='RecipeCreator')
	def initialize(self)->bool:
		if self._initialized:return True
		with logging_context(component='RecipeCreator',operation='initialize')as ctx:
			try:
				success=self._acquire_component_references()
				if not success:self.logger.error('Failed to acquire required component references');return False
				self._initialized=True;self.logger.info('Recipe Creator initialized successfully');return True
			except Exception as e:record_error(message=f"Error initializing Recipe Creator: {e}",exception=e,category=ErrorCategory.EXTENSION,severity=ErrorSeverity.ERROR,source='RecipeCreator.initialize');return False
	def _acquire_component_references(self)->bool:self.stt_processor=self.get_service('stt_processor');self.llm_processor=self.get_service('llm_processor');self.tts_processor=self.get_service('tts_processor');return self.stt_processor is not None and self.llm_processor is not None and self.tts_processor is not None
	@log_operation(component='RecipeCreator')
	@with_error_handling(error_category=ErrorCategory.EXTENSION)
	def start(self)->bool:
		self.state=RecipeState.INITIAL;self.recipe_data=RecipeData()
		if self.running:self.logger.warning('Recipe Creator already running');return False
		if not self.initialized and not self.initialize():self.logger.error('Failed to initialize Recipe Creator');return False
		self._workflow_thread=threading.Thread(target=self._workflow,name='RecipeWorkflow',daemon=True);self._workflow_thread.start()
		if self.state_manager:
			current_state=self.state_manager.get_current_state()
			if current_state in[State.IDLE,State.READY]:self.state_manager.transition_to(State.ACTIVE,'recipe_creator_started')
		self.running=True;self.logger.info('Recipe Creator started');return True
	@log_operation(component='RecipeCreator')
	@with_error_handling(error_category=ErrorCategory.EXTENSION)
	def stop(self)->bool:
		if not self.running:return True
		self.state=RecipeState.CANCELLED;self.running=False
		if self._workflow_thread and self._workflow_thread.is_alive():self._workflow_thread.join(timeout=2.)
		self._cleanup_resources()
		if self.state_manager:
			current_state=self.state_manager.get_current_state()
			if current_state==State.ACTIVE:self.state_manager.transition_to(State.READY,'recipe_creator_stopped')
		self.logger.info('Recipe Creator stopped');return True
	def process_command(self,command:str)->bool:
		if not self.running:return False
		command=command.lower().strip()
		if'cancel'in command or command in['stop','quit','exit']:self.logger.info('Recipe creation cancelled by user');self.state=RecipeState.CANCELLED;return True
		if self.state==RecipeState.NAME_INPUT:
			if command in['yes','correct','right','yeah','yep','sure','okay']:self.logger.info(f"Recipe name confirmed: {self.recipe_data.name}");self.state=RecipeState.DESCRIPTION;return True
			elif command in['no','wrong','incorrect','nope']:self.logger.info('Recipe name rejected, asking again');self.recipe_data.name='';return True
		return False
	def _cleanup_resources(self)->None:
		if hasattr(self,'stt_processor')and self.stt_processor:
			try:self.stt_processor.stop_listening()
			except Exception as e:self.logger.debug(f"Error stopping speech processor: {e}")
	@with_error_handling(error_category=ErrorCategory.EXTENSION)
	def _workflow(self)->None:
		try:
			self._speak("Starting recipe creator. Let's create a new recipe.")
			while self.running and self.state not in[RecipeState.COMPLETED,RecipeState.CANCELLED]:
				if self.state==RecipeState.INITIAL:self.state=RecipeState.NAME_INPUT
				elif self.state==RecipeState.NAME_INPUT:self._handle_name_input_state()
				elif self.state==RecipeState.DESCRIPTION:self._handle_description_state()
				elif self.state==RecipeState.PROCESSING:self._handle_processing_state()
				elif self.state==RecipeState.CREATING:self._handle_creating_state()
				elif self.state==RecipeState.ERROR:self._handle_error_state()
				time.sleep(.1)
			self._finalize_workflow()
		except Exception as e:record_error(message=f"Error in recipe workflow: {e}",exception=e,category=ErrorCategory.EXTENSION,severity=ErrorSeverity.ERROR,source='RecipeCreator._workflow');self._speak('An error occurred while creating the recipe.');self.event_bus.publish('extension_error','recipe_creator')
		finally:self._cleanup_resources();self.running=False
	def _handle_name_input_state(self)->None:
		if not self.recipe_data.name:
			self._speak('What would you like to name this recipe?');success,name=self._recognize_speech(timeout=1e1)
			if success and name:self.recipe_data.name=name;self._speak(f"I heard {name}. Is that correct?")
			else:self._speak("I didn't catch that. Let's try again.")
	def _handle_description_state(self)->None:
		self._speak("Please describe the recipe in detail, including ingredients with quantities, preparation steps in order, and any tips or variations. I'll listen for up to 30 seconds, so take your time.");success,description=self._recognize_speech(timeout=self.speech_timeout)
		if success and description:self.recipe_data.description=description;self._speak('Got it. Processing your recipe.');self.state=RecipeState.PROCESSING
		else:
			self._retry_count+=1
			if self._retry_count<=self._max_retries:self._speak("I didn't catch that. Let's try again.")
			else:self._speak("I'm having trouble understanding. Let's try again later.");self.state=RecipeState.CANCELLED
	def _handle_processing_state(self)->None:success=self._process_with_llm();self.state=RecipeState.CREATING if success else RecipeState.ERROR
	def _handle_creating_state(self)->None:
		success=self._create_document()
		if success:self._speak(f"Recipe '{self.recipe_data.name}' has been created and saved.");self.state=RecipeState.COMPLETED
		else:self._speak('There was an error creating the document.');self.state=RecipeState.CANCELLED
	def _handle_error_state(self)->None:
		self._retry_count+=1
		if self._retry_count<=self._max_retries:self._speak("I encountered an issue. Let's try again.");self.state=RecipeState.DESCRIPTION
		else:self._speak("I'm having trouble completing this recipe after several attempts. Let's try again later.");self.state=RecipeState.CANCELLED
	def _finalize_workflow(self)->None:
		if self.state==RecipeState.CANCELLED:self._speak('Recipe creation cancelled.')
		elif self.state==RecipeState.COMPLETED:self._speak('Recipe creation completed successfully.')
		self.event_bus.publish('extension_completed','recipe_creator')
		if self.state_manager:
			current_state=self.state_manager.get_current_state()
			if current_state==State.ACTIVE:self.state_manager.transition_to(State.READY,'recipe_creator_finished')
	@with_error_handling(error_category=ErrorCategory.EXTENSION)
	def _recognize_speech(self,timeout:float=1e1)->Tuple[bool,str]:
		if self.state_manager:
			current_state=self.state_manager.get_current_state()
			if current_state!=State.ACTIVE:self.state_manager.transition_to(State.ACTIVE,'recipe_speech_recognition')
		try:
			if self.stt_processor and not getattr(self.stt_processor,'listening',False):
				try:self.stt_processor.start_listening();time.sleep(.5)
				except Exception as listen_error:self.logger.error(f"Error starting listening: {listen_error}")
			if self.stt_processor:return self.stt_processor.recognize_speech(timeout=timeout)
			else:self.logger.error('Speech processor not available');return False,''
		except Exception as e:record_error(message=f"Error recognizing speech: {e}",exception=e,category=ErrorCategory.PROCESSING,severity=ErrorSeverity.ERROR,source='RecipeCreator._recognize_speech');return False,''
	@log_operation(component='RecipeCreator')
	@with_error_handling(error_category=ErrorCategory.PROCESSING)
	def _process_with_llm(self)->bool:
		if self.state_manager:
			current_state=self.state_manager.get_current_state()
			if current_state!=State.BUSY:self.state_manager.transition_to(State.BUSY,'recipe_llm_processing')
		with logging_context(component='RecipeCreator',operation='process_with_llm')as ctx:
			try:
				prompt=self._create_llm_prompt();response=self.llm_processor.generate_text(prompt=prompt,max_tokens=1024,temperature=.7,top_p=.95)
				if not response:self.logger.error('Empty response from LLM');return False
				self._parse_llm_response(response);self.logger.info(f"Recipe processed: {len(self.recipe_data.ingredients)} ingredients, {len(self.recipe_data.steps)} steps");return True
			except Exception as e:record_error(message=f"Error processing recipe with LLM: {e}",exception=e,category=ErrorCategory.PROCESSING,severity=ErrorSeverity.ERROR,source='RecipeCreator._process_with_llm');return False
	def _create_llm_prompt(self)->str:return f"""
        Parse the following recipe description into structured format:
        
        Recipe: {self.recipe_data.description}
        
        Extract and format as follows:
        
        INGREDIENTS:
        - [ingredient with quantity]
        - [ingredient with quantity]
        ...
        
        STEPS:
        1. [step 1]
        2. [step 2]
        ...
        
        NOTES:
        [any additional notes or tips]
        """
	def _parse_llm_response(self,response:str)->None:
		current_section=None;ingredients=[];steps=[];notes=[]
		for line in response.strip().split('\n'):
			line=line.strip()
			if not line:continue
			if'INGREDIENTS:'in line:current_section='ingredients';continue
			elif'STEPS:'in line:current_section='steps';continue
			elif'NOTES:'in line:current_section='notes';continue
			if current_section=='ingredients'and(line.startswith('-')or line.startswith('•')):ingredients.append(line[1:].strip())
			elif current_section=='steps'and(line[0].isdigit()and'.'in line[:3]):steps.append(line[line.find('.')+1:].strip())
			elif current_section=='steps'and line.startswith('Step '):steps.append(line[line.find(' ')+1:].strip())
			elif current_section=='notes':notes.append(line)
		self.recipe_data.ingredients=ingredients;self.recipe_data.steps=steps;self.recipe_data.notes='\n'.join(notes)
		if not ingredients:self._extract_ingredients_fallback(response)
	def _extract_ingredients_fallback(self,response:str)->None:
		ingredients=[]
		for line in response.strip().split('\n'):
			if any(unit in line.lower()for unit in['cup','tbsp','tsp','gram','oz']):
				if line not in ingredients:ingredients.append(line.strip())
		if ingredients:self.recipe_data.ingredients=ingredients
	@log_operation(component='RecipeCreator')
	@with_error_handling(error_category=ErrorCategory.PROCESSING)
	def _create_document(self)->bool:
		try:
			if os.path.exists(self.template_path):doc=docx.Document(self.template_path)
			else:self._create_template();doc=docx.Document(self.template_path)
			self._populate_document(doc);filepath=self._save_document(doc);self.logger.info(f"Recipe document saved to {filepath}");return True
		except Exception as e:record_error(message=f"Error creating recipe document: {e}",exception=e,category=ErrorCategory.PROCESSING,severity=ErrorSeverity.ERROR,source='RecipeCreator._create_document');return False
	def _populate_document(self,doc:docx.Document)->None:
		for paragraph in doc.paragraphs:
			if paragraph.text and paragraph.text.strip():paragraph.text=''
		doc.add_heading(self.recipe_data.name,level=1);doc.add_heading('Ingredients',level=2)
		for ingredient in self.recipe_data.ingredients:doc.add_paragraph(f"• {ingredient}",style='ListBullet')
		doc.add_heading('Instructions',level=2)
		for(i,step)in enumerate(self.recipe_data.steps,1):doc.add_paragraph(f"{i}. {step}",style='ListNumber')
		if self.recipe_data.notes:doc.add_heading('Notes',level=2);doc.add_paragraph(self.recipe_data.notes)
	def _save_document(self,doc:docx.Document)->str:safe_name=''.join(c if c.isalnum()or c in' -_'else'_'for c in self.recipe_data.name);filename=f"{safe_name}_{int(time.time())}.docx";filepath=os.path.join(self.output_dir,filename);os.makedirs(self.output_dir,exist_ok=True);doc.save(filepath);return filepath
	@log_operation(component='RecipeCreator')
	@with_error_handling(error_category=ErrorCategory.SYSTEM)
	def _create_template(self)->bool:
		try:
			directory=os.path.dirname(self.template_path)
			if directory:os.makedirs(directory,exist_ok=True)
			doc=docx.Document();doc.add_heading('Recipe Name',level=1);doc.add_heading('Recipe Information',level=2);table=doc.add_table(rows=3,cols=2);table.style='Table Grid';cells=table.rows[0].cells;cells[0].text='Preparation Time';cells[1].text='00 minutes';cells=table.rows[1].cells;cells[0].text='Cooking Time';cells[1].text='00 minutes';cells=table.rows[2].cells;cells[0].text='Servings';cells[1].text='0 servings';doc.add_heading('Ingredients',level=2);doc.add_paragraph('• Ingredient 1',style='ListBullet');doc.add_paragraph('• Ingredient 2',style='ListBullet');doc.add_paragraph('• Ingredient 3',style='ListBullet');doc.add_heading('Instructions',level=2);doc.add_paragraph('1. Step 1',style='ListNumber');doc.add_paragraph('2. Step 2',style='ListNumber');doc.add_paragraph('3. Step 3',style='ListNumber');doc.add_heading('Notes',level=2);doc.add_paragraph('Add any additional notes, tips, or variations here.');doc.save(self.template_path);self.logger.info(f"Recipe template created at {self.template_path}");return True
		except Exception as e:record_error(message=f"Error creating recipe template: {e}",exception=e,category=ErrorCategory.SYSTEM,severity=ErrorSeverity.ERROR,source='RecipeCreator._create_template');return False
	def _speak(self,text:str)->None:
		if self.tts_processor:
			try:self.tts_processor.speak(text)
			except Exception as e:self.logger.error(f"Error in speech synthesis: {e}");self.logger.info(f"Would have said: {text}")
		else:self.logger.warning(f"TTS processor unavailable: {text}")
	def on_enter_active(self,transition:StateTransition)->None:self.logger.debug('System entered ACTIVE state')
	def on_exit_active(self,transition:StateTransition)->None:self.logger.debug('System exited ACTIVE state')
	def on_enter_busy(self,transition:StateTransition)->None:self.logger.debug('System entered BUSY state')
	def on_exit_busy(self,transition:StateTransition)->None:self.logger.debug('System exited BUSY state')