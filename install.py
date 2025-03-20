#!/usr/bin/env python3
"""
Maggie AI Assistant - Comprehensive Installation Script
======================================================

Complete installation system for Maggie AI Assistant, optimized for
AMD Ryzen 9 5900X and NVIDIA RTX 3080 hardware.

This script automates the entire installation process including:
1. System verification and compatibility checks
2. Environment setup and directory structure creation 
3. Dependencies installation with platform-specific optimizations
4. Model downloads and configuration
5. Extension installation 
6. Performance tuning for the target hardware

Features:
- Cross-platform compatibility (Windows and Linux)
- Hardware detection and automatic optimization
- Modular installation with customizable components
- Progress tracking and detailed logging
- Error recovery and graceful failure handling

Examples
--------
Standard installation:
    $ python install.py

Verbose installation:
    $ python install.py --verbose

CPU-only installation (no GPU acceleration):
    $ python install.py --cpu-only

Skip large model downloads:
    $ python install.py --skip-models

Force reinstallation of already installed packages:
    $ python install.py --force-reinstall

"""

# Standard library imports
import argparse
import json
import os
import platform
import shutil
import subprocess
import sys
import time
import urllib.request
import zipfile
from pathlib import Path
from typing import Dict, Any, List, Tuple, Optional, Union, Callable

class ColorOutput:
    """
    Utility class for colorized terminal output.
    
    Parameters
    ----------
    force_enable : bool, optional
        Force enable colors even if terminal doesn't support them, by default False
    
    Attributes
    ----------
    enabled : bool
        Whether color output is enabled
    colors : Dict[str, str]
        Dictionary mapping color names to ANSI color codes
        
    Notes
    -----
    Provides ANSI color codes and formatting for terminal output,
    with automatic detection of terminal compatibility. Available colors:
    - red: Error messages
    - green: Success messages
    - yellow: Warning messages
    - blue: Information messages
    - magenta: Prompt messages
    - cyan: Section headers
    - white: Normal output
    
    Automatically detects color support based on platform and terminal
    capabilities, with options to force enable or disable colors.
    """
    
    def __init__(self, force_enable: bool = False):
        """
        Initialize color output with optional force enable.
        
        Parameters
        ----------
        force_enable : bool, optional
            Force enable colors even if terminal doesn't support them, by default False
        """
        # Auto-detect color support or force enable
        self.enabled = force_enable or self._supports_color()
        
        # Initialize color codes if enabled
        if self.enabled:
            self.colors = {
                "reset": "\033[0m",
                "bold": "\033[1m",
                "red": "\033[91m",
                "green": "\033[92m",
                "yellow": "\033[93m",
                "blue": "\033[94m",
                "magenta": "\033[95m",
                "cyan": "\033[96m",
                "white": "\033[97m",
            }
        else:
            self.colors = {color: "" for color in [
                "reset", "bold", "red", "green", "yellow",
                "blue", "magenta", "cyan", "white"
            ]}
    
    def _supports_color(self) -> bool:
        """
        Detect if the terminal supports color output.
        
        Returns
        -------
        bool
            True if color is supported, False otherwise
        """
        # Windows 10+ supports ANSI colors
        if platform.system() == "Windows":
            if int(platform.release()) >= 10:
                return True
            return False
        
        # Check for NO_COLOR environment variable
        if os.environ.get("NO_COLOR"):
            return False
            
        # Check for FORCE_COLOR environment variable
        if os.environ.get("FORCE_COLOR"):
            return True
            
        # Check if output is a TTY
        return hasattr(sys.stdout, "isatty") and sys.stdout.isatty()
    
    def print(self, message: str, color: Optional[str] = None, bold: bool = False):
        """
        Print a message with color and formatting.
        
        Parameters
        ----------
        message : str
            Message to print
        color : Optional[str], optional
            Color name from available colors, by default None
        bold : bool, optional
            Whether to print in bold, by default False
        """
        formatted = message
        
        if self.enabled:
            if bold and "bold" in self.colors:
                formatted = f"{self.colors['bold']}{formatted}"
            if color and color in self.colors:
                formatted = f"{self.colors[color]}{formatted}"
            if (bold or color) and "reset" in self.colors:
                formatted = f"{formatted}{self.colors['reset']}"
                
        print(formatted)
        
    def input(self, prompt: str, color: Optional[str] = None, bold: bool = False) -> str:
        """
        Get user input with colored prompt.
        
        Parameters
        ----------
        prompt : str
            Input prompt
        color : Optional[str], optional
            Color name from available colors, by default None
        bold : bool, optional
            Whether the prompt should be bold, by default False
            
        Returns
        -------
        str
            User input string
        """
        formatted = prompt
        
        if self.enabled:
            if bold and "bold" in self.colors:
                formatted = f"{self.colors['bold']}{formatted}"
            if color and color in self.colors:
                formatted = f"{self.colors[color]}{formatted}"
            if (bold or color) and "reset" in self.colors:
                formatted = f"{formatted}{self.colors['reset']}"
                
        return input(formatted)


class ProgressTracker:
    """
    Progress tracking utility with progress bar display.
    
    Parameters
    ----------
    color : ColorOutput
        Color output utility for formatted console output
    total_steps : int, optional
        Total number of installation steps, by default 10
    
    Attributes
    ----------
    color : ColorOutput
        Color output utility
    total_steps : int
        Total number of installation steps
    current_step : int
        Current installation step (0-based)
    start_time : float
        Start time of the installation (Unix timestamp)
    
    Notes
    -----
    Provides utilities for tracking and displaying installation progress,
    including progress bars, step completion indicators, and timing.
    
    The class supports:
    - Step-by-step progress tracking with numbered steps
    - Visual progress bars for file downloads and operations
    - Elapsed time tracking and reporting
    - Success/failure status indicators with color coding
    - Installation summary with total time elapsed
    """
    
    def __init__(self, color: ColorOutput, total_steps: int = 10):
        """
        Initialize progress tracker with total steps.
        
        Parameters
        ----------
        color : ColorOutput
            Color output utility
        total_steps : int, optional
            Total number of installation steps, by default 10
        """
        self.color = color
        self.total_steps = total_steps
        self.current_step = 0
        self.start_time = time.time()
        
    def start_step(self, step_name: str):
        """
        Start a new installation step.
        
        Parameters
        ----------
        step_name : str
            Name of the installation step
        """
        self.current_step += 1
        elapsed = time.time() - self.start_time
        
        self.color.print(
            f"\n[{self.current_step}/{self.total_steps}] {step_name} "
            f"(Elapsed: {elapsed:.1f}s)",
            color="cyan", bold=True
        )
        
    def complete_step(self, success: bool = True, message: Optional[str] = None):
        """
        Mark the current step as complete.
        
        Parameters
        ----------
        success : bool, optional
            Whether the step completed successfully, by default True
        message : Optional[str], optional
            Optional completion message, by default None
        """
        if success:
            status = "✓ Complete"
            color = "green"
        else:
            status = "✗ Failed"
            color = "red"
            
        msg = f"  {status}"
        if message:
            msg += f": {message}"
            
        self.color.print(msg, color=color)
        
    def progress_bar(self, iteration: int, total: int, prefix: str = '', 
                   suffix: str = '', decimals: int = 1, length: int = 50,
                   fill: str = '█', print_end: str = '\r'):
        """
        Display a progress bar in the console.
        
        Parameters
        ----------
        iteration : int
            Current iteration (0-based)
        total : int
            Total iterations
        prefix : str, optional
            Prefix string, by default ''
        suffix : str, optional
            Suffix string, by default ''
        decimals : int, optional
            Decimal places for percentage, by default 1
        length : int, optional
            Character length of bar, by default 50
        fill : str, optional
            Bar fill character, by default '█'
        print_end : str, optional
            End character (e.g. "\r", "\n"), by default '\r'
        """
        percent = ("{0:." + str(decimals) + "f}").format(100 * (iteration / float(total)))
        filled_length = int(length * iteration // total)
        bar = fill * filled_length + '-' * (length - filled_length)
        print(f'\r{prefix} |{bar}| {percent}% {suffix}', end=print_end)
        # Print a newline on complete
        if iteration == total:
            print()
            
    def elapsed_time(self) -> float:
        """
        Get elapsed time since installation started.
        
        Returns
        -------
        float
            Elapsed time in seconds
        """
        return time.time() - self.start_time
        
    def display_summary(self, success: bool = True):
        """
        Display installation summary.
        
        Parameters
        ----------
        success : bool, optional
            Whether the installation was successful, by default True
        """
        elapsed = self.elapsed_time()
        
        if success:
            status = "Installation Completed Successfully"
            color = "green"
        else:
            status = "Installation Completed with Errors"
            color = "yellow"
            
        self.color.print(
            f"\n=== {status} ===",
            color=color, bold=True
        )
        self.color.print(f"Total time: {elapsed:.1f} seconds")


class MaggieInstaller:
    """
    Comprehensive installer for Maggie AI Assistant.
    
    Parameters
    ----------
    verbose : bool
        Whether to display verbose output
    cpu_only : bool
        Whether to install CPU-only version (no GPU acceleration)
    skip_models : bool
        Whether to skip downloading large LLM models
    skip_problematic : bool
        Whether to skip problematic dependencies that may cause installation issues
    force_reinstall : bool
        Whether to force reinstallation of already installed packages
    
    Attributes
    ----------
    verbose : bool
        Verbose output flag
    cpu_only : bool
        CPU-only mode flag
    skip_models : bool
        Skip model downloads flag
    skip_problematic : bool
        Skip problematic packages flag
    force_reinstall : bool
        Force reinstallation flag
    base_dir : Path
        Base installation directory
    platform_system : str
        Operating system name ('Windows', 'Linux', etc.)
    platform_machine : str
        Machine architecture ('x86_64', 'arm64', etc.)
    required_dirs : List[str]
        List of required directories to create
    color : ColorOutput
        Color output utility
    progress : ProgressTracker
        Installation progress tracker
    is_admin : bool
        Whether the script is running with admin/root privileges
    has_git : bool
        Whether Git is installed and available
    has_cpp_compiler : bool
        Whether a C++ compiler is installed and available
    hardware_info : Dict[str, Any]
        Detected hardware information
    total_steps : int
        Total number of installation steps
        
    Notes
    -----
    Handles all aspects of installation including system verification,
    dependency management, configuration setup, and extension installation.
    Optimized for AMD Ryzen 9 5900X and NVIDIA RTX 3080 hardware.
    
    The installer creates the complete directory structure as defined in
    self.required_dirs, including all necessary subdirectories for:
    - Core functionality
    - Extensions
    - Models (LLM, STT, TTS)
    - Utilities (config, hardware, llm, stt, tts)
    - Templates
    - Cache directories
    
    It handles platform-specific requirements and optimizations for
    both Windows and Linux operating systems.
    """
    
    def __init__(self, verbose: bool = False, cpu_only: bool = False,
                 skip_models: bool = False, skip_problematic: bool = False,
                 force_reinstall: bool = False):
        """
        Initialize installer with configuration options.
        
        Parameters
        ----------
        verbose : bool, optional
            Whether to display verbose output, by default False
        cpu_only : bool, optional
            Whether to install CPU-only version (no GPU acceleration), by default False
        skip_models : bool, optional
            Whether to skip downloading large LLM models, by default False
        skip_problematic : bool, optional
            Whether to skip problematic dependencies, by default False
        force_reinstall : bool, optional
            Whether to force reinstallation of packages, by default False
        """
        # Configuration options
        self.verbose = verbose
        self.cpu_only = cpu_only
        self.skip_models = skip_models
        self.skip_problematic = skip_problematic
        self.force_reinstall = force_reinstall
        
        # Base installation path
        self.base_dir = Path(os.path.dirname(os.path.abspath(__file__)))
        
        # Platform information
        self.platform_system = platform.system()
        self.platform_machine = platform.machine()
        
        # Required directories for installation matching the ASCII tree
        self.required_dirs = [
            "downloads",
            "logs",
            "maggie",
            "maggie/cache",
            "maggie/cache/tts",
            "maggie/core",
            "maggie/extensions",
            "maggie/models",
            "maggie/models/llm",
            "maggie/models/stt",
            "maggie/models/tts",
            "maggie/templates",
            "maggie/templates/extension",
            "maggie/utils",
            "maggie/utils/hardware",
            "maggie/utils/config",
            "maggie/utils/llm",
            "maggie/utils/stt",
            "maggie/utils/tts",
        ]
        
        # Utilities
        self.color = ColorOutput()
        
        # Total installation steps
        self.total_steps = 8
        self.progress = ProgressTracker(self.color, self.total_steps)
        
        # System capability flags
        self.is_admin = self._check_admin_privileges()
        self.has_git = False
        self.has_cpp_compiler = False
        
        # Hardware information
        self.hardware_info = {
            "cpu": {
                "is_ryzen_9_5900x": False,
                "model": "",
                "cores": 0,
                "threads": 0
            },
            "gpu": {
                "is_rtx_3080": False,
                "model": "",
                "vram_gb": 0,
                "cuda_available": False,
                "cuda_version": ""
            },
            "memory": {
                "total_gb": 0,
                "available_gb": 0,
                "is_32gb": False
            }
        }
    
    def _check_admin_privileges(self) -> bool:
        """
        Check if the script is running with administrative privileges.
        
        Returns
        -------
        bool
            True if running as admin/root, False otherwise
        """
        try:
            if self.platform_system == "Windows":
                import ctypes
                return ctypes.windll.shell32.IsUserAnAdmin() != 0
            else:
                return os.geteuid() == 0
        except:
            return False
    
    def _run_command(self, command: List[str], check: bool = True,
                    shell: bool = False, capture_output: bool = True,
                    cwd: Optional[str] = None) -> Tuple[int, str, str]:
        """
        Run a system command and return the result.
        
        Parameters
        ----------
        command : List[str]
            Command to run as a list of arguments
        check : bool, optional
            Whether to check for command success, by default True
        shell : bool, optional
            Whether to run command in shell, by default False
        capture_output : bool, optional
            Whether to capture stdout/stderr, by default True
        cwd : Optional[str], optional
            Working directory for the command, by default None
            
        Returns
        -------
        Tuple[int, str, str]
            Tuple containing (return_code, stdout, stderr)
        """
        if self.verbose:
            self.color.print(f"Running command: {' '.join(command)}", "cyan")
            
        try:
            if capture_output:
                process = subprocess.Popen(
                    command if not shell else " ".join(command),
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    shell=shell,
                    text=True,
                    cwd=cwd
                )
                
                stdout, stderr = process.communicate()
                return_code = process.returncode
            else:
                process = subprocess.Popen(
                    command if not shell else " ".join(command),
                    shell=shell,
                    cwd=cwd
                )
                
                process.communicate()
                return_code = process.returncode
                stdout, stderr = "", ""
                
            if check and return_code != 0 and capture_output:
                if self.verbose:
                    self.color.print(f"Command failed with code {return_code}: {stderr}", "red")
                    
            return return_code, stdout, stderr
            
        except Exception as e:
            if self.verbose:
                self.color.print(f"Error executing command: {e}", "red")
            return -1, "", str(e)
    
    def _download_file(self, url: str, destination: str, 
                      show_progress: bool = True) -> bool:
        """
        Download a file from a URL with progress tracking.
        
        Parameters
        ----------
        url : str
            URL to download from
        destination : str
            Local path to save the file
        show_progress : bool, optional
            Whether to show download progress, by default True
            
        Returns
        -------
        bool
            True if download successful, False otherwise
        """
        try:
            self.color.print(f"Downloading {url}", "blue")
            
            # Ensure destination directory exists
            os.makedirs(os.path.dirname(destination), exist_ok=True)
            
            with urllib.request.urlopen(url) as response, open(destination, 'wb') as out_file:
                file_size = int(response.info().get('Content-Length', 0))
                downloaded = 0
                block_size = 8192 * 8  # 64KB blocks
                
                if show_progress and file_size > 0:
                    self.color.print(f"Total file size: {file_size / 1024 / 1024:.1f} MB")
                
                while True:
                    buffer = response.read(block_size)
                    if not buffer:
                        break
                        
                    downloaded += len(buffer)
                    out_file.write(buffer)
                    
                    if show_progress and file_size > 0:
                        self.progress.progress_bar(
                            downloaded, file_size, 
                            prefix=f"  Progress:",
                            suffix=f"{downloaded/1024/1024:.1f}/{file_size/1024/1024:.1f} MB"
                        )
            
            self.color.print(f"Download completed: {destination}", "green")
            return True
            
        except Exception as e:
            self.color.print(f"Error downloading file: {e}", "red")
            return False
    
    def _verify_python_version(self) -> bool:
        """
        Verify that the Python version is 3.10.x as required.
        
        Returns
        -------
        bool
            True if Python version is 3.10.x, False otherwise
        """
        version = platform.python_version_tuple()
        
        if int(version[0]) != 3 or int(version[1]) != 10:
            self.color.print(
                f"ERROR: Incompatible Python version: {platform.python_version()}", "red", bold=True
            )
            self.color.print(
                "Maggie requires Python 3.10.x specifically. Other versions are not supported.", "red"
            )
            
            # Provide platform-specific instructions
            if self.platform_system == "Windows":
                self.color.print(
                    "Please install Python 3.10 from: https://www.python.org/downloads/release/python-31011/", 
                    "yellow"
                )
            else:
                self.color.print("Please install Python 3.10 using:", "yellow")
                self.color.print(
                    "sudo apt install python3.10 python3.10-venv python3.10-dev", 
                    "yellow"
                )
                
            return False
            
        self.color.print(f"Python {platform.python_version()} - Compatible ✓", "green")
        return True
    
    def _detect_hardware(self) -> Dict[str, Any]:
        """
        Detect system hardware configuration.
        
        Gathers detailed information about CPU, GPU, and memory
        with specific detection for AMD Ryzen 9 5900X and NVIDIA RTX 3080.
        
        Returns
        -------
        Dict[str, Any]
            Dictionary containing hardware information
        """
        hardware_info = {
            "cpu": self._detect_cpu(),
            "memory": self._detect_memory(),
            "gpu": self._detect_gpu() if not self.cpu_only else {"available": False}
        }
        
        # Log detected hardware
        self.color.print("Hardware Detection:", "cyan", bold=True)
        
        # Log CPU info
        cpu_info = hardware_info["cpu"] 
        if cpu_info["is_ryzen_9_5900x"]:
            self.color.print("  CPU: AMD Ryzen 9 5900X detected ✓", "green")
        else:
            self.color.print(f"  CPU: {cpu_info['model']}", "yellow")
            self.color.print(f"       {cpu_info['cores']} cores / {cpu_info['threads']} threads", "yellow")
            
        # Log Memory info
        memory_info = hardware_info["memory"]
        if memory_info["is_32gb"]:
            self.color.print(f"  RAM: {memory_info['total_gb']:.1f} GB (32GB detected) ✓", "green")
        else:
            self.color.print(f"  RAM: {memory_info['total_gb']:.1f} GB", "yellow")
            
        # Log GPU info
        gpu_info = hardware_info["gpu"]
        if self.cpu_only:
            self.color.print("  GPU: CPU-only mode selected (skipping GPU detection)", "yellow")
        elif gpu_info["is_rtx_3080"]:
            self.color.print(f"  GPU: NVIDIA RTX 3080 detected ✓", "green")
            self.color.print(f"       {gpu_info['vram_gb']:.1f} GB VRAM", "green")
            if gpu_info["cuda_available"]:
                self.color.print(f"       CUDA {gpu_info['cuda_version']} available", "green")
        elif gpu_info["available"]:
            self.color.print(f"  GPU: {gpu_info['model']}", "yellow")
            self.color.print(f"       {gpu_info['vram_gb']:.1f} GB VRAM", "yellow")
            if gpu_info["cuda_available"]:
                self.color.print(f"       CUDA {gpu_info['cuda_version']} available", "yellow")
        else:
            self.color.print("  GPU: No compatible GPU detected", "red")
        
        return hardware_info
    
    def _detect_cpu(self) -> Dict[str, Any]:
        """
        Detect CPU information with Ryzen 9 5900X detection.
        
        Returns
        -------
        Dict[str, Any]
            CPU information including model, cores, and specific CPU detection
        """
        cpu_info = {
            "is_ryzen_9_5900x": False,
            "model": platform.processor() or "Unknown",
            "cores": 0,
            "threads": 0
        }
        
        try:
            import psutil
            cpu_info["cores"] = psutil.cpu_count(logical=False) or 0
            cpu_info["threads"] = psutil.cpu_count(logical=True) or 0
        except ImportError:
            # Fall back to os.cpu_count()
            cpu_info["threads"] = os.cpu_count() or 0
            cpu_info["cores"] = cpu_info["threads"] // 2  # Estimate physical cores
        
        # Check for Ryzen 9 5900X
        model_lower = cpu_info["model"].lower()
        if "ryzen 9" in model_lower and "5900x" in model_lower:
            cpu_info["is_ryzen_9_5900x"] = True
        
        # Additional detection on Windows
        if self.platform_system == "Windows":
            try:
                # Try getting more detailed CPU info via WMI
                import wmi
                c = wmi.WMI()
                for cpu in c.Win32_Processor():
                    cpu_info["model"] = cpu.Name
                    if "Ryzen 9 5900X" in cpu.Name:
                        cpu_info["is_ryzen_9_5900x"] = True
                    break
            except:
                pass
        
        return cpu_info
    
    def _detect_memory(self) -> Dict[str, Any]:
        """
        Detect system memory information.
        
        Returns
        -------
        Dict[str, Any]
            Memory information including total, available, and type
        """
        memory_info = {
            "total_gb": 0,
            "available_gb": 0,
            "is_32gb": False,
            "type": "Unknown"
        }
        
        try:
            import psutil
            mem = psutil.virtual_memory()
            memory_info["total_gb"] = mem.total / (1024**3)
            memory_info["available_gb"] = mem.available / (1024**3)
            
            # Check if close to 32GB
            memory_info["is_32gb"] = 30 <= memory_info["total_gb"] <= 34
            
            # Try to detect memory type on Windows
            if self.platform_system == "Windows":
                try:
                    import wmi
                    c = wmi.WMI()
                    for mem_module in c.Win32_PhysicalMemory():
                        if hasattr(mem_module, 'PartNumber') and mem_module.PartNumber:
                            if "DDR4" in mem_module.PartNumber:
                                memory_info["type"] = "DDR4"
                                if "3200" in mem_module.PartNumber:
                                    memory_info["type"] = "DDR4-3200"
                                break
                except:
                    pass
        except ImportError:
            # Limited info without psutil
            memory_info["total_gb"] = 0
            memory_info["available_gb"] = 0
        
        return memory_info
    
    def _detect_gpu(self) -> Dict[str, Any]:
        """
        Detect GPU information with RTX 3080 detection.
        
        Returns
        -------
        Dict[str, Any]
            GPU information including model, VRAM, and CUDA version
        """
        gpu_info = {
            "available": False,
            "is_rtx_3080": False,
            "model": "Unknown",
            "vram_gb": 0,
            "cuda_available": False,
            "cuda_version": ""
        }
        
        if self.cpu_only:
            return gpu_info
        
        try:
            # Try importing torch to check for CUDA
            returncode, stdout, _ = self._run_command([
                sys.executable, "-c",
                "import torch; print(f'CUDA: {torch.cuda.is_available()}'); "
                "print(f'Device: {torch.cuda.get_device_name(0)}') if torch.cuda.is_available() else None; "
                "print(f'CUDA Version: {torch.version.cuda}') if torch.cuda.is_available() else None"
            ], check=False)
            
            if returncode != 0:
                # torch not installed yet
                if self.verbose:
                    self.color.print("PyTorch not installed yet, CUDA status unknown", "yellow")
                return gpu_info
            
            # Parse stdout from PyTorch
            for line in stdout.splitlines():
                if "CUDA: True" in line:
                    gpu_info["available"] = True
                    gpu_info["cuda_available"] = True
                elif "Device:" in line:
                    gpu_info["model"] = line.split("Device:")[1].strip()
                    if "3080" in gpu_info["model"]:
                        gpu_info["is_rtx_3080"] = True
                elif "CUDA Version:" in line:
                    gpu_info["cuda_version"] = line.split("CUDA Version:")[1].strip()
            
            # Get VRAM if available
            if gpu_info["cuda_available"]:
                returncode, stdout, _ = self._run_command([
                    sys.executable, "-c",
                    "import torch; print(f'VRAM: {torch.cuda.get_device_properties(0).total_memory / (1024**3):.2f}')"
                ], check=False)
                
                if returncode == 0 and "VRAM:" in stdout:
                    vram_str = stdout.split("VRAM:")[1].strip()
                    try:
                        gpu_info["vram_gb"] = float(vram_str)
                    except:
                        pass
        except:
            pass
            
        if self.verbose and not gpu_info["available"] and not self.cpu_only:
            self.color.print("No CUDA-capable GPU detected or PyTorch not installed", "yellow")
            self.color.print("Consider using --cpu-only if no GPU is available", "yellow")
            
        return gpu_info
    
    def _check_tools(self) -> Dict[str, bool]:
        """
        Check for required external tools and utilities.

        Verifies the availability of external dependencies needed for installation:
        1. Git - Required for downloading models and certain packages from repositories
        2. C++ compiler - Required for building packages from source:
        - Windows: Visual C++ compiler (cl.exe)
        - Linux: GCC compiler

        These tools are important for full functionality but installation can proceed
        with reduced capabilities if they are not available.

        Returns
        -------
        Dict[str, bool]
            Dictionary of tool availability status with the following keys:
            - 'git': Whether Git is installed and available
            - 'cpp_compiler': Whether an appropriate C++ compiler is available
        """
        tools_status = {
            "git": False,
            "cpp_compiler": False
        }
        
        # Check for Git
        returncode, stdout, _ = self._run_command(["git", "--version"], check=False)
        if returncode == 0:
            tools_status["git"] = True
            self.has_git = True
            self.color.print(f"Git found: {stdout.strip()}", "green")
        else:
            self.color.print("Git not found - limited functionality will be available", "yellow")
            self.color.print("Install Git for full functionality:", "yellow")
            if self.platform_system == "Windows":
                self.color.print("https://git-scm.com/download/win", "yellow")
            else:
                self.color.print("sudo apt-get install git", "yellow")
        
        # Check for C++ compiler
        if self.platform_system == "Windows":
            # On Windows, check for Visual C++ compiler
            returncode, _, _ = self._run_command(["where", "cl.exe"], check=False)
            if returncode == 0:
                tools_status["cpp_compiler"] = True
                self.has_cpp_compiler = True
                self.color.print("Visual C++ compiler (cl.exe) found", "green")
            else:
                self.color.print("Visual C++ compiler not found", "yellow")
                self.color.print("Some packages may need to be installed from wheels", "yellow")
                if self.verbose:
                    self.color.print("Install Visual C++ Build Tools:", "yellow")
                    self.color.print("https://visualstudio.microsoft.com/visual-cpp-build-tools/", "yellow")
        else:
            # On Linux, check for GCC
            returncode, _, _ = self._run_command(["which", "gcc"], check=False)
            if returncode == 0:
                tools_status["cpp_compiler"] = True
                self.has_cpp_compiler = True
                self.color.print("GCC compiler found", "green")
            else:
                self.color.print("GCC compiler not found", "yellow")
                self.color.print("Some packages may fail to build", "yellow")
                if self.verbose:
                    self.color.print("Install build tools:", "yellow")
                    self.color.print("sudo apt-get install build-essential", "yellow")
        
        return tools_status
    
    def _create_directories(self) -> bool:
        """
        Create all required directories for installation.

        Creates the following directory structure for Maggie AI Assistant:

        .
        ├── downloads/                # Downloaded files
        ├── logs/                     # Application logs
        └── maggie/                   # Main package directory
            ├── cache/                # Cache directory
            │   └── tts/              # Text-to-speech cache
            ├── core/                 # Core functionality
            ├── extensions/           # Extension modules
            ├── models/               # Downloaded model files
            │   ├── llm/              # Large Language Models
            │   ├── stt/              # Speech-to-text models
            │   └── tts/              # Text-to-speech voice models
            ├── templates/            # Template files
            │   └── extension/        # Extension templates       
            └── utils/                # Utility modules
                ├── config/           # Configuration utilities
                ├── hardware/         # Hardware detection utilities
                ├── llm/              # LLM processing utilities
                ├── stt/              # Speech-to-text utilities
                └── tts/              # Text-to-speech utilities

        Parameters
        ----------
        None

        Returns
        -------
        bool
            True if all directories created successfully, False otherwise

        Notes
        -----
        All paths are created relative to the base installation directory.
        Missing parent directories are created automatically. Package directories
        (maggie/, maggie/core/, maggie/extensions/, maggie/utils/, and its submodules)
        include an __init__.py file to make them importable Python modules.
        """
        # Create all required directories
        for directory in self.required_dirs:
            try:
                dir_path = os.path.join(self.base_dir, directory)
                if not os.path.exists(dir_path):
                    os.makedirs(dir_path, exist_ok=True)
                    if self.verbose:
                        self.color.print(f"Created directory: {directory}", "green")
            except Exception as e:
                self.color.print(f"Error creating directory {directory}: {e}", "red")
                return False
        
        # Create __init__.py files for package directories
        package_dirs = [
            "maggie",
            "maggie/core",
            "maggie/extensions",
            "maggie/utils",
            "maggie/utils/config",
            "maggie/utils/hardware",
            "maggie/utils/llm",
            "maggie/utils/stt",
            "maggie/utils/tts",
        ]
        for pkg_dir in package_dirs:
            init_path = os.path.join(self.base_dir, pkg_dir, "__init__.py")
            if not os.path.exists(init_path):
                try:
                    with open(init_path, "w") as f:
                        f.write("# Maggie AI Assistant package\n")
                    if self.verbose:
                        self.color.print(f"Created __init__.py in {pkg_dir}", "green")
                except Exception as e:
                    self.color.print(f"Error creating __init__.py in {pkg_dir}: {e}", "red")
                    return False
        
        self.color.print("All required directories created successfully", "green")
        return True
    
    def _setup_virtual_env(self) -> bool:
        """
        Set up Python virtual environment.
        
        Returns
        -------
        bool
            True if virtual environment created successfully, False otherwise
        """
        venv_dir = os.path.join(self.base_dir, "venv")
        
        # Check if venv already exists
        if os.path.exists(venv_dir):
            self.color.print("Virtual environment already exists", "yellow")
            return True
        
        # Create virtual environment
        python_cmd = sys.executable
        returncode, _, stderr = self._run_command([python_cmd, "-m", "venv", venv_dir])
        
        if returncode != 0:
            self.color.print(f"Error creating virtual environment: {stderr}", "red")
            return False
        
        self.color.print("Virtual environment created successfully", "green")
        return True
    
    def _get_venv_python(self) -> str:
        """
        Get path to Python executable in virtual environment.
        
        Returns
        -------
        str
            Path to Python executable in virtual environment
        """
        if self.platform_system == "Windows":
            return os.path.join(self.base_dir, "venv", "Scripts", "python.exe")
        else:
            return os.path.join(self.base_dir, "venv", "bin", "python")
    
    def _install_basic_dependencies(self, python_cmd: str) -> bool:
        """
        Install basic dependencies needed for further installation.
        
        Parameters
        ----------
        python_cmd : str
            Path to Python executable in virtual environment
            
        Returns
        -------
        bool
            True if installation successful, False otherwise
        """
        # Upgrade pip, setuptools, and wheel
        self.color.print("Upgrading pip, setuptools, and wheel...", "cyan")
        returncode, _, stderr = self._run_command(
            [python_cmd, "-m", "pip", "install", "--upgrade", "pip", "setuptools", "wheel"]
        )
        
        if returncode != 0:
            self.color.print(f"Error upgrading pip, setuptools, and wheel: {stderr}", "red")
            return False
        
        # Install other basic dependencies
        basic_deps = ["urllib3", "tqdm", "numpy", "psutil", "PyYAML", "loguru", "requests"]
        self.color.print(f"Installing basic packages: {', '.join(basic_deps)}...", "cyan")
        returncode, _, stderr = self._run_command(
            [python_cmd, "-m", "pip", "install", "--upgrade"] + basic_deps
        )
        
        if returncode != 0:
            self.color.print(f"Error installing basic packages: {stderr}", "red")
            return False
        
        return True
    
    def _install_pytorch(self, python_cmd: str) -> bool:
        """
        Install PyTorch with appropriate CUDA support.
        
        Parameters
        ----------
        python_cmd : str
            Path to Python executable in virtual environment
            
        Returns
        -------
        bool
            True if installation successful, False otherwise
        """
        if self.cpu_only:
            self.color.print("Installing PyTorch (CPU version)...", "cyan")
            cmd = [
                python_cmd, "-m", "pip", "install", 
                "torch==2.0.1", "torchvision==0.15.2", "torchaudio==2.0.2"
            ]
        else:
            self.color.print("Installing PyTorch with CUDA 11.8 support (optimized for RTX 3080)...", "cyan")
            cmd = [
                python_cmd, "-m", "pip", "install",
                "torch==2.0.1+cu118", "torchvision==0.15.2+cu118", "torchaudio==2.0.2+cu118",
                "--extra-index-url", "https://download.pytorch.org/whl/cu118"
            ]
        
        returncode, _, stderr = self._run_command(cmd)
        
        if returncode != 0:
            self.color.print(f"Error installing PyTorch: {stderr}", "red")
            self.color.print("Continuing with installation, but GPU acceleration may not work", "yellow")
            return False
        
        # Verify installation
        verify_cmd = [
            python_cmd, "-c",
            "import torch; print(f'PyTorch {torch.__version__} installed successfully'); "
            "print(f'CUDA available: {torch.cuda.is_available()}')"
        ]
        returncode, stdout, _ = self._run_command(verify_cmd, check=False)
        
        if returncode == 0:
            for line in stdout.splitlines():
                self.color.print(line, "green")
            return True
        else:
            self.color.print("PyTorch installation verification failed", "yellow")
            return False
    
    def _install_dependencies(self) -> bool:
        """
        Install all dependencies from requirements file.
        
        Parameters
        ----------
        None
        
        Returns
        -------
        bool
            True if installation successful, False otherwise
            
        Notes
        -----
        Performs installation in the following order:
        1. Basic dependencies (pip, setuptools, wheel, etc.)
        2. PyTorch with appropriate CUDA support or CPU-only version
        3. Standard dependencies from requirements.txt (filtering out special cases)
        4. Specialized dependencies that need custom handling
           - PyAudio (with platform-specific handling)
           - Kokoro TTS engine
           - Faster-Whisper speech recognition
           - GPU-specific dependencies
        """
        python_cmd = self._get_venv_python()
        
        # Install basic dependencies first
        if not self._install_basic_dependencies(python_cmd):
            self.color.print("Failed to install basic dependencies", "red")
            return False
        
        # Install PyTorch with appropriate CUDA support
        pytorch_success = self._install_pytorch(python_cmd)
        if not pytorch_success and not self.cpu_only:
            self.color.print("PyTorch with CUDA failed to install", "yellow")
            response = self.color.input(
                "Try installing CPU version instead? (y/n): ", color="magenta"
            )
            if response.lower() == 'y':
                # Retry with CPU version
                cmd = [
                    python_cmd, "-m", "pip", "install", 
                    "torch==2.0.1", "torchvision==0.15.2", "torchaudio==2.0.2"
                ]
                self._run_command(cmd)
        
        # Create a filtered requirements file excluding specialized dependencies
        req_path = os.path.join(self.base_dir, "requirements.txt")
        temp_req_path = os.path.join(self.base_dir, "temp_requirements.txt")
        
        try:
            with open(req_path, "r") as f:
                req_content = f.read()
            
            # Filter out packages that need special handling
            filtered_lines = []
            for line in req_content.splitlines():
                # Skip comments, empty lines, and special packages
                if (not line or line.startswith("#") or 
                    "torch" in line or "cuda" in line or 
                    "PyAudio" in line or "whisper" in line.lower() or
                    "kokoro" in line.lower()):
                    continue
                filtered_lines.append(line)
            
            with open(temp_req_path, "w") as f:
                f.write("\n".join(filtered_lines))
            
            # Install from the filtered requirements file
            self.color.print("Installing standard dependencies...", "cyan")
            returncode, _, stderr = self._run_command([
                python_cmd, "-m", "pip", "install", "-r", temp_req_path
            ])
            
            # Clean up temp file
            os.remove(temp_req_path)
            
            if returncode != 0:
                self.color.print(f"Error installing standard dependencies: {stderr}", "red")
                self.color.print("Continuing with installation of critical components", "yellow")
        
        except Exception as e:
            self.color.print(f"Error processing requirements file: {e}", "red")
            if os.path.exists(temp_req_path):
                os.remove(temp_req_path)
            return False
        
        # Install specialized dependencies
        self._install_specialized_dependencies(python_cmd)
        
        return True
    
    def _install_specialized_dependencies(self, python_cmd: str) -> bool:
        """
        Install specialized dependencies that need special handling.

        Parameters
        ----------
        python_cmd : str
            Path to Python executable in virtual environment
            
        Returns
        -------
        bool
            True if installation successful, False otherwise
            
        Notes
        -----
        Handles installation of packages that cannot be installed with the standard pip
        process due to platform-specific requirements, binaries, or alternative sources:

        1. PyAudio (0.2.13):
           - Windows: Requires prebuilt wheel due to compilation dependencies
           - Linux: Requires portaudio19-dev system package

        2. Kokoro TTS Engine:
           - Installed from GitHub repository, not PyPI
           - Requires different installation approaches with/without Git
           - Has GPU-specific optimizations when CUDA is available

        3. Faster-Whisper (Speech Recognition):
           - Requires specific version compatibility (0.9.0)
           - Has specialized optimizations for GPU acceleration 

        4. Whisper-streaming:
           - Requires copying custom module files to site-packages
           - Not available as a standard PyPI package

        5. GPU-specific dependencies:
           - onnxruntime-gpu only installed when GPU is available
           - Version must be compatible with CUDA runtime
        """
        # 1. Install PyAudio - platform-specific handling
        self._install_pyaudio(python_cmd)
        
        # 2. Install kokoro for TTS
        self._install_kokoro(python_cmd)
        
        # 3. Install faster-whisper for speech recognition
        self._install_whisper(python_cmd)
        
        # 4. Install whisper_streaming for real-time transcription
        self._install_whisper_streaming(python_cmd)
        
        # 5. Install GPU-specific dependencies if using GPU
        if not self.cpu_only:
            self.color.print("Installing GPU-specific dependencies...", "cyan")
            self._run_command([
                python_cmd, "-m", "pip", "install", "onnxruntime-gpu==1.15.1"
            ])
        
        return True
    
    def _install_pyaudio(self, python_cmd: str) -> bool:
        """
        Install PyAudio with platform-specific handling.
        
        Parameters
        ----------
        python_cmd : str
            Path to Python executable in virtual environment
            
        Returns
        -------
        bool
            True if installation successful, False otherwise
        """
        self.color.print("Installing PyAudio...", "cyan")
        
        # First check if already installed
        returncode, _, _ = self._run_command([
            python_cmd, "-c", "import pyaudio; print('PyAudio already installed')"
        ], check=False)
        
        if returncode == 0:
            self.color.print("PyAudio already installed", "green")
            return True
        
        # Platform-specific installation
        if self.platform_system == "Windows":
            # Try to install from wheel
            py_ver = f"{sys.version_info.major}{sys.version_info.minor}"
            wheel_url = f"https://files.pythonhosted.org/packages/27/bc/719d140ee63cf4b0725016531d36743a797ffdbab85e8536922902c9349a/PyAudio-0.2.14-cp310-cp310-win_amd64.whl"
            wheel_path = os.path.join(self.base_dir, "downloads", "wheels", "PyAudio-0.2.14-cp310-cp310-win_amd64.whl")
            
            # Download wheel
            os.makedirs(os.path.dirname(wheel_path), exist_ok=True)
            if not self._download_file(wheel_url, wheel_path):
                self.color.print("Failed to download PyAudio wheel", "red")
                return False
            
            # Install from wheel
            returncode, _, stderr = self._run_command([
                python_cmd, "-m", "pip", "install", wheel_path
            ])
            
            if returncode == 0:
                self.color.print("PyAudio installed successfully from wheel", "green")
                return True
            else:
                self.color.print(f"Error installing PyAudio from wheel: {stderr}", "red")
                
                # Try direct install as fallback
                returncode, _, stderr = self._run_command([
                    python_cmd, "-m", "pip", "install", "PyAudio==0.2.13"
                ])
                
                if returncode == 0:
                    self.color.print("PyAudio installed successfully", "green")
                    return True
                else:
                    self.color.print(f"Failed to install PyAudio: {stderr}", "red")
                    self.color.print("Audio input functionality will be limited", "yellow")
                    return False
        else:
            # On Linux, try to install from pip but may need system dependencies
            returncode, _, stderr = self._run_command([
                python_cmd, "-m", "pip", "install", "PyAudio==0.2.13"
            ])
            
            if returncode == 0:
                self.color.print("PyAudio installed successfully", "green")
                return True
            else:
                self.color.print(f"Error installing PyAudio: {stderr}", "red")
                self.color.print("You may need to install portaudio19-dev:", "yellow")
                self.color.print("sudo apt-get install portaudio19-dev", "yellow")
                self.color.print("Then try: pip install PyAudio==0.2.13", "yellow")
                return False
    
    def _install_kokoro(self, python_cmd: str) -> bool:
        """
        Install kokoro TTS engine.
        
        Parameters
        ----------
        python_cmd : str
            Path to Python executable in virtual environment
            
        Returns
        -------
        bool
            True if installation successful, False otherwise
        """
        self.color.print("Installing kokoro TTS engine...", "cyan")
        
        # First check if already installed
        returncode, _, _ = self._run_command([
            python_cmd, "-c", "import kokoro; print('kokoro already installed')"
        ], check=False)
        
        if returncode == 0:
            self.color.print("kokoro already installed", "green")
            return True
        
        # Make sure dependencies are installed first
        self._run_command([
            python_cmd, "-m", "pip", "install", "numpy", "tqdm", "soundfile"
        ])
        
        # Install onnxruntime (GPU or CPU version)
        if not self.cpu_only and self.hardware_info["gpu"]["cuda_available"]:
            self._run_command([
                python_cmd, "-m", "pip", "install", "onnxruntime-gpu==1.15.1"
            ])
        else:
            self._run_command([
                python_cmd, "-m", "pip", "install", "onnxruntime==1.15.1"
            ])
        
        # Install kokoro using git if available
        if self.has_git:
            returncode, _, stderr = self._run_command([
                python_cmd, "-m", "pip", "install", "git+https://github.com/hexgrad/kokoro.git"
            ])
            
            if returncode == 0:
                self.color.print("kokoro installed successfully from GitHub", "green")
                return True
            else:
                self.color.print(f"Error installing kokoro from GitHub: {stderr}", "red")
                self.color.print("Trying alternative installation method...", "yellow")
        
        # Download repo and install locally if git failed or not available
        kokoro_dir = os.path.join(self.base_dir, "downloads", "kokoro")
        os.makedirs(kokoro_dir, exist_ok=True)
        
        # Download the zip file
        zip_url = "https://github.com/hexgrad/kokoro/archive/refs/heads/main.zip"
        zip_path = os.path.join(self.base_dir, "downloads", "kokoro.zip")
        
        if not self._download_file(zip_url, zip_path):
            self.color.print("Failed to download kokoro repository", "red")
            return False
        
        # Extract the zip file
        try:
            with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                zip_ref.extractall(os.path.dirname(kokoro_dir))
            
            # Move to the right location
            extracted_dir = os.path.join(os.path.dirname(kokoro_dir), "kokoro-main")
            if os.path.exists(kokoro_dir):
                shutil.rmtree(kokoro_dir)
            shutil.move(extracted_dir, kokoro_dir)
            
            # Install from the local directory
            returncode, _, stderr = self._run_command([
                python_cmd, "-m", "pip", "install", kokoro_dir
            ])
            
            if returncode == 0:
                self.color.print("kokoro installed successfully from downloaded repository", "green")
                return True
            else:
                self.color.print(f"Error installing kokoro from local directory: {stderr}", "red")
                if self.skip_problematic:
                    self.color.print("Skipping kokoro installation (TTS may not work)", "yellow")
                    return False
                else:
                    # Final method - offer to skip
                    response = self.color.input(
                        "Failed to install kokoro. Skip this dependency? (y/n): ", 
                        color="magenta"
                    )
                    if response.lower() == "y":
                        self.color.print("Skipping kokoro installation", "yellow")
                        return False
                    else:
                        self.color.print("Installation cannot continue without kokoro", "red")
                        return False
        
        except Exception as e:
            self.color.print(f"Error extracting or installing kokoro: {e}", "red")
            return False
    
    def _install_whisper(self, python_cmd: str) -> bool:
        """
        Install faster-whisper for speech recognition.
        
        Parameters
        ----------
        python_cmd : str
            Path to Python executable in virtual environment
            
        Returns
        -------
        bool
            True if installation successful, False otherwise
        """
        self.color.print("Installing faster-whisper for speech recognition...", "cyan")
        
        # First check if already installed
        returncode, _, _ = self._run_command([
            python_cmd, "-c", "import faster_whisper; print('faster-whisper already installed')"
        ], check=False)
        
        if returncode == 0:
            self.color.print("faster-whisper already installed", "green")
            return True
        
        # Install faster-whisper and requirements
        returncode, _, stderr = self._run_command([
            python_cmd, "-m", "pip", "install", "faster-whisper==0.9.0", "soundfile==0.12.1"
        ])
        
        if returncode != 0:
            self.color.print(f"Error installing faster-whisper: {stderr}", "red")
            self.color.print("Speech recognition functionality may be limited", "yellow")
            return False
        
        self.color.print("faster-whisper installed successfully", "green")
        return True
    
    def _install_whisper_streaming(self, python_cmd: str) -> bool:
        """
        Install whisper-streaming package from GitHub.
        
        Parameters
        ----------
        python_cmd : str
            Path to Python executable in virtual environment
            
        Returns
        -------
        bool
            True if installation successful, False otherwise
            
        Notes
        -----
        Installs the whisper_streaming package using either direct GitHub 
        installation or by downloading and installing from a local source.
        Handles dependencies and ensures proper installation in the virtual
        environment.
        """
        self.color.print("Installing whisper_streaming module...", "cyan")
        
        try:
            # Install dependencies first
            dependencies = ["pyaudio", "numpy", "websockets", "torch", "transformers"]
            self.color.print("Installing whisper_streaming dependencies...", "cyan")
            self._run_command([
                python_cmd, "-m", "pip", "install"] + dependencies
            )
            
            # Download and extract the package
            whisper_streaming_dir = os.path.join(self.base_dir, "downloads", "whisper_streaming")
            os.makedirs(whisper_streaming_dir, exist_ok=True)
            
            # Download the zip file
            zip_url = "https://github.com/ufal/whisper_streaming/archive/refs/heads/main.zip"
            zip_path = os.path.join(self.base_dir, "downloads", "whisper_streaming.zip")
            
            if not self._download_file(zip_url, zip_path):
                self.color.print("Failed to download whisper_streaming repository", "red")
                return False
            
            # Extract the zip file
            try:
                with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                    zip_ref.extractall(os.path.dirname(whisper_streaming_dir))
                
                # Move to the right location
                extracted_dir = os.path.join(os.path.dirname(whisper_streaming_dir), "whisper_streaming-main")
                if os.path.exists(whisper_streaming_dir):
                    shutil.rmtree(whisper_streaming_dir)
                    shutil.move(extracted_dir, whisper_streaming_dir)
                
                    # As a last resort, copy modules directly
                    self._copy_whisper_streaming_modules(python_cmd)
                return True
            except Exception as e:
                self.color.print(f"Error extracting or installing whisper_streaming: {e}", "red")
                return False
                
        except Exception as e:
            self.color.print(f"Error installing whisper_streaming: {e}", "red")
            return False
            
    def _copy_whisper_streaming_modules(self, python_cmd: str) -> bool:
        """
        Copy whisper_streaming modules directly to site-packages.
        
        Parameters
        ----------
        python_cmd : str
            Path to Python executable in virtual environment
            
        Returns
        -------
        bool
            True if copy successful, False otherwise
        """
        try:
            # Find the site-packages directory
            returncode, stdout, _ = self._run_command([
                python_cmd, "-c", 
                "import site; print(site.getsitepackages()[0])"
            ])

            if returncode == 0:
                site_packages = stdout.strip()
                
                # Create whisper_streaming directory in site-packages
                ws_dir = os.path.join(site_packages, "Lib","site-packages", "whisper_streaming")
                os.makedirs(ws_dir, exist_ok=True)
                
                # Copy source files
                source_dir = os.path.join(self.base_dir, "downloads", 
                                        "whisper_streaming")
                
                if os.path.exists(source_dir):
                    # Copy all .py files
                    for file in os.listdir(source_dir):
                        if file.endswith(".py"):
                            src_file = os.path.join(source_dir, file)
                            dst_file = os.path.join(ws_dir, file)
                            shutil.copy2(src_file, dst_file)
                            
                    # Create __init__.py if needed
                    init_py = os.path.join(ws_dir, "__init__.py")
                    if not os.path.exists(init_py):
                        with open(init_py, 'w') as f:
                            f.write("# whisper_streaming package\n")
                    
                    self.color.print("Copied whisper_streaming modules to site-packages", "green")
                    return True
                else:
                    self.color.print(f"Source directory not found: {source_dir}", "red")
                    return False
            else:
                self.color.print("Could not determine site-packages directory", "red")
                return False
        except Exception as e:
            self.color.print(f"Error copying whisper_streaming modules: {e}", "red")
            return False
    
    def _download_af_heart_model(self) -> bool:
        """
        Download the af_heart TTS voice model into maggie/models/tts/.

        Parameters
        ----------
        None

        Returns
        -------
        bool
            True if download successful, False otherwise

        Notes
        -----
        Downloads the af_heart.pt file from Hugging Face and places it in the
        maggie/models/tts/ directory as specified in the directory structure.
        This voice model is used by the TTS component for speech synthesis.
        """
        model_dir = os.path.join(self.base_dir, "maggie", "models", "tts")
        model_path = os.path.join(model_dir, "af_heart.pt")

        # Check if model file already exists
        if os.path.exists(model_path):
            self.color.print("af_heart voice model file already exists", "green")
            return True
        
        # URL for af_heart model file
        model_url = "https://huggingface.co/hexgrad/Kokoro-82M/resolve/main/voices/af_heart.pt"
        
        # Download model file (directory is created by _download_file)
        self.color.print("Downloading af_heart voice model...", "cyan")
        if not self._download_file(model_url, model_path):
            self.color.print("Failed to download af_heart voice model", "red")
            return False
                
        self.color.print("af_heart voice model downloaded successfully", "green")
        return True
    
    def _download_mistral_model(self) -> bool:
        """
        Download Mistral 7B LLM model into maggie/models/llm/.

        Parameters
        ----------
        None

        Returns
        -------
        bool
            True if download successful or user opts to continue, False otherwise

        Notes
        -----
        Clones the Mistral 7B Instruct v0.3 GPTQ model repository from Hugging Face into the
        maggie/models/llm/mistral-7b-instruct-v0.3-GPTQ-4bit/ directory using Git, if available.
        This model requires approximately 5GB of storage space and is used for
        natural language processing and generation tasks.
        """
        if self.skip_models:
            self.color.print("Skipping Mistral model download (--skip-models)", "yellow")
            return True
        
        mistral_dir = os.path.join(self.base_dir, "maggie", "models", "llm", "mistral-7b-instruct-v0.3-GPTQ-4bit")
        
        # Check if model directory already exists and is not empty
        if os.path.exists(mistral_dir) and os.listdir(mistral_dir):
            self.color.print("Mistral model directory already exists", "green")
            return True
        
        # Create model directory
        os.makedirs(mistral_dir, exist_ok=True)
        
        # Ask for confirmation before downloading large model
        if not self.skip_models:
            response = self.color.input(
                "Download Mistral 7B model? This requires ~5GB of storage (y/n): ",
                color="magenta"
            )
            if response.lower() != "y":
                self.color.print("Skipping Mistral model download", "yellow")
                return True
        
        # Download model using Git if available
        if self.has_git:
            self.color.print("Downloading Mistral 7B model using Git (this may take a while)...", "cyan")
            returncode, _, _ = self._run_command([
                "git", "clone", "https://huggingface.co/neuralmagic/Mistral-7B-Instruct-v0.3-GPTQ-4bit",
                mistral_dir
            ], capture_output=False)
            
            if returncode == 0:
                self.color.print("Mistral model downloaded successfully", "green")
                return True
            else:
                self.color.print("Error downloading Mistral model with Git", "red")
                self.color.print("LLM functionality will be limited", "yellow")
                
                # Offer to continue without model
                response = self.color.input(
                    "Continue installation without Mistral model? (y/n): ",
                    color="magenta"
                )
                return response.lower() == "y"
        else:
            self.color.print("Git not found, cannot download Mistral model", "red")
            self.color.print("Install Git and rerun installation, or download model manually:", "yellow")
            self.color.print("https://huggingface.co/neuralmagic/Mistral-7B-Instruct-v0.3-GPTQ-4bit", "yellow")
            
            # Offer to continue without model
            response = self.color.input(
                "Continue installation without Mistral model? (y/n): ",
                color="magenta"
            )
            return response.lower() == "y"
    
    def _create_recipe_template(self) -> bool:
        """
        Create recipe template file for the recipe_creator extension in maggie/templates/.

        Generates a recipe_template.docx file using python-docx and places it in the
        maggie/templates/ directory as specified in the directory structure.

        Returns
        -------
        bool
            True if template created successfully, False otherwise
        """
        template_dir = os.path.join(self.base_dir, "maggie", "templates")
        template_path = os.path.join(template_dir, "recipe_template.docx")
        
        # Check if template already exists
        if os.path.exists(template_path):
            self.color.print("Recipe template already exists", "green")
            return True
        
        # Try to create template using docx
        try:
            python_cmd = self._get_venv_python()
            
            # Try to use python-docx
            returncode, _, _ = self._run_command([
                python_cmd, "-c",
                """
import docx

# Create document
doc = docx.Document()

# Add metadata section
doc.add_heading("Recipe Name", level=1)

# Add metadata section
doc.add_heading("Recipe Information", level=2)
info_table = doc.add_table(rows=3, cols=2)
info_table.style = 'Table Grid'
info_table.cell(0, 0).text = "Preparation Time"
info_table.cell(0, 1).text = "00 minutes"
info_table.cell(1, 0).text = "Cooking Time"
info_table.cell(1, 1).text = "00 minutes"
info_table.cell(2, 0).text = "Servings"
info_table.cell(2, 1).text = "0 servings"

# Add ingredients section
doc.add_heading("Ingredients", level=2)
doc.add_paragraph("• Ingredient 1", style='ListBullet')
doc.add_paragraph("• Ingredient 2", style='ListBullet')
doc.add_paragraph("• Ingredient 3", style='ListBullet')

# Add steps section
doc.add_heading("Instructions", level=2)
doc.add_paragraph("1. Step 1", style='ListNumber')
doc.add_paragraph("2. Step 2", style='ListNumber')
doc.add_paragraph("3. Step 3", style='ListNumber')

# Add notes section
doc.add_heading("Notes", level=2)
doc.add_paragraph("Add any additional notes, tips, or variations here.")

# Save template
doc.save("{}")
                """.format(template_path.replace("\\", "\\\\"))
            ])
            
            if returncode == 0:
                self.color.print("Recipe template created successfully", "green")
                return True
            else:
                # Alternative: Try using main.py with --create-template
                returncode, _, _ = self._run_command([
                    python_cmd, "main.py", "--create-template"
                ])
                
                if returncode == 0:
                    self.color.print("Recipe template created with main.py", "green")
                    return True
                else:
                    self.color.print("Failed to create recipe template", "red")
                    self.color.print("Recipe creator extension may not work properly", "yellow")
                    return False
                
        except Exception as e:
            self.color.print(f"Error creating recipe template: {e}", "red")
            return False
    
    def _setup_config(self) -> bool:
        """
        Set up configuration file using the virtual environment's Python.

        Parameters
        ----------
        None

        Returns
        -------
        bool
            True if configuration set up successfully, False otherwise

        Notes
        -----
        Creates or updates config.yaml using YAML operations within the virtual environment.
        The configuration is optimized for the detected hardware, particularly for
        AMD Ryzen 9 5900X CPU and NVIDIA RTX 3080 GPU if present.

        The configuration includes settings for:
        - System-wide parameters like inactivity_timeout
        - STT (Speech-to-Text) configuration for Whisper and DeepSpeech models
        - Wake word detection settings including access key and sensitivity
        - TTS (Text-to-Speech) parameters including voice model and caching
        - LLM (Large Language Model) settings for model path and GPU usage
        - Logging configuration
        - Extensions configuration
        - Hardware resource allocation (CPU, memory, GPU)
        """
        config_path = os.path.join(self.base_dir, "config.yaml")
        example_path = os.path.join(self.base_dir, "config.yaml.example")

        # Check for alternate example paths
        alt_example_path = os.path.join(self.base_dir, "config-yaml-example.txt")
        if not os.path.exists(example_path) and os.path.exists(alt_example_path):
            try:
                shutil.copy(alt_example_path, example_path)
                self.color.print(f"Created config example from {alt_example_path}", "green")
            except Exception as e:
                self.color.print(f"Error creating config example: {e}", "red")
                return False

        # Write hardware_info to temporary file
        import json
        temp_hardware_file = os.path.join(self.base_dir, "hardware_info.json")
        try:
            with open(temp_hardware_file, 'w') as f:
                json.dump(self.hardware_info, f)
        except Exception as e:
            self.color.print(f"Error writing hardware info: {e}", "red")
            return False

        # Subprocess code to run in virtual environment
        python_cmd = self._get_venv_python()
        
        # Fix: Use a properly normalized path without escape sequences
        normalized_base_dir = os.path.normpath(str(self.base_dir))
        # Fix: Use raw string for path in Python code
        base_dir_str = repr(normalized_base_dir)

        code = f"""import yaml
import json
import os

base_dir = {base_dir_str}
config_path = os.path.join(base_dir, 'config.yaml')
example_path = os.path.join(base_dir, 'config.yaml.example')
hardware_file = os.path.join(base_dir, 'hardware_info.json')

# Load hardware info
with open(hardware_file, 'r') as f:
    hardware_info = json.load(f)

# Check if config exists
if not os.path.exists(config_path):
    if not os.path.exists(example_path):
        print("Error: Configuration example file not found")
        exit(1)
    # Copy example to config
    with open(example_path, 'r') as f:
        config = yaml.safe_load(f)
else:
    with open(config_path, 'r') as f:
        config = yaml.safe_load(f)

# Set TTS voice model
if 'tts' in config and 'voice_model' in config['tts']:
    config['tts']['voice_model'] = 'af_heart.pt'

# Optimize for hardware
if hardware_info['gpu']['is_rtx_3080']:
    if 'llm' in config:
        config['llm']['gpu_layers'] = 32
        config['llm']['gpu_layer_auto_adjust'] = True
    if 'gpu' not in config:
        config['gpu'] = {{}}
    config['gpu']['max_percent'] = 90
    config['gpu']['model_unload_threshold'] = 95
    if 'stt' in config and 'whisper' in config['stt']:
        config['stt']['whisper']['compute_type'] = 'float16'
    if 'tts' in config:
        config['tts']['gpu_acceleration'] = True
        config['tts']['gpu_precision'] = 'mixed_float16'
elif {self.cpu_only}:
    if 'llm' in config:
        config['llm']['gpu_layers'] = 0
        config['llm']['gpu_layer_auto_adjust'] = False
    if 'gpu' not in config:
        config['gpu'] = {{}}
    config['gpu']['max_percent'] = 0
    config['gpu']['model_unload_threshold'] = 0
    if 'tts' in config:
        config['tts']['gpu_acceleration'] = False

if hardware_info['cpu']['is_ryzen_9_5900x']:
    if 'cpu' not in config:
        config['cpu'] = {{}}
    config['cpu']['max_threads'] = 8
    config['cpu']['thread_timeout'] = 30

if hardware_info['memory']['is_32gb']:
    if 'memory' not in config:
        config['memory'] = {{}}
    config['memory']['max_percent'] = 75
    config['memory']['model_unload_threshold'] = 85

# Write config
with open(config_path, 'w') as f:
    yaml.dump(config, f, default_flow_style=False)

# Clean up
os.remove(hardware_file)
"""
        # Run the subprocess
        returncode, _, stderr = self._run_command([python_cmd, "-c", code], cwd=self.base_dir)

        if returncode != 0:
            self.color.print(f"Error setting up configuration: {stderr}", "red")
            return False

        self.color.print("Configuration file created with optimized settings", "green")
        self.color.print("NOTE: You must edit config.yaml to add your Picovoice access key", "yellow")
        return True

    
    def _optimize_config_for_hardware(self, config: Dict[str, Any]) -> None:
        """
        Optimize configuration settings for detected hardware.
        
        Parameters
        ----------
        config : Dict[str, Any]
            Configuration dictionary to optimize

        Returns
        -------
        None
            Modifies the config dictionary in-place
            
        Notes
        -----
        Updates configuration dictionary with optimized settings based on
        detected hardware capabilities, especially for AMD Ryzen 9 5900X 
        and NVIDIA RTX 3080.
        
        For NVIDIA RTX 3080:
        - Sets LLM gpu_layers to 32 (optimal for 10GB VRAM)
        - Enables gpu_layer_auto_adjust for dynamic memory management
        - Sets stt/whisper compute_type to float16 for faster inference
        - Enables TTS GPU acceleration with mixed_float16 precision
        - Sets appropriate GPU memory thresholds
        
        For AMD Ryzen 9 5900X:
        - Configures max_threads to 8 (optimal for 12-core processor)
        - Sets appropriate thread timeout values
        
        For 32GB RAM systems:
        - Sets memory allocation thresholds appropriately
        """
        # GPU optimizations
        if not self.cpu_only and self.hardware_info["gpu"]["is_rtx_3080"]:
            # LLM optimizations for RTX 3080
            if "llm" in config:
                config["llm"]["gpu_layers"] = 32  # Optimal for 10GB VRAM
                config["llm"]["gpu_layer_auto_adjust"] = True
            
            # Add GPU section if it doesn't exist
            if "gpu" not in config:
                config["gpu"] = {}
            
            config["gpu"]["max_percent"] = 90  # Use up to 90% of GPU memory
            config["gpu"]["model_unload_threshold"] = 95  # Unload at 95%
            
            # STT Whisper compute type
            if "stt" in config and "whisper" in config["stt"]:
                config["stt"]["whisper"]["compute_type"] = "float16"  # Use float16 for faster inference
                
            # TTS GPU acceleration
            if "tts" in config:
                config["tts"]["gpu_acceleration"] = True
                config["tts"]["gpu_precision"] = "mixed_float16"
                
        elif self.cpu_only:
            # CPU-only settings
            if "llm" in config:
                config["llm"]["gpu_layers"] = 0
                config["llm"]["gpu_layer_auto_adjust"] = False
            
            if "gpu" not in config:
                config["gpu"] = {}
            
            config["gpu"]["max_percent"] = 0  # Don't use GPU
            config["gpu"]["model_unload_threshold"] = 0
            
            # Disable TTS GPU acceleration
            if "tts" in config:
                config["tts"]["gpu_acceleration"] = False
        
        # CPU optimizations for Ryzen 9 5900X
        if self.hardware_info["cpu"]["is_ryzen_9_5900x"]:
            # Add CPU section if it doesn't exist
            if "cpu" not in config:
                config["cpu"] = {}
            
            config["cpu"]["max_threads"] = 8  # Using 8 of the 12 available cores
            config["cpu"]["thread_timeout"] = 30
        
        # Memory optimizations for 32GB RAM
        if self.hardware_info["memory"]["is_32gb"]:
            if "memory" not in config:
                config["memory"] = {}
            
            config["memory"]["max_percent"] = 75  # Use up to 75% of system memory (24GB)
            config["memory"]["model_unload_threshold"] = 85  # Unload at 85% (27GB)
    
    def _install_extensions(self) -> bool:
        """
        Install available Maggie extensions.
        
        Installs dependencies for the recipe_creator extension, ensuring the template
        in maggie/templates/ is usable.

        Returns
        -------
        bool
            True if extensions installed successfully, False otherwise
        """
        # The default installation includes the recipe_creator extension,
        # which is already part of the repo. Just ensure its dependencies are installed.
        python_cmd = self._get_venv_python()
        
        # Install recipe_creator dependencies
        self.color.print("Installing recipe_creator extension dependencies...", "cyan")
        returncode, _, stderr = self._run_command([
            python_cmd, "-m", "pip", "install", "python-docx>=0.8.11"
        ])
        
        if returncode != 0:
            self.color.print(f"Error installing recipe_creator dependencies: {stderr}", "red")
            self.color.print("Recipe creator extension may not work properly", "yellow")
            return False
        
        self.color.print("Extensions installed successfully", "green")
        return True
        
    def verify_system(self) -> bool:
        """
        Verify system compatibility for installation.
        
        Parameters
        ----------
        None
        
        Returns
        -------
        bool
            True if system meets all requirements, False otherwise
            
        Notes
        -----
        Performs the following compatibility checks:
        1. Python version verification (must be exactly 3.10.x)
        2. Hardware detection (CPU, GPU, memory)
        3. Required external tools (Git, C++ compiler)
        
        The verification is the first step in the installation process,
        and must pass before proceeding with directory creation and 
        dependency installation.
        """
        self.progress.start_step("Verifying system compatibility")
        
        # Check Python version
        python_compatible = self._verify_python_version()
        if not python_compatible:
            self.progress.complete_step(False, "Incompatible Python version")
            return False
        
        # Detect hardware
        self.hardware_info = self._detect_hardware()
        
        # Check for required tools
        self._check_tools()
        
        # System verification passed
        self.progress.complete_step(True)
        return True
    
    def install(self) -> bool:
        """
        Run the complete installation process.
        
        Parameters
        ----------
        None
        
        Returns
        -------
        bool
            True if installation successful, False otherwise
            
        Notes
        -----
        Executes the full installation process in the following steps:
        
        1. System verification - Check compatibility requirements
        2. Directory creation - Set up the full directory structure
        3. Virtual environment - Create Python virtual environment
        4. Dependencies - Install all required packages
        5. Configuration - Set up and optimize config.yaml
        6. Models - Download required model files:
           - TTS voice model (af_heart.pt)
           - LLM model (Mistral 7B)
        7. Extensions - Set up extensions and templates
        8. Finalization - Complete installation and display instructions
        
        Each step is tracked with progress indicators and appropriate
        error handling for graceful failure recovery.
        """
        self.color.print("\n=== Maggie AI Assistant Installation ===", "cyan", bold=True)
        self.color.print(f"Platform: {self.platform_system} ({platform.platform()})", "cyan")
        self.color.print(f"Python: {platform.python_version()}", "cyan")
        
        if self.cpu_only:
            self.color.print("Mode: CPU-only (no GPU acceleration)", "yellow")
        
        # Step 1: Verify system compatibility
        if not self.verify_system():
            return False
        
        # Step 2: Create directories
        self.progress.start_step("Creating directory structure")
        if not self._create_directories():
            self.progress.complete_step(False, "Failed to create directories")
            return False
        self.progress.complete_step(True)
        
        # Step 3: Set up virtual environment
        self.progress.start_step("Setting up virtual environment")
        if not self._setup_virtual_env():
            self.progress.complete_step(False, "Failed to create virtual environment")
            return False
        self.progress.complete_step(True)
        
        # Step 4: Install dependencies
        self.progress.start_step("Installing dependencies")
        if not self._install_dependencies():
            self.progress.complete_step(False, "Some dependencies failed to install")
            if not self.color.input("Continue with installation? (y/n): ", color="magenta").lower() == 'y':
                return False
        else:
            self.progress.complete_step(True)
        
        # Step 5: Set up configuration
        self.progress.start_step("Setting up configuration")
        if not self._setup_config():
            self.progress.complete_step(False, "Failed to set up configuration")
            return False
        self.progress.complete_step(True)
        
        # Step 6: Download models
        self.progress.start_step("Downloading models")
        # Download TTS voice model
        if not self._download_af_heart_model():
            self.color.print("Warning: Failed to download TTS voice model", "yellow")
            self.color.print("Text-to-speech functionality may be limited", "yellow")
        
        # Download LLM model
        if not self.skip_models:
            self._download_mistral_model()
        else:
            self.color.print("Skipping LLM model download (--skip-models)", "yellow")
        
        self.progress.complete_step(True)
        
        # Step 7: Create recipe template
        self.progress.start_step("Setting up extensions")
        if not self._create_recipe_template():
            self.color.print("Warning: Failed to create recipe template", "yellow")
        
        # Install extensions
        if not self._install_extensions():
            self.color.print("Warning: Some extensions may not work properly", "yellow")
        
        self.progress.complete_step(True)
        
        # Step 8: Finalize installation
        self.progress.start_step("Completing installation")
        
        # Display completion message
        self.progress.display_summary(True)
        
        # Print instructions
        self.color.print("\nTo start Maggie AI Assistant:", "cyan", bold=True)
        if self.platform_system == "Windows":
            self.color.print("   .\\venv\\Scripts\\activate", "green")
            self.color.print("   python main.py", "green")
        else:
            self.color.print("   source venv/bin/activate", "green")
            self.color.print("   python main.py", "green")
        
        # Print important notes
        self.color.print("\nImportant Notes:", "cyan", bold=True)
        self.color.print("1. Edit config.yaml to add your Picovoice access key", "yellow")
        self.color.print("   Visit https://console.picovoice.ai/ to obtain a key", "yellow")
        
        if not self.has_git:
            self.color.print("2. Git is not installed - some features may be limited", "yellow")
            self.color.print("   Install Git for full functionality", "yellow")
        
        # Ask if user wants to start Maggie
        self.progress.complete_step(True)
        
        response = self.color.input("\nWould you like to start Maggie now? (y/n): ", color="magenta")
        if response.lower() == 'y':
            self.color.print("\nStarting Maggie AI Assistant...", "cyan", bold=True)
            python_cmd = self._get_venv_python()
            self._run_command([python_cmd, "main.py"], capture_output=False)
        
        return True


def main() -> int:
    """
    Main entry point for the installer.
    
    Parameters
    ----------
    None
    
    Returns
    -------
    int
        Exit code (0 for success, 1 for failure)
        
    Notes
    -----
    Parses command-line arguments and runs the installation process.
    Supported arguments:
    - --verbose: Enable verbose output
    - --cpu-only: Install CPU-only version (no GPU acceleration)
    - --skip-models: Skip downloading large LLM models
    - --skip-problematic: Skip problematic dependencies
    - --force-reinstall: Force reinstallation of packages
    
    The installer creates all necessary directories, installs dependencies,
    downloads models, and configures the system optimally for the 
    detected hardware.
    """
    parser = argparse.ArgumentParser(
        description="Maggie AI Assistant Installer",
        formatter_class=argparse.ArgumentDefaultsHelpFormatter
    )
    
    parser.add_argument(
        "--verbose", action="store_true",
        help="Enable verbose output"
    )
    parser.add_argument(
        "--cpu-only", action="store_true",
        help="Install CPU-only version (no GPU acceleration)"
    )
    parser.add_argument(
        "--skip-models", action="store_true",
        help="Skip downloading large LLM models"
    )
    parser.add_argument(
        "--skip-problematic", action="store_true",
        help="Skip problematic dependencies that may cause installation issues"
    )
    parser.add_argument(
        "--force-reinstall", action="store_true",
        help="Force reinstallation of already installed packages"
    )
    
    args = parser.parse_args()
    
    installer = MaggieInstaller(
        verbose=args.verbose,
        cpu_only=args.cpu_only,
        skip_models=args.skip_models,
        skip_problematic=args.skip_problematic,
        force_reinstall=args.force_reinstall
    )
    
    try:
        success = installer.install()
        return 0 if success else 1
    except KeyboardInterrupt:
        print("\nInstallation cancelled by user")
        return 1
    except Exception as e:
        print(f"\nUnexpected error during installation: {e}")
        return 1


if __name__ == "__main__":
    sys.exit(main())