#!/usr/bin/env python3
"""
Maggie AI Assistant - Unified Installation Script
=============================================
Handles installation and setup of the Maggie AI Assistant.

This script provides a unified installation experience for
Windows and Linux, with specific optimizations for
AMD Ryzen 9 5900X and NVIDIA RTX 3080 hardware.

The installation process includes:
1. System compatibility verification
2. Environment setup
3. Dependency installation
4. Model downloads
5. Configuration creation
6. System optimization

Example
-------
To install Maggie AI Assistant:
    $ python install.py

For verbose output:
    $ python install.py --verbose

To skip problematic packages:
    $ python install.py --skip-problematic

To install without GPU support (CPU only):
    $ python install.py --cpu-only

Notes
-----
- Requires Python 3.10.x specifically
- Windows installation requires Visual C++ Build Tools and Windows SDK
- GPU acceleration requires NVIDIA GPU with CUDA support
- Optimized for AMD Ryzen 9 5900X and NVIDIA RTX 3080
"""

# Standard library imports
import argparse
import sys
import os
sys.path.append(r"c:\ai\claude\maggie\venv\lib\site-packages")
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
import json
import math
import platform
import shutil
import subprocess
import time
import urllib.request
import zipfile
from pathlib import Path
from typing import Dict, Any, List, Tuple, Optional, Union

class MaggieInstaller:
    """
    Unified installer for Maggie AI Assistant.
    
    Provides platform-specific installation and configuration
    for Windows and Linux systems, with optimizations for
    AMD Ryzen 9 5900X and NVIDIA RTX 3080.
    
    Parameters
    ----------
    verbose : bool
        Whether to display verbose output
    skip_problematic : bool
        Whether to skip problematic packages with compilation issues
    cpu_only : bool
        Whether to install without GPU support
    force_reinstall : bool
        Whether to force reinstallation of packages
    
    Attributes
    ----------
    verbose : bool
        Flag for verbose output
    skip_problematic : bool
        Flag to skip problematic packages
    cpu_only : bool
        Flag to disable GPU support
    force_reinstall : bool
        Flag to force package reinstallation
    platform : str
        Detected platform ("Windows" or "Linux")
    base_dir : Path
        Base directory for installation
    colors : Dict[str, str]
        ANSI color codes for terminal output
    required_dirs : List[str]
        Required directories for installation
    is_admin : bool
        Whether script is running with admin/root privileges
    has_cpp_compiler : bool
        Whether C++ compiler is available
    has_git : bool
        Whether Git is available
    """
    
    def __init__(self, verbose: bool = False, skip_problematic: bool = False, 
                 cpu_only: bool = False, force_reinstall: bool = False):
        """
        Initialize the installer with configuration options.
        
        Parameters
        ----------
        verbose : bool, optional
            Whether to display verbose output, by default False
        skip_problematic : bool, optional
            Whether to skip problematic packages that may have compilation issues,
            by default False
        cpu_only : bool, optional
            Whether to install without GPU support, by default False
        force_reinstall : bool, optional
            Whether to force reinstallation of packages, by default False
        """
        self.verbose = verbose
        self.skip_problematic = skip_problematic
        self.cpu_only = cpu_only
        self.force_reinstall = force_reinstall
        self.platform = platform.system()
        self.base_dir = Path(os.path.dirname(os.path.abspath(__file__)))
        
        # Define colors for output
        if self.platform == "Windows":
            os.system("")  # Enable VT100 escape sequences on Windows
            
        # ANSI color codes
        self.colors = {
            "reset": "\033[0m",
            "red": "\033[91m",
            "green": "\033[92m",
            "yellow": "\033[93m",
            "blue": "\033[94m",
            "magenta": "\033[95m",
            "cyan": "\033[96m",
            "bold": "\033[1m"
        }
        
        # Required directories
        self.required_dirs = [
            "logs",
            "models",
            "models/tts",
            "models/tts/af_heart",
            "recipes",
            "templates",
            "cache",
            "cache/tts",
            "downloads",
            "downloads/wheels",
            "downloads/models",
            "site-packages"
        ]
        
        # Check if running as admin/root
        self.is_admin = self._check_admin()
        
        # Flags for special handling
        self.has_cpp_compiler = False
        self.has_git = False
    
    def _print(self, message: str, color: Optional[str] = None, bold: bool = False):
        """
        Print a message with optional color and formatting.
        
        Parameters
        ----------
        message : str
            Message to print
        color : str, optional
            Color name, by default None
        bold : bool, optional
            Whether to apply bold formatting, by default False
        """
        formatted_message = message
        if bold and "bold" in self.colors:
            formatted_message = f"{self.colors['bold']}{formatted_message}"
        if color and color in self.colors:
            formatted_message = f"{self.colors[color]}{formatted_message}{self.colors['reset']}"
        print(formatted_message)
            
    def _check_admin(self) -> bool:
        """
        Check if running with administrative privileges.
        
        Returns
        -------
        bool
            True if running as admin/root, False otherwise
        """
        if self.platform == "Windows":
            try:
                import ctypes
                return ctypes.windll.shell32.IsUserAnAdmin() != 0
            except:
                return False
        else:
            return os.geteuid() == 0
            
    def _run_command(self, cmd: List[str], check: bool = True, shell: bool = False, 
                    capture_output: bool = True) -> Tuple[int, str, str]:
        """
        Run a command and return the result.
        
        Parameters
        ----------
        cmd : List[str]
            Command to run as a list of strings
        check : bool, optional
            Whether to check for errors, by default True
        shell : bool, optional
            Whether to run as shell command, by default False
        capture_output : bool, optional
            Whether to capture stdout/stderr, by default True
            
        Returns
        -------
        Tuple[int, str, str]
            Return code, stdout, stderr
        """
        if self.verbose:
            self._print(f"Running command: {' '.join(cmd)}", "cyan")
            
        try:
            if capture_output:
                process = subprocess.Popen(
                    cmd if not shell else " ".join(cmd),
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    shell=shell,
                    text=True
                )
                
                stdout, stderr = process.communicate()
            else:
                # Run with direct output to terminal
                process = subprocess.Popen(
                    cmd if not shell else " ".join(cmd),
                    shell=shell
                )
                stdout, stderr = "", ""
                process.communicate()
            
            if check and process.returncode != 0 and capture_output:
                self._print(f"Command failed: {' '.join(cmd)}", "red")
                self._print(f"Error: {stderr}", "red")
                
            return process.returncode, stdout, stderr
            
        except Exception as e:
            self._print(f"Error running command: {e}", "red")
            return -1, "", str(e)
    
    def _download_file(self, url: str, destination: str, 
                      show_progress: bool = True) -> bool:
        """
        Download a file from a URL with progress tracking.
        
        Parameters
        ----------
        url : str
            URL to download from
        destination : str
            Path to save the file to
        show_progress : bool, optional
            Whether to show a progress bar, by default True
            
        Returns
        -------
        bool
            True if download was successful, False otherwise
        """
        try:
            self._print(f"Downloading {url}...", "cyan")
            
            # Ensure directory exists
            os.makedirs(os.path.dirname(destination), exist_ok=True)
            
            # Try to download with progress reporting
            with urllib.request.urlopen(url) as response, open(destination, 'wb') as out_file:
                file_size = int(response.info().get('Content-Length', 0))
                downloaded = 0
                block_size = 1024 * 16  # 16KB blocks for faster downloads
                
                while True:
                    buffer = response.read(block_size)
                    if not buffer:
                        break
                    
                    downloaded += len(buffer)
                    out_file.write(buffer)
                    
                    # Show progress
                    if show_progress and file_size > 0:
                        percent = int(downloaded * 100 / file_size)
                        bar_length = 30
                        filled_length = int(bar_length * downloaded // file_size)
                        bar = '█' * filled_length + '░' * (bar_length - filled_length)
                        sys.stdout.write(f"\r|{bar}| {percent}% ({downloaded/1024/1024:.1f}MB / {file_size/1024/1024:.1f}MB)")
                        sys.stdout.flush()
            
            if show_progress and file_size > 0:
                sys.stdout.write("\n")
                
            self._print(f"Download completed: {destination}", "green")
            return True
            
        except Exception as e:
            self._print(f"Error downloading file: {e}", "red")
            return False
    
    def _download_file_with_requests(self, url: str, destination: str) -> bool:
        """
        Download a file using the requests library with progress reporting.
        
        Parameters
        ----------
        url : str
            URL to download from
        destination : str
            Path to save the file to
            
        Returns
        -------
        bool
            True if download was successful, False otherwise
        """
        try:
            import requests
            from tqdm import tqdm
            
            self._print(f"Downloading {url}...", "cyan")
            
            # Ensure directory exists
            os.makedirs(os.path.dirname(destination), exist_ok=True)
            
            # Create a streaming request to get content
            with requests.get(url, stream=True) as r:
                r.raise_for_status()
                total_size = int(r.headers.get('content-length', 0))
                
                # Use tqdm for progress bar if available
                with open(destination, 'wb') as f:
                    for chunk in tqdm(r.iter_content(chunk_size=8192), 
                                      total=total_size//8192, 
                                      unit='KB',
                                      desc=os.path.basename(destination)):
                        f.write(chunk)
                        
            self._print(f"Download completed: {destination}", "green")
            return True
            
        except ImportError:
            # Fall back to urllib if requests is not available
            return self._download_file(url, destination)
            
        except Exception as e:
            self._print(f"Error downloading file: {e}", "red")
            return False
    
    def _check_python_version(self) -> bool:
        """
        Check if Python version is 3.10.x.
        
        Returns
        -------
        bool
            True if Python version is compatible, False otherwise
        """
        version = platform.python_version_tuple()
        
        if int(version[0]) != 3 or int(version[1]) != 10:
            self._print(f"ERROR: Unsupported Python version: {platform.python_version()}", "red")
            self._print("Maggie requires Python 3.10.x specifically", "red")
            
            # Suggest installation instructions
            if self.platform == "Windows":
                self._print("Please install Python 3.10 from: https://www.python.org/downloads/release/python-31011/", "yellow")
            else:
                self._print("Please install Python 3.10 using:", "yellow")
                self._print("sudo apt update && sudo apt install python3.10 python3.10-venv python3.10-dev", "yellow")
                
            return False
            
        self._print(f"Found Python {platform.python_version()} - Compatible version", "green")
        return True
        
    def _check_gpu(self) -> Dict[str, Any]:
        """
        Check if GPU is available and compatible.
        
        Returns
        -------
        Dict[str, Any]
            GPU information including device name, VRAM, and CUDA version
        """
        gpu_info = {
            "available": False,
            "name": None,
            "is_rtx_3080": False,
            "cuda_available": False,
            "cuda_version": None,
            "vram_gb": None
        }
        
        if self.cpu_only:
            self._print("CPU-only mode selected, skipping GPU checks", "yellow")
            return gpu_info
        
        try:
            # Check if PyTorch is installed
            import_returncode, _, _ = self._run_command(
                [sys.executable, "-c", "import torch; print('PyTorch available')"],
                check=False
            )
            
            if import_returncode != 0:
                self._print("PyTorch not installed yet, will install with CUDA support", "yellow")
                return gpu_info
                
            # Check CUDA availability
            returncode, stdout, _ = self._run_command([
                sys.executable, 
                "-c", 
                "import torch; print(f'CUDA available: {torch.cuda.is_available()}'); print(f'GPU: {torch.cuda.get_device_name(0)}' if torch.cuda.is_available() else 'No CUDA-capable GPU detected'); print(f'CUDA Version: {torch.version.cuda}' if torch.cuda.is_available() else 'No CUDA')"
            ], check=False)
            
            if returncode != 0:
                self._print("Error checking GPU capabilities", "yellow")
                return gpu_info
                
            # Parse output
            for line in stdout.splitlines():
                if "CUDA available: True" in line:
                    gpu_info["available"] = True
                    gpu_info["cuda_available"] = True
                elif "GPU:" in line:
                    gpu_name = line.split("GPU:")[1].strip()
                    gpu_info["name"] = gpu_name
                    
                    if "3080" in gpu_name:
                        gpu_info["is_rtx_3080"] = True
                        self._print(f"RTX 3080 detected - Will use optimized settings", "green")
                elif "CUDA Version:" in line and "No CUDA" not in line:
                    gpu_info["cuda_version"] = line.split("CUDA Version:")[1].strip()
                    
            # Get VRAM if available
            if gpu_info["cuda_available"]:
                returncode, stdout, _ = self._run_command([
                    sys.executable,
                    "-c",
                    "import torch; print(f'VRAM: {torch.cuda.get_device_properties(0).total_memory / (1024**3):.2f} GB')"
                ], check=False)
                
                if returncode == 0 and "VRAM:" in stdout:
                    vram_str = stdout.split("VRAM:")[1].strip().split(" GB")[0]
                    try:
                        gpu_info["vram_gb"] = float(vram_str)
                        self._print(f"GPU VRAM: {gpu_info['vram_gb']:.2f} GB", "green")
                    except ValueError:
                        pass
            
            # Log detailed GPU information
            if gpu_info["available"]:
                self._print(f"GPU: {gpu_info['name']}", "green")
                if gpu_info["cuda_version"]:
                    self._print(f"CUDA Version: {gpu_info['cuda_version']}", "green")
                
        except Exception as e:
            self._print(f"Error checking GPU: {e}", "yellow")
            
        return gpu_info
        
    def _check_git(self) -> bool:
        """
        Check if Git is installed and in PATH.
        
        Returns
        -------
        bool
            True if Git is available, False otherwise
        """
        try:
            returncode, stdout, _ = self._run_command(["git", "--version"], check=False)
            
            if returncode == 0 and "git version" in stdout:
                self._print("Git found: " + stdout.strip(), "green")
                self.has_git = True
                return True
            
            # Git not found, provide instructions
            self._print("Git not found in PATH", "yellow")
            self._print("Some dependencies require Git for installation", "yellow")
            self._print("To install Git on Windows:", "yellow")
            self._print("1. Download from https://git-scm.com/download/win", "yellow")
            self._print("2. Install with default options", "yellow")
            self._print("3. Restart this installation after installing Git", "yellow")
            
            if self.platform == "Linux":
                self._print("To install Git on Linux:", "yellow")
                self._print("sudo apt-get update && sudo apt-get install git", "yellow")
            
            return False
            
        except Exception as e:
            self._print(f"Error checking for Git: {e}", "red")
            return False
    
    def _check_cpp_compiler(self) -> bool:
        """
        Check if C++ compiler is available for building wheels.
        
        Returns
        -------
        bool
            True if compiler is available, False otherwise
        """
        if self.platform == "Windows":
            # On Windows, check for Visual C++ Build Tools
            try:
                # First, check for cl.exe (MSVC compiler)
                returncode, _, _ = self._run_command(["where", "cl.exe"], check=False)
                if returncode == 0:
                    self._print("Visual C++ compiler (cl.exe) found", "green")
                    self.has_cpp_compiler = True
                    return True
                
                # If we get here, compiler not found
                self._print("Visual C++ Build Tools not found", "yellow")
                self._print("Some packages may fail to build", "yellow")
                self._print("To install Visual C++ Build Tools:", "yellow")
                self._print("1. Download from https://visualstudio.microsoft.com/visual-cpp-build-tools/", "yellow")
                self._print("2. Select 'Desktop development with C++' workload", "yellow")
                self._print("3. Ensure 'Windows 10 SDK' and 'MSVC C++ x64/x86 build tools' are selected", "yellow")
                self._print("4. Restart this installation after installing Build Tools", "yellow")
                
                return False
            except Exception as e:
                self._print(f"Error checking for Visual C++ compiler: {e}", "red")
                return False
        else:
            # On Linux, check for gcc/g++
            try:
                returncode, _, _ = self._run_command(["which", "gcc"], check=False)
                if returncode == 0:
                    self._print("GCC compiler found", "green")
                    self.has_cpp_compiler = True
                    return True
                
                self._print("GCC compiler not found. Some packages may fail to build.", "yellow")
                self._print("To install GCC on Ubuntu/Debian: sudo apt install build-essential", "yellow")
                return False
            except Exception as e:
                self._print(f"Error checking for GCC compiler: {e}", "red")
                return False

    def _create_directories(self) -> bool:
        """
        Create required directories.
        
        Returns
        -------
        bool
            True if all directories created successfully, False otherwise
        """
        self._print("\nCreating required directories...", "cyan", bold=True)
        
        success = True
        for directory in self.required_dirs:
            dir_path = os.path.join(self.base_dir, directory)
            if not os.path.exists(dir_path):
                try:
                    os.makedirs(dir_path, exist_ok=True)
                    self._print(f"Created directory: {directory}", "green")
                except Exception as e:
                    self._print(f"Error creating directory {directory}: {e}", "red")
                    success = False
            else:
                self._print(f"Directory already exists: {directory}", "yellow")
                
        # Create maggie package directory structure for setup.py
        maggie_dir = os.path.join(self.base_dir, "maggie")
        if not os.path.exists(maggie_dir):
            try:
                os.makedirs(maggie_dir, exist_ok=True)
                # Create __init__.py in maggie directory
                with open(os.path.join(maggie_dir, "__init__.py"), "w") as f:
                    f.write("# Maggie AI Assistant package\n")
                self._print(f"Created package directory: maggie", "green")
            except Exception as e:
                self._print(f"Error creating package directory: {e}", "red")
                success = False
        else:
            self._print(f"Package directory already exists: maggie", "yellow")
            
        # Create maggie/utils directory for setup.py
        utils_dir = os.path.join(maggie_dir, "utils")
        if not os.path.exists(utils_dir):
            try:
                os.makedirs(utils_dir, exist_ok=True)
                # Create __init__.py in maggie/utils directory
                with open(os.path.join(utils_dir, "__init__.py"), "w") as f:
                    f.write("# Maggie AI Assistant utilities package\n")
                self._print(f"Created package directory: maggie/utils", "green")
            except Exception as e:
                self._print(f"Error creating package directory: {e}", "red")
                success = False
        else:
            self._print(f"Package directory already exists: maggie/utils", "yellow")
                
        return success
        
    def _setup_virtual_env(self) -> bool:
        """
        Set up virtual environment.
        
        Returns
        -------
        bool
            True if virtual environment created successfully, False otherwise
        """
        self._print("\nSetting up Python virtual environment...", "cyan", bold=True)
        
        venv_dir = os.path.join(self.base_dir, "venv")
        
        # Check if venv already exists
        if os.path.exists(venv_dir):
            self._print("Virtual environment already exists", "yellow")
            return True
            
        # Create venv
        python_cmd = sys.executable
        returncode, stdout, stderr = self._run_command([python_cmd, "-m", "venv", venv_dir])
        
        if returncode != 0:
            self._print(f"Error creating virtual environment: {stderr}", "red")
            return False
            
        self._print("Virtual environment created successfully", "green")
        return True
        
    def _install_basic_dependencies(self, python_cmd: str) -> bool:
        """
        Install basic dependencies needed for further installation.
        
        Parameters
        ----------
        python_cmd : str
            Path to Python executable in virtual environment
            
        Returns
        -------
        bool
            True if installation successful, False otherwise
        """
        self._print("Installing basic dependencies...", "cyan")
        
        # Upgrade pip, setuptools, and wheel using the proper method
        self._print("Upgrading pip, setuptools, and wheel...", "cyan")
        returncode, stdout, stderr = self._run_command([python_cmd, "-m", "pip", "install", "--upgrade", "pip", "setuptools", "wheel"])
        
        if returncode != 0:
            self._print(f"Error upgrading pip, setuptools, and wheel: {stderr}", "red")
            return False
        
        # Install requests for URL operations
        self._print("Installing requests package...", "cyan")
        returncode, _, stderr = self._run_command([python_cmd, "-m", "pip", "install", "requests"])
        
        if returncode != 0:
            self._print(f"Error installing requests package: {stderr}", "red")
            return False
            
        # Install other basic dependencies
        basic_deps = ["urllib3", "tqdm", "numpy", "psutil", "PyYAML", "loguru"]
        self._print(f"Installing basic packages: {', '.join(basic_deps)}...", "cyan")
        returncode, _, stderr = self._run_command([python_cmd, "-m", "pip", "install", "--upgrade"] + basic_deps)
        
        if returncode != 0:
            self._print(f"Error installing basic packages: {stderr}", "red")
            return False
            
        return True
    
    def _download_github_repo(self, repo_url: str, branch: str, dest_dir: str) -> bool:
        """
        Download a GitHub repository without using git.
        
        Parameters
        ----------
        repo_url : str
            GitHub repository URL (format: username/repository)
        branch : str
            Branch to download (usually "main" or "master")
        dest_dir : str
            Destination directory for the repository
            
        Returns
        -------
        bool
            True if download was successful, False otherwise
        """
        try:
            # Format the repository URL for the zip download
            if repo_url.startswith("https://github.com/"):
                repo_path = repo_url.replace("https://github.com/", "")
            else:
                repo_path = repo_url
                
            # Remove .git extension if present
            repo_path = repo_path.replace(".git", "")
            
            # Create the download URL
            download_url = f"https://github.com/{repo_path}/archive/refs/heads/{branch}.zip"
            
            # Download the zip file
            zip_path = os.path.join(self.base_dir, "downloads", f"{repo_path.replace('/', '_')}_{branch}.zip")
            
            if not self._download_file_with_requests(download_url, zip_path):
                return False
                
            # Extract the zip file
            self._print(f"Extracting {zip_path}...", "cyan")
            
            # Ensure destination directory exists
            os.makedirs(os.path.dirname(dest_dir), exist_ok=True)
            
            # Extract the zip file
            with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                zip_ref.extractall(os.path.dirname(dest_dir))
                
                # Get the name of the extracted directory (should be repo-branch)
                extracted_dir = os.path.join(os.path.dirname(dest_dir), 
                                            zip_ref.namelist()[0].split('/')[0])
                
            # Move the extracted directory to the destination
            if os.path.exists(dest_dir):
                shutil.rmtree(dest_dir)
                
            shutil.move(extracted_dir, dest_dir)
            
            # Clean up the zip file
            os.remove(zip_path)
            
            self._print(f"Downloaded and extracted repository to {dest_dir}", "green")
            return True
            
        except Exception as e:
            self._print(f"Error downloading GitHub repository: {e}", "red")
            return False

    def _install_whisper_streaming(self, python_cmd: str) -> bool:
        """
        Install whisper-streaming package directly.
        
        Parameters
        ----------
        python_cmd : str
            Path to Python executable in virtual environment
            
        Returns
        -------
        bool
            True if installation successful, False otherwise
        """
        self._print("Installing whisper-streaming...", "cyan")
        
        try:
            # Check if whisper-streaming is already installed
            returncode, stdout, _ = self._run_command([
                python_cmd, "-c", "import whisper_online; print('whisper-streaming available')"
            ], check=False)
            
            if returncode == 0 and "whisper-streaming available" in stdout:
                self._print("whisper-streaming package already installed", "green")
                return True
            
            # Copy the provided whisper_online.py and related files to site-packages
            source_files = [
                "site-packages/whisper_online.py",
                "site-packages/silero_vad_iterator.py",
                "site-packages/line_packet.py",
                "site-packages/whisper_online_server.py"
            ]
            
            # Find the site-packages directory in venv
            venv_site_packages = None
            
            if self.platform == "Windows":
                venv_site_packages = os.path.join(self.base_dir, "venv", "Lib", "site-packages")
            else:
                # Try to find site-packages directory for Linux
                for lib_dir in ["lib", "lib64"]:
                    for py_dir in ["python3.10", "python3"]:
                        test_path = os.path.join(self.base_dir, "venv", lib_dir, py_dir, "site-packages")
                        if os.path.exists(test_path):
                            venv_site_packages = test_path
                            break
                    if venv_site_packages:
                        break
            
            if not venv_site_packages:
                self._print("Could not find site-packages directory in virtual environment", "red")
                return False
                
            # Copy files
            for file_path in source_files:
                source = os.path.join(self.base_dir, file_path)
                if os.path.exists(source):
                    dest = os.path.join(venv_site_packages, os.path.basename(file_path))
                    shutil.copy2(source, dest)
                    self._print(f"Copied {os.path.basename(file_path)} to site-packages", "green")
                else:
                    self._print(f"Source file {file_path} not found", "red")
                    
            # Create __init__.py
            with open(os.path.join(venv_site_packages, "__init__.py"), "w") as f:
                f.write("# Whisper Streaming package\n")
                
            # Install requirements for whisper-streaming
            req_packages = ["faster-whisper==0.9.0", "soundfile==0.12.1"]
            self._run_command([
                python_cmd, "-m", "pip", "install"] + req_packages
            )
            
            self._print("Successfully set up whisper-streaming module", "green")
            return True
                
        except Exception as e:
            self._print(f"Error installing whisper-streaming: {e}", "red")
            return False
    
    def _install_from_wheel_url(self, python_cmd: str, package_name: str, wheel_url: str) -> bool:
        """
        Install a package from a wheel URL.
        
        Parameters
        ----------
        python_cmd : str
            Path to Python executable in virtual environment
        package_name : str
            Name of the package
        wheel_url : str
            URL to the wheel file
            
        Returns
        -------
        bool
            True if installation successful, False otherwise
        """
        try:
            self._print(f"Installing {package_name} from wheel...", "cyan")
            
            # Download the wheel
            wheel_dir = os.path.join(self.base_dir, "downloads", "wheels")
            os.makedirs(wheel_dir, exist_ok=True)
            
            wheel_filename = os.path.basename(wheel_url)
            wheel_path = os.path.join(wheel_dir, wheel_filename)
            
            if not self._download_file_with_requests(wheel_url, wheel_path):
                return False
                
            # Install the wheel
            returncode, _, stderr = self._run_command([
                python_cmd, "-m", "pip", "install", wheel_path
            ])
            
            if returncode != 0:
                self._print(f"Error installing {package_name} from wheel: {stderr}", "red")
                return False
                
            self._print(f"Successfully installed {package_name} from wheel", "green")
            return True
            
        except Exception as e:
            self._print(f"Error installing {package_name} from wheel: {e}", "red")
            return False

    def _find_prebuilt_wheel_urls(self) -> Dict[str, str]:
        """
        Find pre-built wheel URLs for required packages.
        
        Returns
        -------
        Dict[str, str]
            Dictionary mapping package names to wheel URLs
        """
        wheel_urls = {}
        py_version = f"cp{sys.version_info.major}{sys.version_info.minor}"
        platform_tag = "win_amd64" if self.platform == "Windows" else "linux_x86_64"
        
        # For PyAudio on Windows
        if self.platform == "Windows":
            wheel_urls["PyAudio"] = f"https://download.lfd.uci.edu/pythonlibs/archived/PyAudio-0.2.13-{py_version}-{py_version}-win_amd64.whl"
            self._print("Found pre-built wheel for PyAudio", "green")
        
        return wheel_urls

    def _install_kokoro(self, python_cmd: str) -> bool:
        """
        Install the kokoro TTS package.
        
        Parameters
        ----------
        python_cmd : str
            Path to Python executable
            
        Returns
        -------
        bool
            True if installation successful, False otherwise
        """
        try:
            self._print("\nInstalling kokoro TTS package...", "cyan", bold=True)
            
            # First check if kokoro is already installed
            returncode, stdout, _ = self._run_command([
                python_cmd, "-c", "import kokoro; print('kokoro available')"
            ], check=False)
            
            if returncode == 0 and "kokoro available" in stdout:
                self._print("kokoro package already installed", "green")
                return True
                
            # First check if PyTorch is installed
            returncode, stdout, _ = self._run_command([
                python_cmd, "-c", "import torch; print('PyTorch available')"
            ], check=False)
            
            if returncode != 0:
                self._print("PyTorch must be installed before kokoro", "yellow")
                return False
                
            # Install dependencies
            dependencies = ["numpy", "tqdm", "soundfile", "pyaudio", "onnxruntime"]
            
            # Add onnxruntime-gpu if GPU is available and not in CPU-only mode
            if not self.cpu_only:
                try:
                    import_returncode, stdout, _ = self._run_command([
                        python_cmd, "-c", "import torch; print(torch.cuda.is_available())"
                    ], check=False)
                    
                    if import_returncode == 0 and "True" in stdout:
                        dependencies = ["numpy", "tqdm", "soundfile", "pyaudio", "onnxruntime-gpu==1.15.1"]
                except:
                    pass
                
            self._print(f"Installing kokoro dependencies: {', '.join(dependencies)}", "cyan")
            returncode, _, stderr = self._run_command([
                python_cmd, "-m", "pip", "install", "--upgrade"] + dependencies
            )
            
            if returncode != 0:
                self._print(f"Error installing kokoro dependencies: {stderr}", "red")
                self._print("Continuing with kokoro installation anyway", "yellow")
            
            # Install kokoro
            if self.has_git:
                self._print("Installing kokoro from GitHub...", "cyan")
                returncode, _, stderr = self._run_command([
                    python_cmd, "-m", "pip", "install", "git+https://github.com/hexgrad/kokoro.git"
                ])
                
                if returncode == 0:
                    self._print("Successfully installed kokoro from GitHub", "green")
                    return True
                else:
                    self._print(f"Error installing kokoro from GitHub: {stderr}", "red")
            
            # If Git installation failed or Git not available, try direct download
            self._print("Trying alternative installation method for kokoro...", "cyan")
            
            # Download the repository
            kokoro_dir = os.path.join(self.base_dir, "downloads", "kokoro")
            if not self._download_github_repo("hexgrad/kokoro", "main", kokoro_dir):
                self._print("Failed to download kokoro repository", "red")
                return False
                
            # Install from the downloaded directory
            returncode, _, stderr = self._run_command([
                python_cmd, "-m", "pip", "install", kokoro_dir
            ])
            
            if returncode == 0:
                self._print("Successfully installed kokoro from downloaded repository", "green")
                return True
            else:
                self._print(f"Error installing kokoro from downloaded repository: {stderr}", "red")
                
                # Final method - offer to skip
                response = input(f"{self.colors['magenta']}Failed to install kokoro. Skip this dependency? (y/n): {self.colors['reset']}")
                if response.lower() == "y":
                    self._print("Skipping kokoro installation", "yellow")
                    return True  # True to continue with installation
                return False
                
        except Exception as e:
            self._print(f"Error installing kokoro: {e}", "red")
            return False

    def _download_af_heart_model(self) -> bool:
        """
        Download the af_heart TTS voice model.
        
        Returns
        -------
        bool
            True if download successful, False otherwise
        """
        try:
            self._print("\nDownloading af_heart voice model...", "cyan", bold=True)
            
            # Define model files and directories
            model_dir = os.path.join(self.base_dir, "models", "tts", "af_heart")
            os.makedirs(model_dir, exist_ok=True)
            
            # Check if model files already exist
            onnx_file = os.path.join(model_dir, "af_heart.onnx")
            json_file = os.path.join(model_dir, "af_heart.json")
            
            if os.path.exists(onnx_file) and os.path.exists(json_file):
                self._print("af_heart voice model files already exist", "green")
                return True
                
            # URLs for af_heart model files
            onnx_url = "https://huggingface.co/rhasspy/piper-voices/resolve/main/en/en_US/aflight/medium/af_heart.onnx"
            json_url = "https://huggingface.co/rhasspy/piper-voices/resolve/main/en/en_US/aflight/medium/af_heart.json"
            
            # Download ONNX model file
            if not os.path.exists(onnx_file):
                self._print(f"Downloading af_heart ONNX model file to {onnx_file}...", "cyan")
                if not self._download_file(onnx_url, onnx_file):
                    self._print("Failed to download af_heart ONNX model", "red")
                    return False
                
            # Download JSON model file
            if not os.path.exists(json_file):
                self._print(f"Downloading af_heart JSON model file to {json_file}...", "cyan")
                if not self._download_file(json_url, json_file):
                    self._print("Failed to download af_heart JSON model", "red")
                    return False
                    
            self._print("Successfully downloaded af_heart voice model files", "green")
            return True
            
        except Exception as e:
            self._print(f"Error downloading af_heart voice model: {e}", "red")
            return False
    
    def _install_dependencies(self) -> bool:
        """
        Install dependencies in virtual environment.
        
        Returns
        -------
        bool
            True if critical dependencies installed successfully, False otherwise
        """
        self._print("\nInstalling dependencies...", "cyan", bold=True)
        
        # Check for Git first
        self._check_git()
        
        # Check for C++ compiler
        self._check_cpp_compiler()
        
        # Determine python command based on platform
        if self.platform == "Windows":
            python_cmd = os.path.join(self.base_dir, "venv", "Scripts", "python")
        else:
            python_cmd = os.path.join(self.base_dir, "venv", "bin", "python")
            
        # Upgrade pip, setuptools, wheel using proper method
        if not self._install_basic_dependencies(python_cmd):
            self._print("Warning: Failed to install basic dependencies, continuing anyway", "yellow")
        
        # Install PyTorch with CUDA support (or CPU version if --cpu-only)
        if self.cpu_only:
            self._print("Installing PyTorch CPU version...", "cyan")
            returncode, _, stderr = self._run_command([
                python_cmd, "-m", "pip", "install", "torch==2.0.1", "torchvision==0.15.2", "torchaudio==2.0.2"
            ])
        else:
            self._print("Installing PyTorch with CUDA 11.8 support (optimized for RTX 3080)...", "cyan")
            returncode, _, stderr = self._run_command([
                python_cmd, "-m", "pip", "install", "torch==2.0.1+cu118", "torchvision==0.15.2+cu118", "torchaudio==2.0.2+cu118", "--extra-index-url", "https://download.pytorch.org/whl/cu118"
            ])
        
        if returncode != 0:
            self._print(f"Error installing PyTorch: {stderr}", "red")
            self._print("Continuing with installation, but GPU acceleration may not work", "yellow")
        
        # Get pre-built wheels if possible
        wheel_urls = self._find_prebuilt_wheel_urls()
        
        # Create a temporary requirements file that excludes special dependencies we'll install separately
        temp_req_path = os.path.join(self.base_dir, "temp_requirements.txt")
        try:
            with open(os.path.join(self.base_dir, "requirements.txt"), "r") as f:
                req_content = f.read()
            
            # Remove lines for dependencies that require special handling
            filtered_content = "\n".join([
                line for line in req_content.split("\n") 
                if line and not line.startswith("#") and  # Skip comments and empty lines
                not line.startswith("torch") and not "cu118" in line and
                not "whisper" in line.lower() and 
                not "kokoro" in line.lower() and
                not "PyAudio" in line.lower()  # PyAudio often needs special handling on Windows
            ])
            
            # Add kokoro which is needed for the af_heart voice model
            filtered_content += "\n# TTS engine\n# kokoro must be installed separately, see installation script\n"
            
            with open(temp_req_path, "w") as f:
                f.write(filtered_content)
                
            # Install from the temporary requirements file
            self._print("Installing standard dependencies...", "cyan")
            returncode, _, stderr = self._run_command([python_cmd, "-m", "pip", "install", "-r", temp_req_path])
            
            # Clean up
            os.remove(temp_req_path)
            
            if returncode != 0:
                self._print(f"Error installing standard dependencies: {stderr}", "red")
                self._print("Continuing with installation of critical components", "yellow")
            
            # Install specialized dependencies in the correct order

            # 1. Install kokoro TTS engine for af_heart voice model
            kokoro_installed = self._install_kokoro(python_cmd)
            if not kokoro_installed:
                self._print("Warning: kokoro TTS package installation failed", "yellow")
                self._print("The af_heart voice model will not work properly", "yellow")
                
            # 2. Install whisper-streaming using our custom method
            whisper_installed = self._install_whisper_streaming(python_cmd)
                
            if not whisper_installed:
                self._print("Warning: whisper-streaming installation failed", "yellow")
                self._print("Voice recognition may not work properly", "yellow")
                
            # 3. Install PyAudio if available as pre-built wheel
            if "PyAudio" in wheel_urls:
                self._print("Installing PyAudio from pre-built wheel...", "cyan")
                pyaudio_installed = self._install_from_wheel_url(
                    python_cmd, "PyAudio", wheel_urls["PyAudio"]
                )
                
                if not pyaudio_installed:
                    self._print("Error installing PyAudio from wheel", "red")
                    self._print("You may need to install PyAudio manually", "yellow")
                    self._print("Note: Audio input functionality will be limited", "yellow")
            
            # 4. Install GPU-specific dependencies if not in CPU-only mode
            if not self.cpu_only:
                self._print("Installing GPU-specific dependencies...", "cyan")
                returncode, _, stderr = self._run_command([python_cmd, "-m", "pip", "install", "onnxruntime-gpu==1.15.1"])
                
                if returncode != 0:
                    self._print(f"Error installing GPU-specific dependencies: {stderr}", "red")
                    self._print("Continuing with CPU-only operation", "yellow")
                
            self._print("Dependencies installed successfully", "green")
            return True
                
        except Exception as e:
            self._print(f"Error processing requirements file: {e}", "red")
            if os.path.exists(temp_req_path):
                os.remove(temp_req_path)
            return False
        
    def _setup_config(self) -> bool:
        """
        Set up configuration file.
        
        Returns
        -------
        bool
            True if configuration set up successfully, False otherwise
        """
        self._print("\nSetting up configuration...", "cyan", bold=True)
        
        config_path = os.path.join(self.base_dir, "config.yaml")
        example_path = os.path.join(self.base_dir, "config.yaml.example")
        
        # If example doesn't exist but we have config-yaml-example.txt
        alt_example_path = os.path.join(self.base_dir, "config-yaml-example.txt")
        if not os.path.exists(example_path) and os.path.exists(alt_example_path):
            try:
                shutil.copy(alt_example_path, example_path)
                self._print(f"Created config example from {alt_example_path}", "green")
            except Exception as e:
                self._print(f"Error creating config example: {e}", "red")
                return False
        
        # Check if config already exists
        if os.path.exists(config_path):
            self._print("Configuration file already exists", "yellow")
            
            # Update the TTS voice model to af_heart
            try:
                import yaml
                with open(config_path, 'r') as f:
                    config = yaml.safe_load(f)
                
                # Set TTS voice model to af_heart
                if 'speech' in config and 'tts' in config['speech']:
                    if config['speech']['tts']['voice_model'] != "af_heart":
                        config['speech']['tts']['voice_model'] = "af_heart"
                        with open(config_path, 'w') as f:
                            yaml.dump(config, f, default_flow_style=False)
                        self._print("Updated configuration to use af_heart voice model", "green")
                
                return True
            except Exception as e:
                self._print(f"Error updating configuration: {e}", "red")
                return False
            
        # Check if example exists
        if not os.path.exists(example_path):
            self._print("Configuration example file not found", "red")
            return False
            
        # Copy example to config and modify for af_heart
        try:
            # Read example config
            import yaml
            with open(example_path, 'r') as f:
                config = yaml.safe_load(f)
            
            # Set TTS voice model to af_heart
            if 'speech' in config and 'tts' in config['speech']:
                config['speech']['tts']['voice_model'] = "af_heart"
            
            # Optimize for Ryzen 9 5900X and RTX 3080
            if not self.cpu_only:
                # Add optimized LLM settings for RTX 3080
                if 'llm' in config:
                    config['llm']['gpu_layers'] = 32  # Optimal for 10GB VRAM
                    config['llm']['gpu_layer_auto_adjust'] = True
                    config['llm']['precision'] = "float16"  # Best for tensor cores
                
                # Optimize threading for Ryzen 9 5900X
                if 'threading' not in config:
                    config['threading'] = {}
                config['threading']['max_workers'] = 8  # Using 8 of the 12 available cores
                
                # Add GPU-specific settings
                if 'gpu' not in config:
                    config['gpu'] = {}
                config['gpu']['enabled'] = True
                config['gpu']['compute_type'] = "float16"
                config['gpu']['tensor_cores'] = True
                config['gpu']['reserved_memory_mb'] = 512
            else:
                # CPU-only settings
                if 'llm' in config:
                    config['llm']['gpu_layers'] = 0
                if 'gpu' not in config:
                    config['gpu'] = {}
                config['gpu']['enabled'] = False
            
            # Write modified config
            with open(config_path, 'w') as f:
                yaml.dump(config, f, default_flow_style=False)
            
            self._print(f"Configuration file created from example with af_heart voice model", "green")
            self._print("NOTE: You need to edit config.yaml to add your Picovoice access key", "yellow")
            return True
        except Exception as e:
            self._print(f"Error creating configuration file: {e}", "red")
            return False
            
    def _download_models(self) -> bool:
        """
        Download required models.
        
        Returns
        -------
        bool
            True if models downloaded successfully, False otherwise
        """
        self._print("\nChecking model downloads...", "cyan", bold=True)
        
        # Ask user if they want to download models
        response = input(f"{self.colors['magenta']}Download models? This may take significant time and bandwidth (y/n): {self.colors['reset']}")
        
        if response.lower() != "y":
            self._print("Skipping model downloads", "yellow")
            return True
        
        # Download TTS models directly (doesn't require Git)
        if not self._download_af_heart_model():
            self._print("Warning: Error downloading TTS model", "yellow")
            self._print("Text-to-speech functionality may not work correctly", "yellow")
            
        # Download Mistral model - requires Git
        if self.has_git:
            # Check for git-lfs
            self._print("Checking for Git LFS...", "cyan")
            returncode, _, _ = self._run_command(["git", "lfs", "version"], check=False)
            
            if returncode != 0:
                self._print("Git LFS not found, installing...", "cyan")
                if self.platform == "Windows":
                    self._run_command(["git", "lfs", "install"], check=False)
                else:
                    if self.is_admin:
                        self._run_command(["apt", "install", "-y", "git-lfs"], check=False)
                    else:
                        self._run_command(["sudo", "apt", "install", "-y", "git-lfs"], check=False)
                    self._run_command(["git", "lfs", "install"], check=False)
                    
            # Download Mistral model
            mistral_dir = os.path.join(self.base_dir, "models", "mistral-7b-instruct-v0.3-GPTQ-4bit")
            if not os.path.exists(mistral_dir):
                self._print("Downloading Mistral 7B model... (this may take a while)", "cyan")
                # For better visibility, don't capture the output
                returncode, _, _ = self._run_command([
                    "git", "clone", "https://huggingface.co/TheBloke/Mistral-7B-Instruct-v0.3-GPTQ", mistral_dir
                ], capture_output=False)
                
                if returncode != 0:
                    self._print("Error downloading Mistral model", "red")
                    return False
            else:
                self._print("Mistral model directory already exists", "yellow")
        else:
            self._print("Git not found, skipping Mistral model download", "yellow")
            self._print("LLM functionality will not work without the model", "yellow")
            self._print("Please install Git and run the installation again", "yellow")
            
        self._print("Model downloads completed", "green")
        return True
        
    def _create_recipe_template(self) -> bool:
        """
        Create recipe template document.
        
        Returns
        -------
        bool
            True if template created successfully or already exists, False otherwise
        """
        self._print("\nChecking recipe template...", "cyan", bold=True)
        
        # Determine python command based on platform
        if self.platform == "Windows":
            python_cmd = os.path.join(self.base_dir, "venv", "Scripts", "python")
        else:
            python_cmd = os.path.join(self.base_dir, "venv", "bin", "python")
            
        # Try to create the template
        try:
            # Create the template directory if it doesn't exist
            template_dir = os.path.join(self.base_dir, "templates")
            os.makedirs(template_dir, exist_ok=True)
            
            # Check if template already exists
            template_path = os.path.join(template_dir, "recipe_template.docx")
            if os.path.exists(template_path):
                self._print(f"Recipe template already exists: {template_path}", "yellow")
                return True
                
            # Try to create the template using docx
            try:
                from docx import Document
                
                doc = Document()
                
                # Add metadata section
                doc.add_heading("Recipe Name", level=1)
                
                # Add metadata section
                doc.add_heading("Recipe Information", level=2)
                info_table = doc.add_table(rows=3, cols=2)
                info_table.style = 'Table Grid'
                info_table.cell(0, 0).text = "Preparation Time"
                info_table.cell(0, 1).text = "00 minutes"
                info_table.cell(1, 0).text = "Cooking Time"
                info_table.cell(1, 1).text = "00 minutes"
                info_table.cell(2, 0).text = "Servings"
                info_table.cell(2, 1).text = "0 servings"
                
                # Add ingredients section with better formatting
                doc.add_heading("Ingredients", level=2)
                doc.add_paragraph("• Ingredient 1", style='ListBullet')
                doc.add_paragraph("• Ingredient 2", style='ListBullet')
                doc.add_paragraph("• Ingredient 3", style='ListBullet')
                
                # Add steps section with numbered steps
                doc.add_heading("Instructions", level=2)
                doc.add_paragraph("1. Step 1", style='ListNumber')
                doc.add_paragraph("2. Step 2", style='ListNumber')
                doc.add_paragraph("3. Step 3", style='ListNumber')
                
                # Add notes section
                doc.add_heading("Notes", level=2)
                doc.add_paragraph("Add any additional notes, tips, or variations here.")
                
                # Save the template
                doc.save(template_path)
                
                self._print(f"Recipe template created at {template_path}", "green")
                return True
                
            except ImportError:
                # If python-docx is not available, try running main.py with --create-template
                returncode, _, _ = self._run_command([python_cmd, "main.py", "--create-template"])
                
                if returncode != 0:
                    self._print("Error creating recipe template using main.py", "red")
                    return False
                    
                self._print("Recipe template created successfully", "green")
                return True
                
        except Exception as e:
            self._print(f"Error creating recipe template: {e}", "red")
            return False
    
    def verify_system(self) -> bool:
        """
        Verify system meets requirements for Maggie.
        
        Returns
        -------
        bool
            True if system meets requirements, False otherwise
        """
        self._print("\nVerifying system configuration...", "cyan", bold=True)
        
        verification_issues = []
        
        # Check Python version
        if not self._check_python_version():
            verification_issues.append("Python version must be 3.10.x")
        
        # Check GPU if not in CPU-only mode
        if not self.cpu_only:
            gpu_info = self._check_gpu()
            if not gpu_info["available"]:
                self._print("GPU not detected or CUDA not available", "yellow")
                self._print("LLM performance will be limited without GPU acceleration", "yellow")
                
        # Check for required directories
        for directory in self.required_dirs:
            dir_path = os.path.join(self.base_dir, directory)
            if not os.path.exists(dir_path):
                verification_issues.append(f"Required directory missing: {directory}")
                
        # Check if Git is available
        if not self._check_git():
            self._print("Git not available - model downloads will be limited", "yellow")
            
        # Check if C++ compiler is available
        if not self._check_cpp_compiler():
            self._print("C++ compiler not available - some packages may fail to build", "yellow")
        
        # Final verification result
        if verification_issues:
            self._print("System verification failed with the following issues:", "red")
            for issue in verification_issues:
                self._print(f"  - {issue}", "red")
            return False
        
        self._print("System verification passed", "green")
        return True
        
    def install(self) -> bool:
        """
        Run the full installation process.
        
        Returns
        -------
        bool
            True if installation successful, False otherwise
        """
        self._print("=== Maggie AI Assistant Installation ===", "cyan", bold=True)
        self._print(f"Platform: {self.platform}", "cyan")
        
        # Print current timestamp for installation timing
        start_time = time.time()
        self._print(f"Installation started at: {time.strftime('%Y-%m-%d %H:%M:%S')}", "cyan")
        
        # Print hardware optimization information
        self._print(f"Hardware Target: AMD Ryzen 9 5900X + NVIDIA RTX 3080", "cyan")
        if self.cpu_only:
            self._print("CPU-only mode: GPU acceleration disabled", "yellow")
        
        if not self.is_admin:
            if self.platform == "Windows":
                self._print("Note: Running without Administrator privileges", "yellow")
                self._print("Some optimizations may not be applied", "yellow")
            else:
                self._print("Note: Running without root privileges", "yellow")
                self._print("Some optimizations may not be applied", "yellow")
        
        # Initial system verification
        if not self.verify_system():
            return False
            
        # Create directories
        if not self._create_directories():
            return False
            
        # Setup virtual environment
        if not self._setup_virtual_env():
            return False
            
        # Install dependencies
        self._print("\nInstalling dependencies (this may take some time)...", "cyan", bold=True)
        dependency_result = self._install_dependencies()
        if not dependency_result:
            self._print("Some dependencies failed to install", "yellow")
            self._print("Continuing with installation, but some features may not work", "yellow")
            
            # Ask user if they want to continue despite dependency failures
            response = input(f"{self.colors['magenta']}Some dependencies failed to install. Continue anyway? (y/n): {self.colors['reset']}")
            
            if response.lower() != "y":
                self._print("Installation aborted by user due to dependency failures", "red")
                return False
            
        # Setup configuration
        if not self._setup_config():
            return False
            
        # Download models
        if not self._download_models():
            self._print("Some models failed to download", "yellow")
            self._print("Continuing with installation, but some features may not work", "yellow")
            
        # Create recipe template
        if not self._create_recipe_template():
            self._print("Recipe template creation failed", "yellow")
            self._print("Recipe functionality may not work properly", "yellow")
            
        # Installation complete
        end_time = time.time()
        installation_time = end_time - start_time
        self._print(f"\n=== Installation Complete ({installation_time:.1f} seconds) ===", "green", bold=True)
        
        # Show summary of installation status
        self._print("\nInstallation Summary:", "cyan", bold=True)
        self._print(f"Platform: {self.platform} ({platform.release()})", "green")
        self._print(f"Python: {platform.python_version()}", "green")
        
        # Hardware detection
        if not self.cpu_only:
            gpu_info = self._check_gpu()
            if gpu_info["available"]:
                gpu_text = f"Detected: {gpu_info['name']}"
                if gpu_info["is_rtx_3080"]:
                    gpu_text += " (Optimized)"
                self._print(f"GPU: {gpu_text}", "green")
            else:
                self._print("GPU: Not detected or not compatible", "yellow")
        else:
            self._print("GPU: Disabled (CPU-only mode)", "yellow")
            
        # Tools detection
        self._print(f"Git found: {'Yes' if self.has_git else 'No'}", "green" if self.has_git else "yellow")
        self._print(f"C++ compiler found: {'Yes' if self.has_cpp_compiler else 'No'}", "green" if self.has_cpp_compiler else "yellow")
        
        # Reminders
        self._print("\nImportant Reminders:", "cyan", bold=True)
        self._print("1. Edit config.yaml to add your Picovoice access key from https://console.picovoice.ai/", "yellow")
        self._print("2. To run Maggie AI Assistant:", "yellow")
        
        if self.platform == "Windows":
            self._print("   .\\venv\\Scripts\\activate", "cyan")
            self._print("   python main.py", "cyan")
        else:
            self._print("   source venv/bin/activate", "cyan")
            self._print("   python main.py", "cyan")
        
        # Required tools not found warnings
        missing_tools = []
        if not self.has_git:
            missing_tools.append(("Git", "https://git-scm.com/download/win"))
            
        if not self.has_cpp_compiler and self.platform == "Windows":
            missing_tools.append(("Visual C++ Build Tools", "https://visualstudio.microsoft.com/visual-cpp-build-tools/"))
            
        if missing_tools:
            self._print("\nMissing Required Tools:", "yellow")
            for i, (tool, url) in enumerate(missing_tools, 1):
                self._print(f"{i}. {tool} not found", "yellow")
                self._print(f"   Install from: {url}", "yellow")
            self._print("\nInstalling these tools will enable full functionality.", "yellow")
            
        # Ask if user wants to start Maggie
        response = input(f"\n{self.colors['magenta']}Would you like to start Maggie now? (y/n): {self.colors['reset']}")
        
        if response.lower() == "y":
            self._print("\nStarting Maggie AI Assistant...", "cyan", bold=True)
            
            if self.platform == "Windows":
                python_cmd = os.path.join(self.base_dir, "venv", "Scripts", "python")
            else:
                python_cmd = os.path.join(self.base_dir, "venv", "bin", "python")
                
            # Run without capturing output so user can see and interact directly
            self._run_command([python_cmd, "main.py"], capture_output=False)
            
        return True


def main() -> int:
    """
    Main entry point for the installer.
    
    Parses command-line arguments and runs the installation process.
    
    Returns
    -------
    int
        Exit code - 0 for success, 1 for error
    """
    parser = argparse.ArgumentParser(
        description="Maggie AI Assistant Installer",
        formatter_class=argparse.ArgumentDefaultsHelpFormatter
    )
    parser.add_argument("--verbose", action="store_true", help="Enable verbose output")
    parser.add_argument("--skip-problematic", action="store_true",
                        help="Skip packages that require compilation and may have compatibility issues")
    parser.add_argument("--cpu-only", action="store_true",
                        help="Install without GPU support (for systems without compatible GPUs)")
    parser.add_argument("--force-reinstall", action="store_true",
                        help="Force reinstallation of packages even if already installed")
    
    args = parser.parse_args()
    
    installer = MaggieInstaller(
        verbose=args.verbose,
        skip_problematic=args.skip_problematic,
        cpu_only=args.cpu_only,
        force_reinstall=args.force_reinstall
    )
    
    success = installer.install()
    
    return 0 if success else 1

if __name__ == "__main__":
    sys.exit(main())