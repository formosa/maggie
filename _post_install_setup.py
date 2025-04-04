#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# _post_install_setup.py
# This script is intended to be run via `poetry run python _post_install_setup.py`
# by the main install_dev.py script AFTER dependencies are installed.

import argparse
import json
import os
import platform
import shutil
import subprocess
import sys
import time
import urllib.request
import zipfile
import tempfile
from pathlib import Path
from typing import Dict, Any, List, Tuple, Optional, Union, Callable
from collections.abc import MutableMapping

# --- Imports for libraries installed by Poetry ---
try:
    # Import specific exceptions for better handling
    from huggingface_hub import (
        snapshot_download,
        hf_hub_download,
        HfHubHTTPError,
        RepositoryNotFoundError,
        EntryNotFoundError
    )
    hf_hub_available = True
except ImportError:
    print("ERROR: huggingface-hub not found. Model downloads will fail.", file=sys.stderr)
    snapshot_download = None
    hf_hub_download = None
    hf_hub_available = False
    class HfHubHTTPError(Exception): pass
    class RepositoryNotFoundError(HfHubHTTPError): pass
    class EntryNotFoundError(HfHubHTTPError): pass
try:
    import yaml
    try: from yaml import CLoader as Loader, CDumper as Dumper
    except ImportError: from yaml import Loader, Dumper
except ImportError:
    print("ERROR: PyYAML not found. Config setup will fail.", file=sys.stderr)
    yaml = None; Loader = None; Dumper = None
try:
    import docx
except ImportError:
    print("ERROR: python-docx not found. Template creation will fail.", file=sys.stderr)
    docx = None
try:
    import psutil
except ImportError:
    psutil = None
    # Warning printed in __init__ or detect method
try:
    import torch
except ImportError as e:
    torch = None
    # Warning printed in __init__ or detect method
except UserWarning as w:
    torch = None
    # Warning printed in __init__ or detect method

if platform.system() == "Windows":
    try:
        import wmi
    except ImportError:
        wmi = None
    except Exception:
        wmi = None
        # Warning printed in __init__ or detect method
else:
    wmi = None


# --- Helper Classes (ColorOutput, ProgressTracker) ---
class ColorOutput:
    """Handles colored terminal output."""
    # (Code remains the same)
    def __init__(self, force_enable: bool = False):
        self.enabled = force_enable or self._supports_color()
        if self.enabled: self.colors = {'reset': '\x1b[0m', 'bold': '\x1b[1m', 'red': '\x1b[91m', 'green': '\x1b[92m', 'yellow': '\x1b[93m', 'blue': '\x1b[94m', 'magenta': '\x1b[95m', 'cyan': '\x1b[96m', 'white': '\x1b[97m'}
        else: self.colors = {color: '' for color in ['reset', 'bold', 'red', 'green', 'yellow', 'blue', 'magenta', 'cyan', 'white']}
    def _supports_color(self) -> bool:
        if platform.system() == 'Windows':
            try:
                import ctypes; kernel32 = ctypes.windll.kernel32; ENABLE_VIRTUAL_TERMINAL_PROCESSING = 0x0004
                stdout = kernel32.GetStdHandle(-11); mode = ctypes.c_ulong()
                if kernel32.GetConsoleMode(stdout, ctypes.byref(mode)):
                    if (mode.value & ENABLE_VIRTUAL_TERMINAL_PROCESSING) == ENABLE_VIRTUAL_TERMINAL_PROCESSING: return True
                    mode.value |= ENABLE_VIRTUAL_TERMINAL_PROCESSING
                    if kernel32.SetConsoleMode(stdout, mode): return True
            except:
                 if int(platform.release()) >= 10: return True
            return False
        if os.environ.get('NO_COLOR'): return False
        if os.environ.get('FORCE_COLOR'): return True
        return hasattr(sys.stdout, 'isatty') and sys.stdout.isatty()
    def print(self, message: str, color: Optional[str] = None, bold: bool = False):
        formatted = message
        if self.enabled:
            if bold and 'bold' in self.colors: formatted = f"{self.colors['bold']}{formatted}"
            if color and color in self.colors: formatted = f"{self.colors[color]}{formatted}"
            if (bold or color) and 'reset' in self.colors: formatted = f"{formatted}{self.colors['reset']}"
        print(formatted)
    def input(self, prompt: str, color: Optional[str] = None, bold: bool = False) -> str:
        formatted = prompt
        if self.enabled:
            if bold and 'bold' in self.colors: formatted = f"{self.colors['bold']}{formatted}"
            if color and color in self.colors: formatted = f"{self.colors[color]}{formatted}"
            if (bold or color) and 'reset' in self.colors: formatted = f"{formatted}{self.colors['reset']}"
        sys.stdout.flush(); return input(formatted)

class ProgressTracker:
    """Tracks and displays installation progress."""
    def __init__(self, color: ColorOutput, total_steps: int = 10, initial_step: int = 0):
        self.color = color; self.total_steps = total_steps; self.current_step = initial_step; self.start_time = time.time()
    def start_step(self, step_name: str):
        self.current_step += 1; elapsed = time.time() - self.start_time
        self.color.print(f"\n[{self.current_step}/{self.total_steps}] {step_name} (Elapsed: {elapsed:.1f}s)", color='cyan', bold=True)
    def complete_step(self, success: bool = True, message: Optional[str] = None):
        if success: status = '✓ Complete'; color = 'green'
        else: status = '✗ Failed'; color = 'red'
        msg = f"  {status}"
        if message: msg += f": {message}"
        self.color.print(msg, color=color)
    def elapsed_time(self) -> float: return time.time() - self.start_time
    # *** display_summary method restored ***
    def display_summary(self, success: bool = True):
        """Displays the final installation summary."""
        elapsed = self.elapsed_time()
        if success: status = 'Installation Completed Successfully'; color = 'green'
        else: status = 'Installation Completed with Errors or Warnings'; color = 'yellow'
        self.color.print(f"\n=== {status} ===", color=color, bold=True)
        self.color.print(f"Total time: {elapsed:.1f} seconds")
    # *** End of restoration ***

# --- Post-Install Setup Class ---

class PostInstallSetup:
    """Performs setup steps after dependencies are installed."""
    def __init__(self, verbose: bool = False, cpu_only: bool = False, skip_models: bool = False):
        self.verbose = verbose
        self.cpu_only = cpu_only
        self.skip_models = skip_models
        self.base_dir = Path(__file__).parent.resolve()
        self.platform_system = platform.system()
        self.color = ColorOutput()
        self.progress = ProgressTracker(self.color, total_steps=7, initial_step=3) # Start at step 4/7
        self.hardware_info = { # Initialize default structure
            'cpu': {'is_ryzen_9_5900x': False, 'model': '', 'cores': 0, 'threads': 0},
            'gpu': {'is_rtx_3080': False, 'model': 'Unknown', 'vram_gb': 0, 'cuda_available': False, 'cuda_version': '', 'cudnn_available': False, 'cudnn_version': ''},
            'memory': {'total_gb': 0, 'available_gb': 0, 'is_32gb': False}
        }
        self.has_git = self._check_git_exists()
        self._setup_hf_token() # Attempt to set HF token from file

    def _setup_hf_token(self):
        """Reads token from _access.token and sets HF_TOKEN environment variable."""
        # (Code remains the same)
        token_file = self.base_dir / "_access.token"
        self.color.print("Checking for Hugging Face token file...", "blue")
        try:
            if token_file.is_file():
                with open(token_file, 'r') as f: token = f.readline().strip()
                if token:
                    os.environ['HF_TOKEN'] = token
                    self.color.print("  Found _access.token, HF_TOKEN environment variable set.", "green")
                    if self.verbose:
                         token_display = token[:4] + "..." + token[-4:] if len(token) > 8 else token[:4] + "..."
                         self.color.print(f"    Token: {token_display}", "cyan")
                else: self.color.print("  Found _access.token but it is empty.", "yellow")
            else:
                self.color.print("  _access.token file not found.", "yellow")
                self.color.print("  Downloads for private/gated models might fail.", "yellow")
                self.color.print("  Consider creating '_access.token' or using 'huggingface-cli login'.", "yellow")
        except Exception as e: self.color.print(f"  Error reading _access.token file: {e}", "red")

    def _check_git_exists(self) -> bool:
        """Checks if git command is available."""
        try:
            subprocess.run(['git', '--version'], check=True, capture_output=True, text=True)
            return True
        except (subprocess.CalledProcessError, FileNotFoundError):
            return False

    def _run_command(self, command: List[str], check: bool = True, shell: bool = False, capture_output: bool = True, cwd: Optional[Union[str, Path]] = None, env: Optional[Dict] = None) -> Tuple[int, str, str]:
        """Runs a shell command."""
        # (Same as in bootstrap script)
        if self.verbose: self.color.print(f"Running command: {' '.join(command)} in {cwd or self.base_dir}", 'cyan')
        try:
            full_env = os.environ.copy()
            if env: full_env.update(env)
            process = subprocess.Popen(
                command if not shell else ' '.join(command),
                stdout=subprocess.PIPE if capture_output else None,
                stderr=subprocess.PIPE if capture_output else None,
                shell=shell, text=True, encoding='utf-8', errors='replace',
                cwd=str(cwd or self.base_dir), env=full_env
            )
            stdout, stderr = process.communicate()
            return_code = process.returncode
            if check and return_code != 0:
                 error_msg = f"Command '{' '.join(command)}' failed with code {return_code}"
                 if stderr: error_msg += f": {stderr.strip()}"
                 if self.verbose: self.color.print(error_msg, 'red')
            return return_code, stdout.strip() if stdout else '', stderr.strip() if stderr else ''
        except FileNotFoundError:
            if self.verbose: self.color.print(f"Error: Command not found: {command[0]}", 'red')
            return -1, '', f"Command not found: {command[0]}"
        except Exception as e:
            if self.verbose: self.color.print(f"Error executing command '{' '.join(command)}': {e}", 'red')
            return -1, '', str(e)

    def _download_file(self, url: str, destination: str, show_progress: bool = True) -> bool:
        """Downloads a file from a URL with improved progress."""
        # (Same as previous corrected version)
        dest_path = Path(destination)
        try:
            self.color.print(f"Downloading {url}", 'blue')
            dest_path.parent.mkdir(parents=True, exist_ok=True)
            req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'})
            with urllib.request.urlopen(req, timeout=60) as response, open(dest_path, 'wb') as out_file:
                content_length_header = response.info().get('Content-Length')
                file_size = int(content_length_header) if content_length_header else 0
                downloaded = 0
                block_size = 8192 * 16
                start_time = time.time()
                if show_progress and file_size > 0:
                    self.color.print(f"Total file size: {file_size / 1024 / 1024:.1f} MB")
                    progress_bar_width = 40; last_percent_reported = -1
                    while True:
                        buffer = response.read(block_size)
                        if not buffer: break
                        downloaded += len(buffer); out_file.write(buffer)
                        percent = int(downloaded * 100 / file_size)
                        if percent > last_percent_reported:
                           last_percent_reported = percent
                           filled_width = int(progress_bar_width * downloaded / file_size)
                           bar = '█' * filled_width + '-' * (progress_bar_width - filled_width)
                           elapsed = time.time() - start_time
                           speed_mbps = (downloaded / (1024 * 1024)) / elapsed if elapsed > 0 else 0
                           print(f"\r  Progress: |{bar}| {percent}% ({downloaded/1024/1024:.1f}/{file_size/1024/1024:.1f} MB) {speed_mbps:.1f} MB/s  ", end="")
                    print()
                else:
                    self.color.print("Downloading (size unknown)...", "blue")
                    while True: buffer = response.read(block_size); out_file.write(buffer)
            self.color.print(f"Download completed: {dest_path}", 'green'); return True
        except urllib.error.URLError as e:
             self.color.print(f"Error downloading file (URL Error: {e.reason}): {url}", 'red')
             if dest_path.exists():
                 try: dest_path.unlink(missing_ok=True)
                 except OSError: pass
             return False
        except Exception as e:
            self.color.print(f"Error downloading file {url}: {e}", 'red')
            if dest_path.exists():
                try: dest_path.unlink(missing_ok=True)
                except OSError: pass
            return False

    def _detect_hardware(self) -> None:
        """Detects CPU, Memory, and GPU hardware using installed libraries."""
        # (Code remains the same)
        self.color.print('Detecting Hardware Configuration...', 'cyan', bold=True)
        self.hardware_info['cpu'] = self._detect_cpu()
        self.hardware_info['memory'] = self._detect_memory()
        self.hardware_info['gpu'] = self._detect_gpu() if not self.cpu_only else {'available': False, 'cuda_available': False}
        self._print_hardware_summary()


    def _detect_cpu(self)->Dict[str,Any]:
        """Detects CPU information."""
        # (Code remains the same)
        cpu_info = {'is_ryzen_9_5900x': False, 'model':'Unknown','cores':0,'threads':0}
        try: cpu_info['model'] = platform.processor() or 'Unknown'
        except Exception: pass
        try:
             if psutil: # Check if import succeeded
                 cpu_info['cores']=psutil.cpu_count(logical=False) or 0
                 cpu_info['threads']=psutil.cpu_count(logical=True) or 0
             else: raise ImportError # Trigger fallback if psutil unavailable
        except (ImportError, Exception):
             if self.verbose: self.color.print("psutil unavailable, using os.cpu_count().", "yellow")
             cpu_info['threads'] = os.cpu_count() or 0
             cpu_info['cores'] = cpu_info['threads'] // 2 if cpu_info['threads'] > 1 else cpu_info['threads']

        if self.platform_system == 'Windows':
            try:
                if wmi: # Check if import succeeded
                    c = wmi.WMI()
                    processor = c.Win32_Processor()[0]
                    cpu_info['model'] = processor.Name.strip()
                    if cpu_info['cores'] == 0 and hasattr(processor, 'NumberOfCores'): cpu_info['cores'] = processor.NumberOfCores
                    if cpu_info['threads'] == 0 and hasattr(processor, 'NumberOfLogicalProcessors'): cpu_info['threads'] = processor.NumberOfLogicalProcessors
                # else: No warning needed if import failed at top
            except Exception as e:
                 if self.verbose: self.color.print(f"WMI CPU detection error: {e}", "yellow")

        model_lower = cpu_info['model'].lower()
        if 'ryzen 9' in model_lower and '5900x' in model_lower: cpu_info['is_ryzen_9_5900x'] = True
        return cpu_info

    def _detect_memory(self)->Dict[str,Any]:
        """Detects memory information."""
        # (Code remains the same)
        memory_info={'total_gb':0,'available_gb':0,'is_32gb':False}
        if not psutil: # Check if import succeeded at top
             self.color.print("psutil library not available. Cannot determine RAM details.", "yellow")
             return memory_info
        try:
            mem=psutil.virtual_memory()
            memory_info['total_gb']=mem.total / (1024**3)
            memory_info['available_gb']=mem.available / (1024**3)
            memory_info['is_32gb'] = memory_info['total_gb'] >= 30.0
        except Exception as e:
             if self.verbose: self.color.print(f"Memory detection error: {e}", "yellow")
        return memory_info

    def _detect_gpu(self)->Dict[str,Any]:
        """Detects GPU information using PyTorch."""
        # (Code remains the same)
        gpu_info = {'available': False, 'is_rtx_3080': False, 'model': 'Unknown','vram_gb': 0, 'cuda_available': False, 'cuda_version': '', 'cudnn_available': False, 'cudnn_version': ''}
        if self.cpu_only: return gpu_info

        if not torch: # Check if torch import succeeded at top
             self.color.print("PyTorch library not available. Cannot detect GPU.", "red")
             return gpu_info

        try:
            gpu_info['cuda_available'] = torch.cuda.is_available()
            if gpu_info['cuda_available']:
                device_count = torch.cuda.device_count()
                if device_count > 0:
                    props = torch.cuda.get_device_properties(0)
                    gpu_info['model'] = props.name
                    gpu_info['available'] = True # Mark as available if CUDA found
                    gpu_info['vram_gb'] = props.total_memory / (1024**3)
                    gpu_info['cuda_version'] = torch.version.cuda
                    gpu_info['cudnn_available'] = torch.backends.cudnn.is_available()
                    if gpu_info['cudnn_available']: gpu_info['cudnn_version'] = str(torch.backends.cudnn.version())
                    if '3080' in gpu_info['model']: gpu_info['is_rtx_3080'] = True
            if not gpu_info['cuda_available']:
                 self.color.print('No CUDA-capable GPU detected by PyTorch.', 'yellow')
                 self.color.print('Ensure NVIDIA drivers, CUDA Toolkit 11.8, and cuDNN 8.9.7 are installed and compatible.', 'yellow')

        except Exception as e:
            self.color.print(f"PyTorch GPU detection error: {e!r}", "red")
            # Fallback check using nvidia-smi
            smi_path = "nvidia-smi"
            if self.platform_system == "Windows":
                program_files = os.environ.get("ProgramFiles", "C:\\Program Files")
                smi_path_win = Path(program_files) / "NVIDIA Corporation" / "NVSMI" / "nvidia-smi.exe"
                if smi_path_win.exists(): smi_path = str(smi_path_win)
            rc_smi, _, _ = self._run_command([smi_path], check=False, capture_output=False)
            if rc_smi == 0: self.color.print('nvidia-smi found. GPU likely present but PyTorch cannot use it.', 'yellow'); gpu_info['available'] = True

        return gpu_info

    def _print_hardware_summary(self):
        """Prints hardware summary previously gathered."""
        # (Code remains the same as previous version)
        self.color.print('Hardware Configuration Summary:', 'cyan', bold=True)
        cpu_info = self.hardware_info['cpu']; mem_info = self.hardware_info['memory']; gpu_info = self.hardware_info['gpu']
        self.color.print(f"  CPU: {cpu_info.get('model', 'Unknown')} ({cpu_info.get('cores', 'N/A')}c / {cpu_info.get('threads', 'N/A')}t)", 'green' if cpu_info.get('is_ryzen_9_5900x') else 'yellow')
        self.color.print(f"  RAM: {mem_info.get('total_gb', 0):.1f} GB", 'green' if mem_info.get('is_32gb') else 'yellow')
        if self.cpu_only: self.color.print('  GPU: CPU-only mode selected', 'yellow')
        elif gpu_info.get('cuda_available'): # Only print details if CUDA was actually found by torch
             self.color.print(f"  GPU: {gpu_info.get('model', 'Unknown')} ✓", 'green')
             self.color.print(f"       {gpu_info.get('vram_gb', 0):.1f} GB VRAM", 'green')
             self.color.print(f"       CUDA {gpu_info.get('cuda_version', 'N/A')} | cuDNN {gpu_info.get('cudnn_version', 'N/A')}", 'green' if gpu_info.get('cudnn_available') else 'yellow')
        elif gpu_info.get('available'): # GPU might exist (nvidia-smi found) but not usable by torch
             self.color.print(f"  GPU: {gpu_info.get('model', 'Unknown')} (Detected, but PyTorch cannot use CUDA)", 'yellow')
        else: # No GPU detected at all
             self.color.print('  GPU: No CUDA-capable GPU detected', 'red')

    def _setup_config(self)->bool:
        """Creates or updates the config.yaml based on hardware detection."""
        # *** MODIFIED: Update LLM model path and type ***
        config_path = self.base_dir / 'config.yaml'; example_path = self.base_dir / 'config.yaml.example'
        if not example_path.exists(): example_path = self.base_dir / 'config.yaml.txt'
        if not example_path.exists(): example_path = self.base_dir / 'config-yaml-example.txt'
        if not config_path.exists() and not example_path.exists(): self.color.print("ERROR: Config/Example file not found.", 'red'); return False

        self.color.print("Setting up configuration file (config.yaml)...", "cyan")

        if not yaml: # Check if import failed at top
             self.color.print("Error: PyYAML library not available. Cannot setup config.", 'red')
             return False

        try:
            config = {}
            # Load existing or copy example
            if config_path.exists():
                self.color.print(f"Loading existing config: {config_path}", "blue")
                try:
                    with open(config_path, 'r', encoding='utf-8') as f: config = yaml.load(f, Loader=Loader)
                    if not isinstance(config, MutableMapping): self.color.print("Warning: Existing config invalid.", "yellow"); config = {}
                except Exception as e: self.color.print(f"Error loading config: {e}.", "red"); config = {}
            elif example_path and example_path.exists():
                self.color.print(f"Creating config from example: {example_path}", "blue")
                try:
                    shutil.copyfile(example_path, config_path)
                    with open(config_path, 'r', encoding='utf-8') as f: config = yaml.load(f, Loader=Loader)
                    if not isinstance(config, MutableMapping): self.color.print("Warning: Example config invalid.", "yellow"); config = {}
                except Exception as e: self.color.print(f"Error copying/loading example: {e}.", "red"); config = {}
            else: self.color.print("No config/example found. Starting fresh.", "yellow"); config = {}

            # Helper to ensure nested keys exist
            def ensure_keys(d, keys):
                current = d
                for key in keys:
                    if not isinstance(current, MutableMapping): return None
                    current = current.setdefault(key, {})
                return current

            # Apply settings
            llm_config = ensure_keys(config, ['llm']); gpu_config = ensure_keys(config, ['gpu']); cpu_config = ensure_keys(config, ['cpu']); mem_config = ensure_keys(config, ['memory']); stt_config = ensure_keys(config, ['stt']); stt_whisper_config = ensure_keys(stt_config, ['whisper']) if stt_config else None; tts_config = ensure_keys(config, ['tts'])

            # *** Use the NEW Phi-4 directory name and type ***
            llm_dir_name = 'microsoft/Phi-4-multimodal-instruct-onnx' # Use repo structure
            if llm_config is not None:
                llm_config.update({
                    'model_path': f'maggie/models/llm/{llm_dir_name}',
                    'model_type': 'phi-msft', # Use specific type for Phi-4
                    'use_autogptq': False # ONNX model likely doesn't use AutoGPTQ
                })
                llm_config.pop('rtx_3080_optimized', None) # Remove potentially irrelevant flags

            # *** Use the NEW Whisper directory name and size ***
            whisper_dir_name = 'faster-whisper-large-v3-turbo-ct2'
            if stt_whisper_config is not None:
                 stt_whisper_config['model_path'] = f'maggie/models/stt/{whisper_dir_name}'
                 stt_whisper_config['model_size'] = 'large-v3-turbo' # Update model size

            if tts_config is not None: tts_config.update({'model_path': 'maggie/models/tts', 'voice_model': 'af_heart'})

            gpu_hw = self.hardware_info.get('gpu', {}); cpu_hw = self.hardware_info.get('cpu', {}); mem_hw = self.hardware_info.get('memory', {})
            if self.cpu_only:
                self.color.print("Applying CPU-only optimizations...", "blue")
                if llm_config: llm_config.update({'gpu_layers': 0, 'gpu_layer_auto_adjust': False})
                if gpu_config: gpu_config.update({'max_percent': 0, 'model_unload_threshold': 0}); gpu_config.pop('rtx_3080_optimized', None)
                if tts_config: tts_config['gpu_acceleration'] = False
                if stt_whisper_config: stt_whisper_config['compute_type'] = 'int8'
            else:
                if gpu_hw.get('cuda_available'):
                    self.color.print("Applying GPU optimizations...", "blue")
                    if llm_config: llm_config.update({'gpu_layer_auto_adjust': False}) # ONNX might not support auto-adjust
                    if tts_config: tts_config.update({'gpu_acceleration': True})
                    if stt_whisper_config: stt_whisper_config['compute_type'] = 'float16' # Default GPU type
                    if gpu_hw.get('is_rtx_3080'):
                        self.color.print("Applying RTX 3080 specific optimizations...", "blue")
                        if gpu_config: gpu_config.update({'max_percent': 90, 'model_unload_threshold': 95, 'rtx_3080_optimized': True})
                        if tts_config: tts_config['gpu_precision'] = 'mixed_float16'
                    else:
                        self.color.print("Applying generic GPU optimizations...", "blue")
                        if gpu_config: gpu_config.update({'max_percent': 85, 'model_unload_threshold': 90}); gpu_config.pop('rtx_3080_optimized', None)
                        if tts_config: tts_config['gpu_precision'] = 'float16'
                else:
                     self.color.print("CUDA not available via PyTorch. Forcing CPU settings.", "yellow")
                     if llm_config: llm_config.update({'gpu_layers': 0, 'gpu_layer_auto_adjust': False}); llm_config.pop('rtx_3080_optimized', None)
                     if gpu_config: gpu_config.update({'max_percent': 0, 'model_unload_threshold': 0}); gpu_config.pop('rtx_3080_optimized', None)
                     if tts_config: tts_config['gpu_acceleration'] = False
                     if stt_whisper_config: stt_whisper_config['compute_type'] = 'int8'
            if cpu_hw.get('is_ryzen_9_5900x'):
                self.color.print("Applying Ryzen 9 5900X optimizations...", "blue")
                if cpu_config: cpu_config.update({'max_threads': 8, 'thread_timeout': 30, 'ryzen_9_5900x_optimized': True})
            else:
                if cpu_config: cpu_config.pop('ryzen_9_5900x_optimized', None)
            if mem_hw.get('is_32gb'):
                self.color.print("Applying 32GB+ RAM optimizations...", "blue")
                if mem_config: mem_config.update({'max_percent': 75, 'model_unload_threshold': 85, 'xpg_d10_memory': True})
            else:
                if mem_config: mem_config.pop('xpg_d10_memory', None)

            # Write config
            with open(config_path, 'w', encoding='utf-8') as f:
                yaml.dump(config, f, Dumper=Dumper, default_flow_style=False, sort_keys=False, allow_unicode=True)
            self.color.print(f"Configuration saved to {config_path} ✓", 'green')
            self.color.print('NOTE: Edit config.yaml to add Picovoice Access Key.', 'yellow')
            return True

        except Exception as e:
            self.color.print(f"Error setting up configuration file: {e}", 'red')
            return False
        # *** End of modification ***

    def _download_hf_snapshot(self, repo_user: str, repo_name: str, target_dir: Optional[Path] = None, parent_dir: Optional[Path] = None, force_download: bool = False, **kwargs) -> Optional[str]:
        """
        Downloads a model snapshot from Hugging Face Hub using snapshot_download.
        Relies on HF_TOKEN environment variable for authentication if needed.
        """
        # *** Integrated user function (modified) ***
        if not snapshot_download:
            self.color.print("ERROR: huggingface-hub not available.", "red")
            return None

        repo_id = f"{repo_user}/{repo_name}"

        # Determine local directory path
        local_dir = None
        if target_dir:
            local_dir = Path(target_dir)
        elif parent_dir:
            # Use user/repo structure for clarity if parent_dir is used
            local_dir = Path(parent_dir) / repo_user / repo_name
        else:
            # Default to subfolder in current dir if neither is specified (less ideal)
            local_dir = Path.cwd() / repo_user / repo_name

        self.color.print(f"Downloading snapshot for {repo_id} to {local_dir}...", 'cyan')

        try:
            local_dir.mkdir(parents=True, exist_ok=True)
            # snapshot_download uses HF_TOKEN env var automatically if set
            # token= parameter removed from here, relies on _setup_hf_token
            downloaded_path = snapshot_download(
                repo_id=repo_id,
                local_dir=str(local_dir),
                force_download=force_download,
                local_dir_use_symlinks=False, # Deprecated but keep for safety
                resume_download=True,
                **kwargs # Pass other args like allow_patterns, ignore_patterns
            )
            self.color.print(f"Snapshot download complete for {repo_id} ✓", 'green')
            # Verify directory is not empty after download
            if not any(local_dir.iterdir()):
                 self.color.print(f"Warning: Downloaded directory {local_dir} appears empty.", "yellow")
                 # Decide if this is an error - snapshot_download should raise if failed
                 # return None
            return downloaded_path # Return the actual path

        except (RepositoryNotFoundError, EntryNotFoundError) as e_nf:
            self.color.print(f"Download failed for {repo_id}: Repository or specific entry not found.", "red")
            self.color.print(f"  -> Check repo ID or login via 'huggingface-cli login' if private: https://huggingface.co/{repo_id}", "yellow")
            return None
        except Exception as e:
            self.color.print(f"An error occurred while downloading snapshot for {repo_id}: {e!r}", "red")
            return None
        # *** End of Integration/Modification ***

    def _download_whisper_model(self)->bool:
        """Downloads the specified faster-whisper model using snapshot download."""
        # *** MODIFIED: Use new repo and _download_hf_snapshot ***
        if self.skip_models: self.color.print('Skipping Whisper model download (--skip-models)', 'yellow'); return True

        repo_user = "deepdml"
        repo_name = "faster-whisper-large-v3-turbo-ct2"
        model_dir_name = "faster-whisper-large-v3-turbo-ct2" # Simpler dir name
        model_dir = self.base_dir / 'maggie/models/stt' / model_dir_name

        # Check if directory exists and prompt user if non-empty
        if model_dir.exists() and any(model_dir.iterdir()):
             self.color.print(f"Whisper model directory already exists: {model_dir}", 'yellow')
             response = self.color.input("  Download again (y) or keep existing (n)? [n]: ", color='magenta')
             if response.lower() != 'y':
                 self.color.print("Keeping existing model files.", 'green')
                 return True
             else:
                 self.color.print("Removing existing directory before download...", 'yellow')
                 shutil.rmtree(model_dir, ignore_errors=True)

        # Use the snapshot download helper
        # Specify typical CTranslate2 model file patterns
        downloaded_path = self._download_hf_snapshot(
            repo_user=repo_user,
            repo_name=repo_name,
            target_dir=model_dir, # Use target_dir directly
            allow_patterns=["*.bin", "*.json", "*.txt", "*.py"], # Common CTranslate2 patterns
            ignore_patterns=["onnx/*", "*.onnx", "*.ot", "*.pb", "tf_*", "flax_*", "*.h5"], # Ignore other formats
        )

        if downloaded_path and Path(downloaded_path).exists():
             # Basic verification: Check for common CTranslate2 files
             essential_files = ['model.bin', 'config.json', 'tokenizer.json', 'vocabulary.txt'] # Or vocab.json? Check repo.
             files_in_dir = {f.name for f in Path(downloaded_path).iterdir()}
             missing = [f for f in essential_files if f not in files_in_dir]
             if not missing:
                 self.color.print(f"{repo_id} model downloaded and verified ✓", 'green')
                 return True
             else:
                 self.color.print(f"Whisper download appears incomplete. Missing: {', '.join(missing)}", 'yellow')
                 return False
        else:
            self.color.print(f"Failed to download or verify {repo_id} model.", 'red')
            return False
        # *** End of modification ***

    def _download_kokoro_onnx_models(self)->bool:
        """Downloads the necessary ONNX model files for kokoro using direct GitHub links."""
        # *** MODIFIED: Use direct download logic from user snippet ***
        if self.skip_models: self.color.print('Skipping kokoro model download (--skip-models)', 'yellow'); return True

        self.color.print("Downloading Kokoro TTS models (from GitHub Releases)...", "cyan")
        model_dir= self.base_dir / 'maggie' / 'models' / 'tts' # Use pathlib
        model_dir.mkdir(parents=True, exist_ok=True)

        # List of models to download with their URLs and minimum expected sizes
        models=[
            {
                'name':'kokoro-v1.0.onnx',
                'url':'https://github.com/thewh1teagle/kokoro-onnx/releases/download/model-files-v1.0/kokoro-v1.0.onnx',
                'min_size':10*1024*1024  # 10 MB minimum size check
            },
            {
                'name':'voices-v1.0.bin',
                'url':'https://github.com/thewh1teagle/kokoro-onnx/releases/download/model-files-v1.0/voices-v1.0.bin',
                'min_size':5*1024*1024   # 5 MB minimum size check
            }
            # Add tokens.txt or others here if they are also in GitHub releases
        ]

        all_successful=True
        for model in models:
            model_path = model_dir / model['name']
            min_size = model['min_size']

            # Check if file already exists and has sufficient size
            if model_path.exists():
                try:
                    file_size=model_path.stat().st_size
                    if file_size>=min_size:
                        self.color.print(f"{model['name']} already exists ({file_size/(1024*1024):.2f} MB) ✓",'green')
                        continue
                    else:
                        self.color.print(f"Existing {model['name']} has incorrect size: {file_size/(1024*1024):.2f} MB. Re-downloading...",'yellow')
                except Exception as e_stat:
                     self.color.print(f"Error checking existing {model['name']}: {e_stat}. Re-downloading...", 'yellow')

            # Download the file using the internal helper
            self.color.print(f"Downloading {model['name']}...",'cyan')
            if self._download_file(model['url'], str(model_path)):
                try:
                    file_size=model_path.stat().st_size
                    if file_size>=min_size:
                        self.color.print(f"{model['name']} download successful ({file_size/(1024*1024):.2f} MB)",'green')
                    else:
                        self.color.print(f"Downloaded file has incorrect size: {file_size/(1024*1024):.2f} MB",'yellow')
                        model_path.unlink(missing_ok=True)
                        all_successful=False
                except FileNotFoundError:
                     self.color.print(f"Verification failed: {model['name']} not found after download attempt.", "red")
                     all_successful=False
                except Exception as e_verify:
                     self.color.print(f"Error verifying {model['name']}: {e_verify}", 'red')
                     all_successful=False
            else:
                self.color.print(f"Failed to download {model['name']}",'red')
                all_successful=False

        # Still attempt espeak-ng-data download from HF Hub
        self.color.print("Downloading espeak-ng-data (required by kokoro)...", 'cyan')
        espeak_dir = model_dir / 'espeak-ng-data'
        espeak_ok = False
        if espeak_dir.exists() and any(espeak_dir.iterdir()):
            self.color.print("espeak-ng-data directory already exists ✓", "green")
            espeak_ok = True
        elif hf_hub_available:
             try:
                 snapshot_download(repo_id="rhasspy/espeak-ng-data", repo_type="dataset", local_dir=str(espeak_dir), resume_download=True)
                 if espeak_dir.is_dir() and any(espeak_dir.iterdir()):
                     self.color.print("espeak-ng-data downloaded successfully ✓", "green")
                     espeak_ok = True
                 else:
                     self.color.print("espeak-ng-data download finished but directory is empty/missing.", "red")
             except (RepositoryNotFoundError, EntryNotFoundError) as e_nf:
                 self.color.print(f"Download failed for espeak-ng-data: Repository not found or private.", "red")
                 self.color.print(f"  -> Check repo ID/type or login via 'huggingface-cli login'.", "yellow")
             except Exception as e:
                 self.color.print(f"Download failed for espeak-ng-data: {e!r}", "red")
        else:
             self.color.print("huggingface-hub library not available, cannot download espeak-ng-data.", "red")

        if not espeak_ok: all_successful = False


        if all_successful:
            self.color.print('Kokoro models downloaded successfully ✓','green')
            return True
        else:
            self.color.print('Some Kokoro models/data failed to download','yellow')
            return False
        # *** End of modification ***


    def _download_llm_model(self)->bool: # Renamed from _download_mistral_model
        """Downloads the specified LLM model using snapshot download."""
        # *** MODIFIED: Use _download_hf_snapshot with Phi-4 repo ***
        if self.skip_models:
            self.color.print('Skipping LLM model download (--skip-models)', 'yellow')
            return True

        repo_user = "microsoft"
        repo_name = "Phi-4-multimodal-instruct-onnx"
        repo_id = f"{repo_user}/{repo_name}" # For messages
        # Define parent directory for models
        parent_dir = self.base_dir / 'maggie/models/llm'
        # Construct the expected final directory path using user/repo structure
        llm_dir = parent_dir / repo_user / repo_name

        # Check if directory exists and prompt user if non-empty
        if llm_dir.exists() and any(llm_dir.iterdir()):
             self.color.print(f"LLM model directory already exists: {llm_dir}", 'yellow')
             response = self.color.input("  Download again (y) or keep existing (n)? [n]: ", color='magenta')
             if response.lower() != 'y':
                 self.color.print("Keeping existing model files.", 'green')
                 return True
             else:
                 self.color.print("Removing existing directory before download...", 'yellow')
                 shutil.rmtree(llm_dir, ignore_errors=True)

        # Use the snapshot download helper
        downloaded_path = self._download_hf_snapshot(
            repo_user=repo_user,
            repo_name=repo_name,
            parent_dir=parent_dir, # Will create microsoft/Phi-4... structure
            # Add ignore patterns if needed, e.g., ignore_patterns=["*.pt", "*.safetensors"]
        )

        if downloaded_path and Path(downloaded_path).exists() and any(Path(downloaded_path).iterdir()):
            # Basic verification: just check if the directory is not empty
            self.color.print(f"{repo_id} model downloaded successfully ✓", 'green')
            return True
        else:
            self.color.print(f"Failed to download or verify {repo_id} model.", 'red')
            return False
        # *** End of modification ***

    def _create_recipe_template(self)->bool:
        """Creates a default recipe template if it doesn't exist using python-docx."""
        # (Code remains the same)
        template_dir = self.base_dir / 'maggie/templates'
        template_path = template_dir / 'recipe_template.docx'
        template_dir.mkdir(parents=True, exist_ok=True)

        if template_path.exists():
            self.color.print('Recipe template already exists ✓', 'green')
            return True

        self.color.print("Creating default recipe template...", "cyan")
        if not docx: # Check if import failed at top
            self.color.print("Error: python-docx library not available. Cannot create template.", 'red')
            return False
        try:
            doc = docx.Document(); doc.add_heading("Recipe Name", level=1); doc.add_paragraph()
            doc.add_heading("Recipe Information", level=2); info_table = doc.add_table(rows=3, cols=2); info_table.style = 'Table Grid'
            info_table.cell(0, 0).text = "Preparation Time"; info_table.cell(0, 1).text = "00 minutes"
            info_table.cell(1, 0).text = "Cooking Time"; info_table.cell(1, 1).text = "00 minutes"
            info_table.cell(2, 0).text = "Servings"; info_table.cell(2, 1).text = "0 servings"; doc.add_paragraph()
            doc.add_heading("Ingredients", level=2); doc.add_paragraph("• Ingredient 1", style='List Bullet'); doc.add_paragraph("• Ingredient 2", style='List Bullet'); doc.add_paragraph("• Ingredient 3", style='List Bullet'); doc.add_paragraph()
            doc.add_heading("Instructions", level=2); doc.add_paragraph("Step 1", style='List Number'); doc.add_paragraph("Step 2", style='List Number'); doc.add_paragraph("Step 3", style='List Number'); doc.add_paragraph()
            doc.add_heading("Notes", level=2); doc.add_paragraph("Add any additional notes, tips, or variations here.")
            doc.save(template_path)
            self.color.print(f"Template saved to {template_path} ✓", 'green')
            return True
        except Exception as e:
            self.color.print(f"Error creating docx template: {e}", 'red')
            if template_path.exists():
                try: os.remove(template_path)
                except: pass
            return False

    def run_all_steps(self):
        """Runs all post-install steps."""
        # (Code remains the same)
        self.progress.start_step('Detecting Hardware')
        self._detect_hardware()
        self.progress.complete_step(True) # Assume detection itself worked for progress

        self.progress.start_step('Setting up configuration file')
        if not self._setup_config():
            self.progress.complete_step(False, 'Failed to set up configuration')
            return False # Stop if config fails
        self.progress.complete_step(True)

        self.progress.start_step('Downloading models')
        models_ok = True
        # *** REMOVED CALL to _download_af_heart_model ***
        if not self._download_kokoro_onnx_models():
             self.color.print('Warning: Failed kokoro download.', 'yellow') # Updated message slightly
             models_ok = False # Decide if this should be fatal
        if not self._download_whisper_model():
             self.color.print('Warning: Failed Whisper download.', 'yellow')
             models_ok = False # Decide if this should be fatal
        # *** Calls the renamed _download_llm_model ***
        if not self._download_llm_model():
             self.color.print('Warning: Failed LLM download.', 'yellow')
             models_ok = False # Decide if this should be fatal

        self.progress.complete_step(models_ok)

        self.progress.start_step('Setting up templates & completing installation')
        template_ok = self._create_recipe_template()
        if not template_ok: self.color.print('Warning: Failed recipe template creation.', 'yellow')

        # *** Use self.progress.display_summary (fixed) ***
        self.progress.display_summary(models_ok) # Summary reflects model download success
        self.color.print('\n--- Important Notes ---', 'cyan', bold=True)
        self.color.print('1. Dependencies installed via Poetry.', 'green')
        self.color.print(f"2. Non-Python Requirements: CUDA 11.8, cuDNN 8.9.7 (for GPU).", 'yellow')
        self.color.print(f"   (Detected CUDA: {self.hardware_info['gpu'].get('cuda_version', 'N/A')}, cuDNN: {self.hardware_info['gpu'].get('cudnn_version', 'N/A')})", 'yellow')
        self.color.print('3. Edit config.yaml for Picovoice Access Key.', 'yellow')
        self.color.print('4. Ensure Kokoro voices are available via the installed package.', 'yellow')
        if not self.has_git: self.color.print('5. Git not installed - model updates affected.', 'yellow')
        # C++ compiler check result isn't stored here, but was checked in bootstrap
        # if not self.has_cpp_compiler: self.color.print('6. C++ Compiler not found - building packages may fail.', 'yellow')
        self.color.print('\n--- To start Maggie AI Assistant ---', 'cyan', bold=True)
        self.color.print('   Run: poetry run python main.py', 'green')
        self.progress.complete_step(True) # Mark final step complete
        return models_ok


# --- Main Execution (for _post_install_setup.py) ---

def post_install_main() -> int:
    """Parses arguments passed from bootstrap and runs post-install steps."""
    parser = argparse.ArgumentParser(description='Maggie AI Assistant Post-Install Setup')
    # Add arguments that need to be passed from install_dev.py
    parser.add_argument('--verbose', '-v', action='store_true', help='Enable verbose output')
    parser.add_argument('--cpu-only', action='store_true', help='CPU-only mode was selected')
    parser.add_argument('--skip-models', action='store_true', help='Skip downloading AI models')
    args = parser.parse_args()

    setup_handler = PostInstallSetup(
        verbose=args.verbose,
        cpu_only=args.cpu_only,
        skip_models=args.skip_models
    )

    try:
        # Check if essential libraries were imported successfully at the top
        if not yaml or not docx:
             setup_handler.color.print("Critical Error: Essential libraries (PyYAML or python-docx) missing.", "red", bold=True)
             return 1
        # huggingface_hub check happens within download methods

        success = setup_handler.run_all_steps()
        return 0 if success else 1 # Return 0 on success, 1 on failure
    except KeyboardInterrupt:
        print('\n\nPost-installation setup cancelled by user.')
        return 1
    except Exception as e:
        setup_handler.color.print(f"\n\nAN UNEXPECTED ERROR OCCURRED during post-install setup:", 'red', bold=True)
        setup_handler.color.print(f"Error Type: {type(e).__name__}", 'red')
        setup_handler.color.print(f"Error Details: {e}", 'red')
        if args.verbose:
            import traceback
            setup_handler.color.print("Traceback:", 'red')
            traceback.print_exc()
        return 1

if __name__ == '__main__':
    # This script should only be run via `poetry run python _post_install_setup.py`
    sys.exit(post_install_main())
